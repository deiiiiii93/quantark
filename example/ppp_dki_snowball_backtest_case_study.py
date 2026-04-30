"""
PPP/NPP Snowball delta-hedging case study.

The script compares PPP-DKI, NPP-DKI, and PPP-EKI Parachute Snowballs on
CSI1000 (000852.SH), hedged with CFFEX IM futures. It can fetch/cache AKShare
history, generate deterministic stress paths, solve per-product fair coupons,
run daily OTC backtests, and export CSV/XLSX/DOCX/HTML/chart artifacts.

Usage:
    python example/ppp_dki_snowball_backtest_case_study.py
    python example/ppp_dki_snowball_backtest_case_study.py --refresh-data
    python example/ppp_dki_snowball_backtest_case_study.py --synthetic-only --scenario-days 60
"""

from __future__ import annotations

import argparse
import html as html_lib
import json
import math
import os
import sys
import warnings
from calendar import FRIDAY, monthcalendar
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Callable, Optional

import numpy as np
import pandas as pd

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from asset.equity.param import MCParams, PDEParams, QuadParams
from asset.equity.product.option import (
    AccrualConfig,
    BarrierConfig,
    PayoffConfig,
    SnowballOption,
    create_parachute_snowball,
)
from backtest.otc import (
    AutocallableBacktestConfig,
    AutocallableBacktestEngine,
    AutocallableDeltaHedgeStrategy,
    AutocallableEngineConfig,
    AutocallableMarketDataSet,
    FuturesRollPolicy,
    SignedDividendYield,
    SurfaceGridConfig,
)
from backtest.otc.engine_factory import create_pricing_engine
from param import FlatRateCurve, FlatVolSurface, SpotQuote
from priceenv import PricingEnvironment
from util.enum import CouponPayType, ObservationType, ProtectionType
from util.enum.engine_enums import EngineType, PDEMethod
from util.exceptions import ValidationError


PROJECT_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_OUTPUT_ROOT = PROJECT_ROOT / "output" / "case_study"
UNDERLYING_SYMBOL = "000852"
UNDERLYING_NAME = "CSI1000"
FUTURES_PREFIX = "IM"
FUTURES_MULTIPLIER = 200.0
DEFAULT_NOTIONAL = 10_000_000.0


class CaseStudyError(RuntimeError):
    """Raised when the case-study runner cannot complete."""


@dataclass(frozen=True)
class CachedMarketFrames:
    """Normalized market frames persisted by the case-study cache."""

    spot_data: pd.DataFrame
    futures_data: pd.DataFrame


@dataclass(frozen=True)
class CaseStudyTerms:
    """Economic terms shared by all structures."""

    notional: float = DEFAULT_NOTIONAL
    ko_ratio: float = 1.03
    ki_ratio: float = 0.75
    ppp_protection_rate: float = 0.25
    rate: float = 0.02
    dividend_yield: float = 0.08
    volatility: float = 0.25


@dataclass(frozen=True)
class FairCouponResult:
    """Fair-coupon calibration audit record."""

    coupon: float
    pv: float
    lower: float
    upper: float
    pv_lower: float
    pv_upper: float
    iterations: int
    solved: bool


def parse_args(argv: Optional[list[str]] = None) -> argparse.Namespace:
    today = datetime.now().strftime("%Y%m%d")
    parser = argparse.ArgumentParser(
        description="Run the PPP/NPP Snowball backtest case study."
    )
    parser.add_argument("--output-dir", type=Path, default=None)
    parser.add_argument("--refresh-data", action="store_true")
    parser.add_argument("--cache-only", action="store_true")
    parser.add_argument("--synthetic-only", action="store_true")
    parser.add_argument("--scenario-days", type=int, default=0)
    parser.add_argument("--history-years", type=float, default=3.0)
    parser.add_argument("--start-date", default=None)
    parser.add_argument("--end-date", default=None)
    parser.add_argument("--engine", choices=("quad", "pde", "mc"), default="quad")
    parser.add_argument("--rate", type=float, default=0.02)
    parser.add_argument("--dividend-yield", type=float, default=0.08)
    parser.add_argument("--volatility", type=float, default=0.25)
    parser.add_argument("--notional", type=float, default=DEFAULT_NOTIONAL)
    parser.add_argument("--ko", type=float, default=1.03)
    parser.add_argument("--ki", type=float, default=0.75)
    parser.add_argument("--ppp-protection", type=float, default=0.25)
    parser.add_argument("--quad-grid", type=int, default=1001)
    parser.add_argument("--quad-std-devs", type=float, default=6.0)
    parser.add_argument("--pde-grid", type=int, default=140)
    parser.add_argument("--pde-steps", type=int, default=140)
    parser.add_argument("--mc-paths", type=int, default=2000)
    parser.add_argument("--mc-steps", type=int, default=252)
    parser.add_argument("--mc-seed", type=int, default=42)
    parser.add_argument("--event-probabilities", action="store_true")
    parser.add_argument("--surfaces", action="store_true")
    args = parser.parse_args(argv)
    if args.output_dir is None:
        args.output_dir = DEFAULT_OUTPUT_ROOT / f"ppp_dki_snowball_backtest_{today}"
    return args


def load_akshare():
    try:
        import akshare as ak  # type: ignore
    except ImportError as exc:
        raise CaseStudyError(
            "AKShare is not installed. Install akshare or rerun with "
            "--synthetic-only / an existing cache."
        ) from exc
    return ak


def _pick_column(df: pd.DataFrame, candidates: list[str], label: str) -> str:
    for column in candidates:
        if column in df.columns:
            return column
    raise CaseStudyError(f"Missing {label} column. Available columns: {list(df.columns)}")


def normalize_index_spot(raw: pd.DataFrame) -> pd.DataFrame:
    date_col = _pick_column(raw, ["date", "日期"], "index date")
    close_col = _pick_column(raw, ["close", "收盘", "收盘价"], "index close")
    out = pd.DataFrame(
        {
            "date": pd.to_datetime(raw[date_col]).dt.normalize(),
            "spot": pd.to_numeric(raw[close_col], errors="coerce"),
        }
    )
    out = out.dropna().sort_values("date").drop_duplicates("date", keep="last")
    if out.empty:
        raise CaseStudyError("Index data is empty after normalization")
    return out.reset_index(drop=True)


def fetch_csi1000_spot(ak) -> pd.DataFrame:
    try:
        raw = ak.stock_zh_index_daily(symbol=f"sh{UNDERLYING_SYMBOL}")
        return normalize_index_spot(raw)
    except Exception:
        raw = ak.index_zh_a_hist(symbol=UNDERLYING_SYMBOL, period="daily")
        return normalize_index_spot(raw)


def third_friday(year: int, month: int) -> pd.Timestamp:
    fridays = [week[FRIDAY] for week in monthcalendar(year, month) if week[FRIDAY] != 0]
    return pd.Timestamp(year=year, month=month, day=fridays[2])


def add_months(date: pd.Timestamp, months: int) -> pd.Timestamp:
    month_index = date.month - 1 + months
    year = date.year + month_index // 12
    month = month_index % 12 + 1
    return pd.Timestamp(year=year, month=month, day=1)


def contract_expiry(contract: str) -> pd.Timestamp:
    suffix = str(contract).replace(FUTURES_PREFIX, "")
    if len(suffix) != 4 or not suffix.isdigit():
        raise CaseStudyError(f"Cannot infer expiry from futures contract {contract!r}")
    year = 2000 + int(suffix[:2])
    month = int(suffix[2:])
    return third_friday(year, month)


def futures_contract_symbols(
    start_date: pd.Timestamp, end_date: pd.Timestamp, extra_months: int = 3
) -> list[str]:
    cursor = pd.Timestamp(start_date).normalize().replace(day=1)
    end = add_months(pd.Timestamp(end_date).normalize().replace(day=1), extra_months)
    symbols = []
    while cursor <= end:
        symbols.append(f"{FUTURES_PREFIX}{cursor:%y%m}")
        cursor = add_months(cursor, 1)
    return symbols


def normalize_im_futures(raw: pd.DataFrame, contract: str) -> pd.DataFrame:
    date_col = _pick_column(raw, ["date", "日期"], "futures date")
    close_col = _pick_column(raw, ["close", "收盘", "收盘价"], "futures close")
    out = pd.DataFrame(
        {
            "date": pd.to_datetime(raw[date_col]).dt.normalize(),
            "contract": contract,
            "futures_price": pd.to_numeric(raw[close_col], errors="coerce"),
            "expiry_date": contract_expiry(contract),
            "multiplier": FUTURES_MULTIPLIER,
        }
    )
    out = out.dropna().sort_values("date").drop_duplicates(
        ["date", "contract"], keep="last"
    )
    return out[out["futures_price"] > 0].reset_index(drop=True)


def fetch_im_futures(
    ak, start_date: pd.Timestamp, end_date: pd.Timestamp
) -> pd.DataFrame:
    frames: list[pd.DataFrame] = []
    for contract in futures_contract_symbols(start_date, end_date):
        try:
            raw = ak.futures_zh_daily_sina(symbol=contract)
            frame = normalize_im_futures(raw, contract)
        except Exception as exc:
            print(f"[warn] Skipping {contract}: {exc}")
            continue
        mask = (frame["date"] >= start_date) & (frame["date"] <= end_date)
        if mask.any():
            frames.append(frame.loc[mask])
    if not frames:
        raise CaseStudyError("No IM futures data was fetched for the requested window")
    return pd.concat(frames, ignore_index=True).sort_values(
        ["date", "expiry_date", "contract"]
    )


def cache_paths(cache_dir: Path) -> dict[str, Path]:
    return {
        "spot": cache_dir / "csi1000_spot.csv",
        "futures": cache_dir / "im_futures.csv",
    }


def read_cache(cache_dir: Path) -> CachedMarketFrames:
    paths = cache_paths(cache_dir)
    if not paths["spot"].exists() or not paths["futures"].exists():
        raise CaseStudyError(f"Cache is missing under {cache_dir}")
    return CachedMarketFrames(
        spot_data=pd.read_csv(paths["spot"], parse_dates=["date"]),
        futures_data=pd.read_csv(paths["futures"], parse_dates=["date", "expiry_date"]),
    )


def write_cache(cache_dir: Path, frames: CachedMarketFrames) -> None:
    cache_dir.mkdir(parents=True, exist_ok=True)
    paths = cache_paths(cache_dir)
    frames.spot_data.to_csv(paths["spot"], index=False)
    frames.futures_data.to_csv(paths["futures"], index=False)


def latest_window(
    spot_data: pd.DataFrame,
    *,
    history_years: float,
    start_date: Optional[str],
    end_date: Optional[str],
) -> tuple[pd.Timestamp, pd.Timestamp]:
    spot = spot_data.sort_values("date")
    end = (
        pd.Timestamp(end_date).normalize()
        if end_date is not None
        else pd.Timestamp(spot["date"].max()).normalize()
    )
    start = (
        pd.Timestamp(start_date).normalize()
        if start_date is not None
        else end - pd.DateOffset(days=int(round(history_years * 365)))
    )
    if end < start:
        raise CaseStudyError("end-date must be on or after start-date")
    available = spot[(spot["date"] >= start) & (spot["date"] <= end)]
    if available.empty:
        raise CaseStudyError("No spot data in requested history window")
    return (
        pd.Timestamp(available["date"].min()).normalize(),
        pd.Timestamp(available["date"].max()).normalize(),
    )


def load_market_cache_or_fetch(
    *,
    cache_dir: Path,
    refresh_data: bool,
    cache_only: bool,
    synthetic_only: bool,
    history_years: float,
    start_date: Optional[str],
    end_date: Optional[str],
    terms: CaseStudyTerms,
    scenario_days: int,
) -> CachedMarketFrames:
    if synthetic_only:
        dates = pd.bdate_range("2024-01-02", periods=max(int(scenario_days), 12))
        spot = pd.DataFrame(
            {
                "date": dates,
                "spot": 100.0 * (1.0 + 0.02 * np.sin(np.linspace(0, 3, len(dates)))),
            }
        )
        futures = synthetic_futures_chain(spot, terms.rate, terms.dividend_yield)
        return CachedMarketFrames(spot_data=spot, futures_data=futures)

    if not refresh_data:
        try:
            cached = read_cache(cache_dir)
            latest_window(
                cached.spot_data,
                history_years=history_years,
                start_date=start_date,
                end_date=end_date,
            )
            return cached
        except CaseStudyError:
            if cache_only:
                raise

    if cache_only:
        raise CaseStudyError(f"Cache is missing under {cache_dir}")

    try:
        ak = load_akshare()
        spot = fetch_csi1000_spot(ak)
        start, end = latest_window(
            spot,
            history_years=history_years,
            start_date=start_date,
            end_date=end_date,
        )
        spot = restrict_window(spot, start, end)
        try:
            futures = fetch_im_futures(ak, start, end)
        except Exception as exc:
            print(f"[warn] Live IM futures fetch failed; using synthetic futures. Reason: {exc}")
            futures = synthetic_futures_chain(spot, terms.rate, terms.dividend_yield)
        frames = CachedMarketFrames(spot_data=spot, futures_data=futures)
        write_cache(cache_dir, frames)
        return frames
    except Exception as exc:
        try:
            cached = read_cache(cache_dir)
            print(f"[warn] Live AKShare fetch failed; using cache. Reason: {exc}")
            return cached
        except CaseStudyError:
            raise CaseStudyError(
                "Could not load AKShare data and no usable cache exists. "
                f"Original error: {exc}"
            ) from exc


def restrict_window(
    df: pd.DataFrame, start_date: pd.Timestamp, end_date: pd.Timestamp
) -> pd.DataFrame:
    out = df.copy()
    out["date"] = pd.to_datetime(out["date"]).dt.normalize()
    mask = (out["date"] >= start_date) & (out["date"] <= end_date)
    return out.loc[mask].reset_index(drop=True)


def constant_vol_data(dates: pd.Series, volatility: float) -> pd.DataFrame:
    return pd.DataFrame(
        {
            "date": pd.to_datetime(dates).dt.normalize(),
            "volatility": float(volatility),
        }
    )


def rate_data(dates: pd.Series, rate: float) -> pd.DataFrame:
    return pd.DataFrame({"date": pd.to_datetime(dates).dt.normalize(), "rate": rate})


def synthetic_futures_chain(
    spot_data: pd.DataFrame,
    rate: float,
    dividend_yield: float,
    contracts_ahead: int = 4,
) -> pd.DataFrame:
    rows: list[dict[str, Any]] = []
    spot = spot_data.sort_values("date").copy()
    for _, row in spot.iterrows():
        date = pd.Timestamp(row["date"]).normalize()
        base_month = date.replace(day=1)
        added = 0
        month_offset = 0
        while added < contracts_ahead:
            month = add_months(base_month, month_offset)
            expiry = third_friday(month.year, month.month)
            month_offset += 1
            if expiry <= date:
                continue
            ttm = max((expiry - date).days / 365.0, 1.0 / 365.0)
            price = float(row["spot"]) * math.exp((rate - dividend_yield) * ttm)
            rows.append(
                {
                    "date": date,
                    "contract": f"{FUTURES_PREFIX}{month:%y%m}",
                    "futures_price": price,
                    "expiry_date": expiry,
                    "multiplier": FUTURES_MULTIPLIER,
                }
            )
            added += 1
    return pd.DataFrame(rows).sort_values(["date", "expiry_date", "contract"])


def build_market_dataset(
    spot_data: pd.DataFrame,
    futures_data: pd.DataFrame,
    terms: CaseStudyTerms,
) -> AutocallableMarketDataSet:
    spot = spot_data.sort_values("date").reset_index(drop=True)
    return AutocallableMarketDataSet.from_dataframes(
        spot_data=spot,
        vol_data=constant_vol_data(spot["date"], terms.volatility),
        rate_data=rate_data(spot["date"], terms.rate),
        futures_data=futures_data,
        metadata={
            "spot_symbol": f"{UNDERLYING_SYMBOL}.SH",
            "futures_prefix": FUTURES_PREFIX,
            "vol_source": "flat",
            "rate_source": "flat",
            "pricing_dividend_yield": terms.dividend_yield,
        },
    )


def generate_scenario_spot_data(base_spot: pd.DataFrame) -> dict[str, pd.DataFrame]:
    spot = base_spot.sort_values("date").reset_index(drop=True)
    dates = spot["date"]
    s0 = float(spot["spot"].iloc[0])
    n = len(spot)
    t = np.linspace(0.0, 1.0, n)

    rapid_up = s0 * np.where(t < 0.25, 1.0 + 1.0 * t, 1.25)
    rapid_down = s0 * np.where(t < 0.25, 1.0 - 1.4 * t, 0.65)
    high_osc = s0 * (
        1.0
        + 0.14 * np.sin(10.0 * math.pi * t)
        + 0.05 * np.sin(37.0 * math.pi * t)
    )
    low_osc = s0 * (1.0 + 0.035 * np.sin(8.0 * math.pi * t))

    scenarios = {
        "historical": spot["spot"].to_numpy(dtype=float),
        "rapid_up": rapid_up,
        "rapid_down": rapid_down,
        "high_oscillation": np.clip(high_osc, 0.55 * s0, 1.35 * s0),
        "low_oscillation": np.clip(low_osc, 0.90 * s0, 1.10 * s0),
    }
    return {
        name: pd.DataFrame({"date": dates, "spot": values})
        for name, values in scenarios.items()
    }


def ko_observation_times(tenor: float, count: int = 36) -> list[float]:
    return [float((i + 1) / count * tenor) for i in range(count)]


def daily_ki_times(dates: pd.Series, start_date: pd.Timestamp) -> list[float]:
    start = pd.Timestamp(start_date).normalize()
    times = [
        max((pd.Timestamp(date).normalize() - start).days / 365.0, 0.0)
        for date in dates
    ]
    return [float(t) for t in times[1:] if t > 0.0]


def _base_payoff_config(
    *,
    protection_type: ProtectionType,
    protection_rate: float,
    coupon: float,
) -> PayoffConfig:
    return PayoffConfig(
        rebate_rate=float(coupon),
        include_principal=False,
        participation_rate=1.0,
        protection_type=protection_type,
        protection_rate=float(protection_rate),
    )


def build_case_study_products(
    *,
    initial_spot: float,
    issue_date: pd.Timestamp,
    dates: pd.Series,
    terms: CaseStudyTerms,
    coupons: Optional[dict[str, float]] = None,
) -> dict[str, SnowballOption]:
    coupons = coupons or {}
    maturity = max(
        (pd.Timestamp(dates.iloc[-1]).normalize() - pd.Timestamp(issue_date).normalize()).days
        / 365.0,
        1.0 / 365.0,
    )
    multiplier = float(terms.notional) / float(initial_spot)
    ko_barrier = float(initial_spot) * float(terms.ko_ratio)
    ki_barrier = float(initial_spot) * float(terms.ki_ratio)
    ko_times = ko_observation_times(maturity)
    dki_times = daily_ki_times(dates, issue_date)

    def make_dki(label: str, protection_type: ProtectionType) -> SnowballOption:
        product = SnowballOption(
            initial_price=float(initial_spot),
            strike=float(initial_spot),
            maturity=maturity,
            initial_date=pd.Timestamp(issue_date).to_pydatetime(),
            contract_multiplier=multiplier,
            barrier_config=BarrierConfig(
                ko_barrier=ko_barrier,
                ko_rate=float(coupons.get(label, 0.0)),
                ko_observation_type=ObservationType.DISCRETE,
                ko_observation_dates=ko_times,
                ki_barrier=ki_barrier,
                ki_observation_type=ObservationType.DISCRETE,
                ki_observation_dates=dki_times,
                ki_continuous=False,
            ),
            payoff_config=_base_payoff_config(
                protection_type=protection_type,
                protection_rate=terms.ppp_protection_rate,
                coupon=float(coupons.get(label, 0.0)),
            ),
            accrual_config=AccrualConfig(
                coupon_pay_type=CouponPayType.INSTANT,
                is_annualized=True,
            ),
        )
        return product

    ppp_eki_parachute = create_parachute_snowball(
        initial_price=float(initial_spot),
        strike=float(initial_spot),
        maturity=maturity,
        contract_multiplier=multiplier,
        ko_barrier=ko_barrier,
        ko_rate=float(coupons.get("PPP-EKI-Parachute", 0.0)),
        ki_barrier=ki_barrier,
        num_observations=len(ko_times),
        ko_observation_dates=ko_times,
        ki_observation_type=ObservationType.DISCRETE,
        ki_observation_dates=[maturity],
        ki_continuous=False,
        include_principal=False,
        participation_rate=1.0,
        protection_type=ProtectionType.PARTIAL,
        protection_rate=terms.ppp_protection_rate,
        rebate_rate=float(coupons.get("PPP-EKI-Parachute", 0.0)),
    )
    ppp_eki_parachute.initial_date = pd.Timestamp(issue_date).to_pydatetime()

    return {
        "PPP-DKI": make_dki("PPP-DKI", ProtectionType.PARTIAL),
        "NPP-DKI": make_dki("NPP-DKI", ProtectionType.NONE),
        "PPP-EKI-Parachute": ppp_eki_parachute,
    }


def create_engine_config(args: argparse.Namespace) -> AutocallableEngineConfig:
    quad_params = QuadParams(
        grid_points=int(args.quad_grid),
        num_std_devs=float(args.quad_std_devs),
    )
    if args.engine == "quad":
        return AutocallableEngineConfig(
            pricing_engine_type=EngineType.QUADRATURE,
            quad_params=quad_params,
        )
    if args.engine == "pde":
        return AutocallableEngineConfig(
            pricing_engine_type=EngineType.PDE,
            method=EngineType.PDE(PDEMethod.CRANK_NICOLSON),
            pde_params=PDEParams(
                grid_size=int(args.pde_grid),
                time_steps=int(args.pde_steps),
                max_time_steps=max(int(args.pde_steps), 200),
            ),
            quad_params=quad_params,
            surface_engine_type=EngineType.QUADRATURE,
        )
    return AutocallableEngineConfig(
        pricing_engine_type=EngineType.MONTE_CARLO,
        mc_params=MCParams(
            num_paths=int(args.mc_paths),
            time_steps=int(args.mc_steps),
            seed=int(args.mc_seed),
        ),
        quad_params=quad_params,
        surface_engine_type=EngineType.QUADRATURE,
    )


def validate_quad_grid_for_case(args: argparse.Namespace, base_spot: pd.DataFrame) -> None:
    """Guard production 3Y autocallable runs against under-resolved quadrature grids."""
    if args.engine != "quad":
        return
    if base_spot.empty:
        return
    start = pd.Timestamp(base_spot["date"].iloc[0]).normalize()
    end = pd.Timestamp(base_spot["date"].iloc[-1]).normalize()
    tenor_years = max((end - start).days / 365.0, 0.0)
    if tenor_years >= 2.5 and int(args.quad_grid) < 1001:
        raise CaseStudyError(
            "quad-grid must be at least 1001 for 3Y autocallables. "
            f"Received --quad-grid {args.quad_grid}; rerun with --quad-grid 1001 "
            "or above, or use --engine pde/mc for a separate approximation check."
        )


def pricing_environment(
    spot: float, valuation_date: pd.Timestamp, terms: CaseStudyTerms
) -> PricingEnvironment:
    return PricingEnvironment(
        spot_quote=SpotQuote(spot=float(spot), asset_name=UNDERLYING_NAME),
        vol_surface=FlatVolSurface(volatility=float(terms.volatility)),
        rate_curve=FlatRateCurve(rate=float(terms.rate)),
        div_yield=SignedDividendYield(float(terms.dividend_yield)),
        valuation_date=pd.Timestamp(valuation_date).to_pydatetime(),
    )


def solve_fair_coupon(
    *,
    product_builder: Callable[[float], SnowballOption],
    engine_config: AutocallableEngineConfig,
    env: PricingEnvironment,
    lower: float = 0.0,
    upper: float = 0.80,
    tolerance: float = 1e-4,
    max_iterations: int = 40,
) -> float:
    return solve_fair_coupon_result(
        product_builder=product_builder,
        engine_config=engine_config,
        env=env,
        lower=lower,
        upper=upper,
        tolerance=tolerance,
        max_iterations=max_iterations,
    ).coupon


def solve_fair_coupon_result(
    *,
    product_builder: Callable[[float], SnowballOption],
    engine_config: AutocallableEngineConfig,
    env: PricingEnvironment,
    lower: float = 0.0,
    upper: float = 0.80,
    tolerance: float = 1e-4,
    max_iterations: int = 40,
) -> FairCouponResult:
    def value(coupon: float) -> float:
        product = product_builder(float(coupon))
        engine = create_pricing_engine(product, engine_config)
        pv = float(engine.price(product, env))
        if not math.isfinite(pv):
            raise CaseStudyError(
                f"Fair-coupon calibration produced non-finite PV at coupon={coupon:.8f}. "
                "Change the pricing engine or example terms instead of accepting a zero coupon."
            )
        return pv

    low = float(lower)
    high = float(upper)
    f_low = value(low)
    f_high = value(high)
    while f_low * f_high > 0.0 and high < 5.0:
        high *= 1.5
        f_high = value(high)

    if f_low * f_high > 0.0:
        raise CaseStudyError(
            "Fair coupon is not bracketed under the current product terms: "
            f"PV({low:.6f})={f_low:,.6f}, PV({high:.6f})={f_high:,.6f}. "
            "Adjust KO/KI/protection terms or use a different pricing engine; "
            "the runner will not silently return a zero or boundary coupon."
        )

    if abs(f_low) <= tolerance:
        return FairCouponResult(
            coupon=low,
            pv=f_low,
            lower=low,
            upper=high,
            pv_lower=f_low,
            pv_upper=f_high,
            iterations=0,
            solved=True,
        )
    if abs(f_high) <= tolerance:
        return FairCouponResult(
            coupon=high,
            pv=f_high,
            lower=low,
            upper=high,
            pv_lower=f_low,
            pv_upper=f_high,
            iterations=0,
            solved=True,
        )

    if not math.isclose(f_high, f_low, rel_tol=0.0, abs_tol=1e-14):
        linear_coupon = low - f_low * (high - low) / (f_high - f_low)
        bracket_low = min(low, high)
        bracket_high = max(low, high)
        if bracket_low <= linear_coupon <= bracket_high:
            f_linear = value(linear_coupon)
            if abs(f_linear) <= tolerance:
                return FairCouponResult(
                    coupon=float(linear_coupon),
                    pv=f_linear,
                    lower=low,
                    upper=high,
                    pv_lower=f_low,
                    pv_upper=f_high,
                    iterations=1,
                    solved=True,
                )

    iterations = 0
    for _ in range(max_iterations):
        iterations += 1
        mid = 0.5 * (low + high)
        f_mid = value(mid)
        if abs(f_mid) <= tolerance:
            return FairCouponResult(
                coupon=float(mid),
                pv=f_mid,
                lower=low,
                upper=high,
                pv_lower=f_low,
                pv_upper=f_high,
                iterations=iterations,
                solved=True,
            )
        if f_low * f_mid <= 0.0:
            high = mid
            f_high = f_mid
        else:
            low = mid
            f_low = f_mid
    coupon = float(0.5 * (low + high))
    pv = value(coupon)
    return FairCouponResult(
        coupon=coupon,
        pv=pv,
        lower=low,
        upper=high,
        pv_lower=f_low,
        pv_upper=f_high,
        iterations=iterations,
        solved=abs(pv) <= tolerance,
    )


def solve_case_study_coupons(
    *,
    dates: pd.Series,
    initial_spot: float,
    issue_date: pd.Timestamp,
    terms: CaseStudyTerms,
    engine_config: AutocallableEngineConfig,
) -> tuple[dict[str, float], dict[str, FairCouponResult]]:
    env = pricing_environment(initial_spot, issue_date, terms)

    coupons: dict[str, float] = {}
    results: dict[str, FairCouponResult] = {}
    labels = ["PPP-DKI", "NPP-DKI", "PPP-EKI-Parachute"]
    for label in labels:
        def builder(coupon: float, label: str = label) -> SnowballOption:
            return build_case_study_products(
                initial_spot=initial_spot,
                issue_date=issue_date,
                dates=dates,
                terms=terms,
                coupons={label: coupon},
            )[label]

        result = solve_fair_coupon_result(
            product_builder=builder,
            engine_config=engine_config,
            env=env,
        )
        coupons[label] = result.coupon
        results[label] = result
    return coupons, results


def build_backtest_config(
    *,
    product: SnowballOption,
    market_data: AutocallableMarketDataSet,
    engine_config: AutocallableEngineConfig,
    terms: CaseStudyTerms,
    args: argparse.Namespace,
    scenario: str,
    product_label: str,
) -> AutocallableBacktestConfig:
    dates = market_data.dates
    return AutocallableBacktestConfig(
        product=product,
        market_data=market_data,
        engine_config=engine_config,
        strategy=AutocallableDeltaHedgeStrategy(
            delta_threshold=0.0,
            hedge_ratio=1.0,
            round_contracts=True,
        ),
        roll_policy=FuturesRollPolicy(roll_days_before_expiry=5),
        product_quantity=-1.0,
        underlying=UNDERLYING_NAME,
        start_date=pd.Timestamp(dates[0]).to_pydatetime(),
        end_date=pd.Timestamp(dates[-1]).to_pydatetime(),
        initial_product_price=0.0,
        fixed_dividend_yield=float(terms.dividend_yield),
        surface_config=SurfaceGridConfig(spot_nodes=5, spot_width=0.05, q_nodes=3, q_width=0.005),
        calculate_surfaces=bool(args.surfaces),
        calculate_event_probabilities=bool(args.event_probabilities),
        metadata={
            "case_study": "ppp_dki_snowball_backtest",
            "scenario": scenario,
            "product": product_label,
            "engine": args.engine,
        },
    )


def _write_frame(df: pd.DataFrame, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    df.to_csv(path)


def run_single_backtest(
    *,
    product_label: str,
    product: SnowballOption,
    scenario: str,
    market_data: AutocallableMarketDataSet,
    terms: CaseStudyTerms,
    engine_config: AutocallableEngineConfig,
    args: argparse.Namespace,
    output_dir: Path,
) -> tuple[Any, dict[str, Any]]:
    config = build_backtest_config(
        product=product,
        market_data=market_data,
        engine_config=engine_config,
        terms=terms,
        args=args,
        scenario=scenario,
        product_label=product_label,
    )
    results = AutocallableBacktestEngine(config).run()
    run_dir = output_dir / "intermediate" / scenario / product_label.replace("/", "_")
    _write_frame(results.states_df, run_dir / "states.csv")
    _write_frame(results.greeks_df, run_dir / "greeks.csv")
    _write_frame(results.rebalance_df, run_dir / "rebalances.csv")
    _write_frame(results.trades_df, run_dir / "trades.csv")
    _write_frame(results.actions_df, run_dir / "actions.csv")
    _write_frame(results.daily_event_summary_df, run_dir / "daily_event_summary.csv")
    _write_frame(results.event_probability_df, run_dir / "event_probabilities.csv")

    summary = results.get_summary()
    states = results.states_df
    summary.update(
        {
            "scenario": scenario,
            "product": product_label,
            "final_total_pnl": 0.0 if states.empty else float(states["total_pnl"].iloc[-1]),
            "min_total_pnl": 0.0 if states.empty else float(states["total_pnl"].min()),
            "max_drawdown": calculate_drawdown(states["total_pnl"]).min() if not states.empty else 0.0,
            "num_trades": int(len(results.trades_df)),
            "num_actions": int(len(results.actions_df)),
            "lifecycle_actions": ""
            if results.actions_df.empty
            else ",".join(sorted(results.actions_df["action_type"].astype(str).unique())),
            "final_hedge_contracts": 0.0
            if states.empty
            else float(states["futures_contracts"].iloc[-1]),
        }
    )
    return results, summary


def calculate_drawdown(series: pd.Series) -> pd.Series:
    running_max = series.cummax()
    return series - running_max


def _tagged_frame(df: pd.DataFrame, scenario: str, product: str) -> pd.DataFrame:
    out = df.reset_index().copy()
    out.insert(0, "product", product)
    out.insert(0, "scenario", scenario)
    return out


def build_consolidated_frames(run_results: dict[tuple[str, str], Any]) -> dict[str, pd.DataFrame]:
    frames: dict[str, list[pd.DataFrame]] = {
        "daily_greeks": [],
        "daily_pnl": [],
        "hedge_actions": [],
        "lifecycle_events": [],
    }
    for (scenario, product), results in run_results.items():
        greeks = results.greeks_df
        states = results.states_df
        rebalances = results.rebalance_df
        trades = results.trades_df
        actions = results.actions_df
        if not greeks.empty:
            frames["daily_greeks"].append(_tagged_frame(greeks, scenario, product))
        if not states.empty:
            pnl_cols = [
                "portfolio_value",
                "product_mtm",
                "hedge_mtm",
                "cash",
                "product_pnl",
                "hedge_pnl",
                "total_pnl",
                "spot",
                "futures_contracts",
                "active_contract",
                "pricing_q",
                "implied_q",
            ]
            available = [col for col in pnl_cols if col in states.columns]
            frames["daily_pnl"].append(_tagged_frame(states[available], scenario, product))
        if not rebalances.empty:
            frames["hedge_actions"].append(_tagged_frame(rebalances, scenario, product))
        if not trades.empty:
            frames["hedge_actions"].append(_tagged_frame(trades, scenario, product))
        if not actions.empty:
            frames["lifecycle_events"].append(_tagged_frame(actions, scenario, product))

    consolidated: dict[str, pd.DataFrame] = {}
    for name, parts in frames.items():
        if not parts:
            consolidated[name] = pd.DataFrame()
            continue
        with warnings.catch_warnings():
            warnings.filterwarnings(
                "ignore",
                message="The behavior of DataFrame concatenation with empty or all-NA entries is deprecated.*",
                category=FutureWarning,
            )
            consolidated[name] = pd.concat(parts, ignore_index=True)
    return consolidated


def terms_frame(
    terms: CaseStudyTerms,
    coupons: dict[str, float],
    initial_spot: float,
    coupon_results: Optional[dict[str, FairCouponResult]] = None,
    issue_date: Optional[pd.Timestamp] = None,
    dates: Optional[pd.Series] = None,
) -> pd.DataFrame:
    normalized_issue_date = pd.Timestamp(issue_date).normalize() if issue_date is not None else None
    normalized_dates = pd.Series(dtype="datetime64[ns]")
    if dates is not None:
        normalized_dates = pd.to_datetime(pd.Series(dates), errors="coerce").dropna().dt.normalize()
    if normalized_issue_date is None and not normalized_dates.empty:
        normalized_issue_date = pd.Timestamp(normalized_dates.iloc[0]).normalize()

    maturity = 3.0
    if normalized_issue_date is not None and not normalized_dates.empty:
        maturity = max(
            (pd.Timestamp(normalized_dates.iloc[-1]).normalize() - normalized_issue_date).days
            / 365.0,
            1.0 / 365.0,
        )

    ko_times = ko_observation_times(maturity)
    maturity_date = (
        pd.Timestamp(normalized_dates.iloc[-1]).normalize()
        if not normalized_dates.empty
        else (
            normalized_issue_date + pd.Timedelta(days=int(round(maturity * 365.0)))
            if normalized_issue_date is not None
            else None
        )
    )

    def _date_text(date: Optional[pd.Timestamp]) -> str:
        if date is None or pd.isna(date):
            return "n/a"
        return pd.Timestamp(date).strftime("%Y-%m-%d")

    def _schedule_date_from_time(time: float) -> Optional[pd.Timestamp]:
        if normalized_issue_date is None:
            return None
        target = normalized_issue_date + pd.Timedelta(days=int(round(time * 365.0)))
        if normalized_dates.empty:
            return target
        candidates = normalized_dates[normalized_dates >= target]
        return pd.Timestamp(candidates.iloc[0]).normalize() if not candidates.empty else maturity_date

    ko_dates = [_schedule_date_from_time(time) for time in ko_times]
    dki_dates = (
        normalized_dates[normalized_dates > normalized_issue_date]
        if normalized_issue_date is not None and not normalized_dates.empty
        else pd.Series(dtype="datetime64[ns]")
    )

    def _date_schedule_rows(
        term_prefix: str,
        schedule_dates: list[Optional[pd.Timestamp]] | pd.Series,
        *,
        chunk_size: int = 18,
    ) -> list[dict[str, Any]]:
        date_texts = [_date_text(date) for date in list(schedule_dates)]
        if not date_texts:
            return [{"term": term_prefix, "value": "n/a"}]
        return [
            {
                "term": f"{term_prefix}.{idx // chunk_size + 1:02d}",
                "value": ", ".join(date_texts[idx : idx + chunk_size]),
            }
            for idx in range(0, len(date_texts), chunk_size)
        ]

    ko_window = (
        f"{len(ko_times)} monthly discrete observations; "
        f"{_date_text(ko_dates[0] if ko_dates else None)} to {_date_text(ko_dates[-1] if ko_dates else None)}"
    )
    dki_window = (
        f"{len(dki_dates)} trading-day discrete observations; "
        f"{_date_text(dki_dates.iloc[0] if not dki_dates.empty else None)} to "
        f"{_date_text(dki_dates.iloc[-1] if not dki_dates.empty else None)}"
    )
    rows = [
        {"term": "underlying", "value": f"{UNDERLYING_SYMBOL}.SH"},
        {"term": "hedging_instrument", "value": "IM.CFE"},
        {"term": "initial_spot", "value": initial_spot},
        {"term": "issue_date", "value": _date_text(normalized_issue_date)},
        {"term": "maturity_date", "value": _date_text(maturity_date)},
        {"term": "tenor_years_realized", "value": maturity},
        {"term": "pv_convention", "value": "principal_excluded_zero_upfront"},
        {"term": "payoff_include_principal", "value": False},
        {"term": "initial_product_price", "value": "0.0"},
        {"term": "notional", "value": terms.notional},
        {"term": "ko_ratio", "value": terms.ko_ratio},
        {"term": "ki_ratio", "value": terms.ki_ratio},
        {"term": "ko_barrier_level", "value": float(initial_spot) * float(terms.ko_ratio)},
        {"term": "ki_barrier_level", "value": float(initial_spot) * float(terms.ki_ratio)},
        {
            "term": "observation_schedule_date_basis",
            "value": "KO target dates are mapped to the first available path date on or after each monthly observation time; DKI dates use every trading date after issue date",
        },
        {"term": "ko_observation.PPP-DKI", "value": ko_window},
        {"term": "ko_observation.NPP-DKI", "value": ko_window},
        {
            "term": "ko_observation.PPP-EKI-Parachute",
            "value": f"{ko_window}; final KO barrier equals KI barrier ({terms.ki_ratio:.2%})",
        },
        {"term": "ki_observation.PPP-DKI", "value": dki_window},
        {"term": "ki_observation.NPP-DKI", "value": dki_window},
        {
            "term": "ki_observation.PPP-EKI-Parachute",
            "value": f"final-only discrete KI observation on {_date_text(maturity_date)}",
        },
        {
            "term": "final_observation_precedence",
            "value": "KO has precedence at a final KO/KI tie; spot >= KI exits by KO, spot < KI settles through KI/maturity payoff",
        },
        {
            "term": "post_ki_ko_allowed",
            "value": "True; later KO remains allowed after KI unless a product explicitly disables it",
        },
        {"term": "ppp_protection_rate", "value": terms.ppp_protection_rate},
        {"term": "rate", "value": terms.rate},
        {"term": "dividend_yield", "value": terms.dividend_yield},
        {"term": "volatility", "value": terms.volatility},
        {"term": "futures_multiplier", "value": FUTURES_MULTIPLIER},
    ]
    for product_label in ["PPP-DKI", "NPP-DKI", "PPP-EKI-Parachute"]:
        rows.extend(_date_schedule_rows(f"ko_observation_dates.{product_label}", ko_dates))
    for product_label in ["PPP-DKI", "NPP-DKI"]:
        rows.extend(_date_schedule_rows(f"ki_observation_dates.{product_label}", dki_dates, chunk_size=18))
    rows.extend(
        _date_schedule_rows(
            "ki_observation_dates.PPP-EKI-Parachute",
            [maturity_date],
            chunk_size=18,
        )
    )
    for label, coupon in coupons.items():
        rows.append({"term": f"fair_coupon.{label}", "value": coupon})
        if coupon_results and label in coupon_results:
            result = coupon_results[label]
            rows.extend(
                [
                    {"term": f"fair_coupon_pv.{label}", "value": result.pv},
                    {"term": f"fair_coupon_solved.{label}", "value": result.solved},
                    {"term": f"fair_coupon_iterations.{label}", "value": result.iterations},
                    {"term": f"fair_coupon_bracket_pv_low.{label}", "value": result.pv_lower},
                    {"term": f"fair_coupon_bracket_pv_high.{label}", "value": result.pv_upper},
                ]
            )
    return pd.DataFrame(rows)


def write_csv_outputs(
    output_dir: Path,
    summaries: list[dict[str, Any]],
    consolidated: dict[str, pd.DataFrame],
) -> None:
    data_dir = output_dir / "data"
    data_dir.mkdir(parents=True, exist_ok=True)
    pd.DataFrame(summaries).to_csv(data_dir / "summary_metrics.csv", index=False)
    for name, frame in consolidated.items():
        frame.to_csv(data_dir / f"{name}.csv", index=False)


def write_excel_output(
    path: Path,
    summaries: list[dict[str, Any]],
    terms: pd.DataFrame,
    consolidated: dict[str, pd.DataFrame],
) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        pd.DataFrame(summaries).to_excel(writer, sheet_name="Summary", index=False)
        terms.to_excel(writer, sheet_name="Terms", index=False)
        for name, frame in consolidated.items():
            sheet = {
                "daily_greeks": "Daily_Greeks",
                "daily_pnl": "Daily_PnL",
                "hedge_actions": "Hedge_Actions",
                "lifecycle_events": "Lifecycle_Events",
            }[name]
            frame.to_excel(writer, sheet_name=sheet, index=False)


def require_matplotlib():
    os.environ.setdefault("MPLCONFIGDIR", "/tmp/matplotlib")
    import matplotlib.pyplot as plt  # type: ignore

    return plt


def save_line_chart(
    frame: pd.DataFrame,
    *,
    path: Path,
    value_col: str,
    title: str,
    y_label: str,
) -> Optional[Path]:
    if frame.empty or value_col not in frame.columns:
        return None
    plt = require_matplotlib()
    path.parent.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(10, 5))
    plot_frame = frame.copy()
    plot_frame["date"] = pd.to_datetime(plot_frame["date"])
    for (scenario, product), group in plot_frame.groupby(["scenario", "product"]):
        label = f"{scenario} / {product}"
        ax.plot(group["date"], group[value_col], linewidth=1.4, label=label)
    ax.set_title(title)
    ax.set_xlabel("Date")
    ax.set_ylabel(y_label)
    ax.grid(True, alpha=0.25)
    ax.legend(fontsize=7, ncol=2)
    fig.autofmt_xdate()
    fig.tight_layout()
    fig.savefig(path, dpi=150)
    plt.close(fig)
    return path


def save_bar_chart(
    frame: pd.DataFrame,
    *,
    path: Path,
    title: str,
    y_col: str,
) -> Optional[Path]:
    if frame.empty or y_col not in frame.columns:
        return None
    plt = require_matplotlib()
    path.parent.mkdir(parents=True, exist_ok=True)
    pivot = frame.pivot(index="scenario", columns="product", values=y_col)
    fig, ax = plt.subplots(figsize=(9, 5))
    pivot.plot(kind="bar", ax=ax)
    ax.set_title(title)
    ax.set_ylabel(y_col)
    ax.grid(True, axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(path, dpi=150)
    plt.close(fig)
    return path


def save_price_path_chart(
    frame: pd.DataFrame,
    *,
    path: Path,
    terms: pd.DataFrame,
) -> Optional[Path]:
    if frame.empty or not {"scenario", "date", "spot"}.issubset(frame.columns):
        return None
    plt = require_matplotlib()
    path.parent.mkdir(parents=True, exist_ok=True)
    terms_map = dict(zip(terms["term"].astype(str), terms["value"]))
    initial_spot = float(terms_map.get("initial_spot", np.nan))
    ko_level = initial_spot * float(terms_map.get("ko_ratio", np.nan))
    ki_level = initial_spot * float(terms_map.get("ki_ratio", np.nan))

    plot_frame = (
        frame[["scenario", "date", "spot"]]
        .drop_duplicates(["scenario", "date"])
        .copy()
    )
    plot_frame["date"] = pd.to_datetime(plot_frame["date"])
    fig, ax = plt.subplots(figsize=(10, 5))
    for scenario, group in plot_frame.groupby("scenario", sort=False):
        ax.plot(group["date"], group["spot"], linewidth=1.6, label=str(scenario))
    if math.isfinite(initial_spot):
        ax.axhline(initial_spot, color="#475467", linestyle="-", linewidth=1.0, alpha=0.65, label="Initial spot")
    if math.isfinite(ko_level):
        ax.axhline(ko_level, color="#087443", linestyle="--", linewidth=1.0, alpha=0.8, label="KO")
    if math.isfinite(ki_level):
        ax.axhline(ki_level, color="#a63838", linestyle="--", linewidth=1.0, alpha=0.8, label="KI")
    ax.set_title("Underlying Price Paths by Scenario")
    ax.set_xlabel("Date")
    ax.set_ylabel("000852.SH spot")
    ax.grid(True, alpha=0.25)
    ax.legend(fontsize=8, ncol=3)
    fig.autofmt_xdate()
    fig.tight_layout()
    fig.savefig(path, dpi=150)
    plt.close(fig)
    return path


def generate_charts(
    output_dir: Path,
    summaries: list[dict[str, Any]],
    consolidated: dict[str, pd.DataFrame],
    terms: pd.DataFrame,
) -> dict[str, Path]:
    chart_dir = output_dir / "charts"
    summary_df = pd.DataFrame(summaries)
    paths = {
        "price_paths": save_price_path_chart(
            consolidated["daily_pnl"],
            path=chart_dir / "scenario_price_paths.png",
            terms=terms,
        ),
        "pnl": save_line_chart(
            consolidated["daily_pnl"],
            path=chart_dir / "total_pnl.png",
            value_col="total_pnl",
            title="Total PnL by Scenario and Product",
            y_label="PnL",
        ),
        "delta": save_line_chart(
            consolidated["daily_greeks"],
            path=chart_dir / "post_hedge_delta.png",
            value_col="post_hedge_delta",
            title="Post-Hedge Delta",
            y_label="Delta",
        ),
        "gamma": save_line_chart(
            consolidated["daily_greeks"],
            path=chart_dir / "post_hedge_gamma.png",
            value_col="post_hedge_gamma",
            title="Post-Hedge Gamma",
            y_label="Gamma",
        ),
        "hedge": save_line_chart(
            consolidated["daily_pnl"],
            path=chart_dir / "futures_contracts.png",
            value_col="futures_contracts",
            title="Futures Hedge Contracts",
            y_label="Contracts",
        ),
        "scenario_pnl": save_bar_chart(
            summary_df,
            path=chart_dir / "scenario_final_pnl.png",
            title="Final Total PnL",
            y_col="final_total_pnl",
        ),
    }
    return {key: value for key, value in paths.items() if value is not None}


def write_docx_report(
    path: Path,
    summaries: list[dict[str, Any]],
    terms: pd.DataFrame,
    charts: dict[str, Path],
) -> Path:
    try:
        from docx import Document  # type: ignore
        from docx.shared import Inches  # type: ignore
    except ImportError as exc:
        raise CaseStudyError("python-docx is required to write the DOCX report") from exc

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("PPP/NPP Snowball Backtest Case Study", 0)
    doc.add_paragraph(
        "Delta-neutral hedge comparison for PPP-DKI, NPP-DKI, and PPP-EKI "
        "Parachute Snowballs on 000852.SH hedged with IM futures."
    )

    doc.add_heading("Terms", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Term"
    table.rows[0].cells[1].text = "Value"
    for _, row in terms.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["term"])
        cells[1].text = str(row["value"])

    doc.add_heading("Scenario Metrics", level=1)
    summary_df = pd.DataFrame(summaries)
    cols = ["scenario", "product", "final_total_pnl", "min_total_pnl", "num_trades", "lifecycle_actions"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    for i, col in enumerate(cols):
        table.rows[0].cells[i].text = col
    for _, row in summary_df[cols].iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(cols):
            value = row[col]
            if isinstance(value, float):
                value = f"{value:,.2f}"
            cells[i].text = str(value)

    doc.add_heading("Charts", level=1)
    for title, chart_path in charts.items():
        doc.add_paragraph(title.replace("_", " ").title())
        doc.add_picture(str(chart_path), width=Inches(6.5))

    doc.save(path)
    return path


def write_html_dashboard(
    path: Path,
    summaries: list[dict[str, Any]],
    terms: pd.DataFrame,
    consolidated: dict[str, pd.DataFrame],
) -> Path:
    import plotly.graph_objects as go  # type: ignore
    from plotly.offline import plot  # type: ignore

    path.parent.mkdir(parents=True, exist_ok=True)
    summary_df = pd.DataFrame(summaries)
    daily_pnl = consolidated["daily_pnl"].copy()
    daily_greeks = consolidated["daily_greeks"].copy()
    hedge_actions = consolidated["hedge_actions"].copy()
    lifecycle_events = consolidated["lifecycle_events"].copy()

    def fmt_number(value: Any, digits: int = 0) -> str:
        try:
            if pd.isna(value):
                return "-"
            return f"{float(value):,.{digits}f}"
        except (TypeError, ValueError):
            return "-"

    def fmt_signed(value: Any, digits: int = 0) -> str:
        try:
            if pd.isna(value):
                return "-"
            number = float(value)
            return f"{number:+,.{digits}f}"
        except (TypeError, ValueError):
            return "-"

    def fmt_percent(value: Any, digits: int = 2) -> str:
        try:
            if pd.isna(value):
                return "-"
            return f"{float(value) * 100.0:.{digits}f}%"
        except (TypeError, ValueError):
            return "-"

    def esc(value: Any) -> str:
        return html_lib.escape("" if value is None else str(value))

    def table_html(
        frame: pd.DataFrame,
        columns: list[str],
        *,
        limit: Optional[int] = None,
        numeric: Optional[set[str]] = None,
        signed: Optional[set[str]] = None,
    ) -> str:
        numeric = numeric or set()
        signed = signed or set()
        if frame.empty:
            return '<div class="empty-state">No records</div>'
        view = frame.loc[:, [col for col in columns if col in frame.columns]].copy()
        if limit is not None:
            view = view.head(limit)
        headers = "".join(f"<th>{esc(col.replace('_', ' ').title())}</th>" for col in view.columns)
        body_rows = []
        for _, row in view.iterrows():
            cells = []
            for col in view.columns:
                value = row[col]
                if col in signed:
                    text = fmt_signed(value)
                    cls = "num neg" if float(value or 0.0) < 0 else "num pos"
                elif col in numeric:
                    text = fmt_number(value)
                    cls = "num"
                elif isinstance(value, float):
                    text = fmt_number(value, 4)
                    cls = "num"
                else:
                    text = esc(value)
                    cls = ""
                cells.append(f'<td class="{cls}">{text}</td>')
            body_rows.append("<tr>" + "".join(cells) + "</tr>")
        return f"<table><thead><tr>{headers}</tr></thead><tbody>{''.join(body_rows)}</tbody></table>"

    def observation_schedule_payload(frame: pd.DataFrame) -> dict[str, list[str]]:
        schedules: dict[str, list[str]] = {}
        if frame.empty or not {"term", "value"}.issubset(frame.columns):
            return schedules
        schedule_rows = frame[frame["term"].astype(str).str.startswith(("ko_observation_dates.", "ki_observation_dates."))]
        for _, row in schedule_rows.iterrows():
            term = str(row["term"])
            prefix = term.rsplit(".", 1)[0]
            key = prefix.replace("_dates", "", 1)
            dates = [item.strip() for item in str(row["value"]).split(",") if item.strip()]
            schedules.setdefault(key, []).extend(dates)
        return schedules

    observation_schedules = observation_schedule_payload(terms)

    def terms_table_html(frame: pd.DataFrame) -> str:
        if frame.empty:
            return '<div class="empty-state">No records</div>'
        visible = frame[
            ~frame["term"].astype(str).str.startswith(("ko_observation_dates.", "ki_observation_dates."))
        ].copy()
        headers = "<th>Term</th><th>Value</th>"
        rows = []
        for _, row in visible.iterrows():
            term = str(row["term"])
            value = row["value"]
            if isinstance(value, float):
                value_text = fmt_number(value, 4)
                value_cls = "num"
            else:
                value_text = esc(value)
                value_cls = ""
            if term in observation_schedules:
                value_html = (
                    '<div class="term-value-action">'
                    f"<span>{value_text}</span>"
                    f'<button class="obs-schedule-button" type="button" data-observation-key="{esc(term)}">View dates</button>'
                    "</div>"
                )
                value_cls = ""
            else:
                value_html = value_text
            rows.append(
                "<tr>"
                f'<td class="term-name">{esc(term)}</td>'
                f'<td class="{value_cls}">{value_html}</td>'
                "</tr>"
            )
        return f"<table><thead><tr>{headers}</tr></thead><tbody>{''.join(rows)}</tbody></table>"

    def metric_card(label: str, value: str, detail: str = "", tone: str = "") -> str:
        return (
            f'<div class="metric {tone}"><span>{esc(label)}</span>'
            f"<strong>{esc(value)}</strong><small>{esc(detail)}</small></div>"
        )

    def plot_div(fig: Any, *, include_plotlyjs: bool = False) -> str:
        fig.update_layout(
            template="plotly_white",
            paper_bgcolor="rgba(0,0,0,0)",
            plot_bgcolor="rgba(250,252,255,0.92)",
            margin=dict(l=54, r=24, t=54, b=44),
            font=dict(family="Aptos, Segoe UI, sans-serif", size=12, color="#182235"),
            legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="left", x=0),
        )
        return plot(
            fig,
            include_plotlyjs=include_plotlyjs,
            output_type="div",
            config={"responsive": True, "displaylogo": False},
        )

    if not daily_pnl.empty:
        daily_pnl["date"] = pd.to_datetime(daily_pnl["date"])
        daily_pnl = daily_pnl.sort_values(["scenario", "product", "date"])
        daily_pnl["daily_product_pnl"] = daily_pnl.groupby(["scenario", "product"])[
            "product_pnl"
        ].diff().fillna(daily_pnl["product_pnl"])
        daily_pnl["daily_hedge_pnl"] = daily_pnl.groupby(["scenario", "product"])[
            "hedge_pnl"
        ].diff().fillna(daily_pnl["hedge_pnl"])
        daily_pnl["daily_total_pnl"] = daily_pnl.groupby(["scenario", "product"])[
            "total_pnl"
        ].diff().fillna(daily_pnl["total_pnl"])
        daily_pnl["drawdown"] = daily_pnl.groupby(["scenario", "product"])[
            "total_pnl"
        ].transform(lambda s: s - s.cummax())
    if not daily_greeks.empty:
        daily_greeks["date"] = pd.to_datetime(daily_greeks["date"])
        daily_greeks = daily_greeks.sort_values(["scenario", "product", "date"])
    if not hedge_actions.empty and "date" in hedge_actions.columns:
        hedge_actions["date"] = pd.to_datetime(hedge_actions["date"])
    if not lifecycle_events.empty and "date" in lifecycle_events.columns:
        lifecycle_events["date"] = pd.to_datetime(lifecycle_events["date"])

    summary_for_display = summary_df.copy()
    if not daily_pnl.empty:
        risk_agg = daily_pnl.groupby(["scenario", "product"]).agg(
            worst_daily_pnl=("daily_total_pnl", "min"),
            worst_drawdown=("drawdown", "min"),
            final_product_pnl=("product_pnl", "last"),
            final_hedge_pnl=("hedge_pnl", "last"),
            final_contracts=("futures_contracts", "last"),
        )
        summary_for_display = summary_for_display.merge(
            risk_agg.reset_index(), on=["scenario", "product"], how="left"
        )
    if not daily_greeks.empty:
        greek_agg = daily_greeks.groupby(["scenario", "product"]).agg(
            max_abs_post_delta_cash=("post_hedge_delta_cash_1pct", lambda s: float(s.abs().max())),
            max_abs_gamma_cash=("post_hedge_gamma_cash_1pct", lambda s: float(s.abs().max())),
        )
        summary_for_display = summary_for_display.merge(
            greek_agg.reset_index(), on=["scenario", "product"], how="left"
        )
    if not hedge_actions.empty:
        trades_only = hedge_actions[hedge_actions.get("trade_type").notna()] if "trade_type" in hedge_actions.columns else pd.DataFrame()
        if not trades_only.empty:
            trade_agg = trades_only.groupby(["scenario", "product"]).agg(
                gross_trade_notional=("notional", "sum"),
                max_trade_notional=("notional", "max"),
                roll_count=("trade_type", lambda s: int(s.astype(str).str.startswith("roll_").sum())),
            )
            summary_for_display = summary_for_display.merge(
                trade_agg.reset_index(), on=["scenario", "product"], how="left"
            )

    best_row = summary_for_display.loc[summary_for_display["final_total_pnl"].idxmax()]
    worst_row = summary_for_display.loc[summary_for_display["final_total_pnl"].idxmin()]
    worst_drawdown_row = summary_for_display.loc[summary_for_display["worst_drawdown"].idxmin()] if "worst_drawdown" in summary_for_display else worst_row
    action_counts = (
        lifecycle_events["action_type"].astype(str).value_counts().to_dict()
        if not lifecycle_events.empty and "action_type" in lifecycle_events.columns
        else {}
    )
    roll_issues = (
        hedge_actions[
            hedge_actions.get("reason", pd.Series(dtype=str)).astype(str)
            == "futures_roll_missing_old_contract"
        ]
        if not hedge_actions.empty and "reason" in hedge_actions.columns
        else pd.DataFrame()
    )

    terms_map = dict(zip(terms["term"].astype(str), terms["value"]))
    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    scenario_count = int(summary_for_display["scenario"].nunique()) if not summary_for_display.empty else 0
    product_count = int(summary_for_display["product"].nunique()) if not summary_for_display.empty else 0

    executive_cards = "".join(
        [
            metric_card("Run Set", f"{len(summary_for_display)} runs", f"{scenario_count} scenarios x {product_count} products"),
            metric_card("Best Final PnL", fmt_signed(best_row["final_total_pnl"]), f"{best_row['scenario']} / {best_row['product']}", "good"),
            metric_card("Worst Final PnL", fmt_signed(worst_row["final_total_pnl"]), f"{worst_row['scenario']} / {worst_row['product']}", "bad"),
            metric_card("Worst Drawdown", fmt_signed(worst_drawdown_row.get("worst_drawdown", 0.0)), f"{worst_drawdown_row['scenario']} / {worst_drawdown_row['product']}", "warn"),
            metric_card("Lifecycle Events", str(sum(action_counts.values())), ", ".join(f"{k}:{v}" for k, v in sorted(action_counts.items())) or "none"),
            metric_card("Roll Data Gaps", str(len(roll_issues)), "proxy close rows used" if len(roll_issues) else "none observed"),
        ]
    )

    final_pnl_matrix = summary_for_display.pivot(
        index="scenario", columns="product", values="final_total_pnl"
    )
    pnl_heatmap = go.Figure(
        data=go.Heatmap(
            z=final_pnl_matrix.values,
            x=list(final_pnl_matrix.columns),
            y=list(final_pnl_matrix.index),
            colorscale="RdYlGn",
            zmid=0,
            colorbar=dict(title="PnL"),
            text=np.vectorize(lambda x: fmt_signed(x))(final_pnl_matrix.values),
            hovertemplate="Scenario=%{y}<br>Product=%{x}<br>Final PnL=%{text}<extra></extra>",
        )
    )
    pnl_heatmap.update_layout(title="Final PnL Matrix")

    drawdown_matrix = summary_for_display.pivot(
        index="scenario", columns="product", values="worst_drawdown"
    )
    drawdown_heatmap = go.Figure(
        data=go.Heatmap(
            z=drawdown_matrix.values,
            x=list(drawdown_matrix.columns),
            y=list(drawdown_matrix.index),
            colorscale="Reds_r",
            colorbar=dict(title="Drawdown"),
            text=np.vectorize(lambda x: fmt_signed(x))(drawdown_matrix.values),
            hovertemplate="Scenario=%{y}<br>Product=%{x}<br>Drawdown=%{text}<extra></extra>",
        )
    )
    drawdown_heatmap.update_layout(title="Worst Drawdown Matrix")

    pnl_trajectory = go.Figure()
    if not daily_pnl.empty:
        for (scenario, product), group in daily_pnl.groupby(["scenario", "product"]):
            pnl_trajectory.add_trace(
                go.Scatter(
                    x=group["date"],
                    y=group["total_pnl"],
                    mode="lines",
                    name=f"{scenario} / {product}",
                    hovertemplate="%{x|%Y-%m-%d}<br>PnL=%{y:,.0f}<extra></extra>",
                )
            )
    pnl_trajectory.update_layout(title="Cumulative PnL Paths", xaxis_title="Date", yaxis_title="PnL")

    attribution = go.Figure()
    if {"final_product_pnl", "final_hedge_pnl"}.issubset(summary_for_display.columns):
        labels = summary_for_display["scenario"].astype(str) + " / " + summary_for_display["product"].astype(str)
        attribution.add_trace(go.Bar(x=labels, y=summary_for_display["final_product_pnl"], name="Product PnL"))
        attribution.add_trace(go.Bar(x=labels, y=summary_for_display["final_hedge_pnl"], name="Hedge PnL"))
    attribution.update_layout(title="Final Product vs Hedge PnL", barmode="relative", xaxis_tickangle=-35, yaxis_title="PnL")

    hedge_contracts = go.Figure()
    if not daily_pnl.empty:
        for (scenario, product), group in daily_pnl.groupby(["scenario", "product"]):
            hedge_contracts.add_trace(
                go.Scatter(
                    x=group["date"],
                    y=group["futures_contracts"],
                    mode="lines",
                    name=f"{scenario} / {product}",
                )
            )
    hedge_contracts.update_layout(title="Futures Inventory", xaxis_title="Date", yaxis_title="Contracts")

    hedge_turnover = go.Figure()
    if "gross_trade_notional" in summary_for_display.columns:
        labels = summary_for_display["scenario"].astype(str) + " / " + summary_for_display["product"].astype(str)
        hedge_turnover.add_trace(
            go.Bar(
                x=labels,
                y=summary_for_display["gross_trade_notional"].fillna(0.0),
                marker_color="#4267ac",
                name="Gross trade notional",
            )
        )
    hedge_turnover.update_layout(title="Hedge Turnover", xaxis_tickangle=-35, yaxis_title="Gross Notional")

    delta_cash = go.Figure()
    if not daily_greeks.empty:
        for (scenario, product), group in daily_greeks.groupby(["scenario", "product"]):
            delta_cash.add_trace(
                go.Scatter(
                    x=group["date"],
                    y=group["post_hedge_delta_cash_1pct"],
                    mode="lines",
                    name=f"{scenario} / {product}",
                )
            )
    delta_cash.update_layout(title="Residual Delta Cash for 1% Spot Move", xaxis_title="Date", yaxis_title="Cash PnL / 1%")

    gamma_cash = go.Figure()
    if not daily_greeks.empty:
        for (scenario, product), group in daily_greeks.groupby(["scenario", "product"]):
            gamma_cash.add_trace(
                go.Scatter(
                    x=group["date"],
                    y=group["post_hedge_gamma_cash_1pct"],
                    mode="lines",
                    name=f"{scenario} / {product}",
                )
            )
    gamma_cash.update_layout(title="Gamma Cash for 1% Spot Move", xaxis_title="Date", yaxis_title="Cash Gamma / 1%")

    lifecycle_fig = go.Figure()
    if not lifecycle_events.empty:
        event_rank = {name: idx for idx, name in enumerate(sorted(lifecycle_events["action_type"].astype(str).unique()))}
        for action_type, group in lifecycle_events.groupby("action_type"):
            lifecycle_fig.add_trace(
                go.Scatter(
                    x=group["date"],
                    y=[event_rank[str(action_type)]] * len(group),
                    mode="markers",
                    name=str(action_type),
                    marker=dict(size=10),
                    text=group["scenario"].astype(str) + " / " + group["product"].astype(str),
                    hovertemplate="%{x|%Y-%m-%d}<br>%{text}<br>Spot=%{customdata:,.2f}<extra></extra>",
                    customdata=group["spot"].to_numpy() if "spot" in group.columns else None,
                )
            )
        lifecycle_fig.update_yaxes(
            tickmode="array",
            tickvals=list(event_rank.values()),
            ticktext=list(event_rank.keys()),
        )
    lifecycle_fig.update_layout(title="Lifecycle Event Timeline", xaxis_title="Date", yaxis_title="Event")

    price_paths = go.Figure()
    if not daily_pnl.empty:
        path_frame = (
            daily_pnl[["scenario", "date", "spot"]]
            .drop_duplicates(["scenario", "date"])
            .sort_values(["scenario", "date"])
        )
        for scenario, group in path_frame.groupby("scenario", sort=False):
            price_paths.add_trace(
                go.Scatter(
                    x=group["date"],
                    y=group["spot"],
                    mode="lines",
                    name=str(scenario),
                    hovertemplate="%{x|%Y-%m-%d}<br>Spot=%{y:,.2f}<extra></extra>",
                )
            )
        try:
            initial_spot = float(terms_map.get("initial_spot", np.nan))
            ko_level = initial_spot * float(terms_map.get("ko_ratio", np.nan))
            ki_level = initial_spot * float(terms_map.get("ki_ratio", np.nan))
        except (TypeError, ValueError):
            initial_spot = ko_level = ki_level = np.nan
        if math.isfinite(initial_spot):
            price_paths.add_hline(y=initial_spot, line_color="#475467", line_width=1, annotation_text="S0", annotation_position="top left")
        if math.isfinite(ko_level):
            price_paths.add_hline(y=ko_level, line_color="#087443", line_dash="dash", line_width=1, annotation_text="KO 103%", annotation_position="top left")
        if math.isfinite(ki_level):
            price_paths.add_hline(y=ki_level, line_color="#a63838", line_dash="dash", line_width=1, annotation_text="KI 75%", annotation_position="bottom left")
    price_paths.update_layout(
        title="Underlying Price Paths by Scenario",
        xaxis_title="Date",
        yaxis_title="000852.SH spot",
    )

    top_losses = (
        daily_pnl.sort_values("daily_total_pnl").head(12)
        if not daily_pnl.empty and "daily_total_pnl" in daily_pnl.columns
        else pd.DataFrame()
    )
    trader_exceptions = roll_issues.sort_values("date").head(20) if not roll_issues.empty else pd.DataFrame()
    risk_table = summary_for_display.sort_values("worst_drawdown").head(15)
    trader_table = summary_for_display.sort_values("gross_trade_notional", ascending=False) if "gross_trade_notional" in summary_for_display.columns else summary_for_display

    terms_html = terms_table_html(terms)
    executive_table = table_html(
        summary_for_display.sort_values(["scenario", "product"]),
        ["scenario", "product", "final_total_pnl", "worst_drawdown", "worst_daily_pnl", "num_trades", "lifecycle_actions"],
        signed={"final_total_pnl", "worst_drawdown", "worst_daily_pnl"},
        numeric={"num_trades"},
    )
    trader_summary_table = table_html(
        trader_table,
        ["scenario", "product", "gross_trade_notional", "max_trade_notional", "roll_count", "final_contracts", "num_trades"],
        numeric={"gross_trade_notional", "max_trade_notional", "roll_count", "final_contracts", "num_trades"},
    )
    trader_exceptions_table = table_html(
        trader_exceptions,
        ["date", "scenario", "product", "trade_type", "contract", "quantity", "price", "reason"],
        limit=20,
        numeric={"quantity", "price"},
    )
    risk_summary_table = table_html(
        risk_table,
        ["scenario", "product", "worst_drawdown", "worst_daily_pnl", "max_abs_post_delta_cash", "max_abs_gamma_cash", "final_total_pnl"],
        signed={"worst_drawdown", "worst_daily_pnl", "final_total_pnl"},
        numeric={"max_abs_post_delta_cash", "max_abs_gamma_cash"},
    )
    loss_days_table = table_html(
        top_losses,
        ["date", "scenario", "product", "daily_total_pnl", "total_pnl", "spot", "futures_contracts"],
        limit=12,
        signed={"daily_total_pnl", "total_pnl"},
        numeric={"spot", "futures_contracts"},
    )
    lifecycle_table = table_html(
        lifecycle_events.sort_values("date") if not lifecycle_events.empty else lifecycle_events,
        ["date", "scenario", "product", "action_type", "spot", "barrier", "cashflow", "knocked_in_before", "knocked_in_after"],
        limit=80,
        signed={"cashflow"},
        numeric={"spot", "barrier"},
    )

    def json_records(frame: pd.DataFrame) -> str:
        if frame.empty:
            return "[]"
        clean = frame.replace([np.inf, -np.inf], np.nan).copy()
        records: list[dict[str, Any]] = []
        for record in clean.to_dict("records"):
            row: dict[str, Any] = {}
            for key, value in record.items():
                if isinstance(value, pd.Timestamp):
                    row[key] = value.strftime("%Y-%m-%d")
                elif isinstance(value, datetime):
                    row[key] = value.strftime("%Y-%m-%d")
                elif isinstance(value, (np.integer, np.floating)):
                    number = value.item()
                    row[key] = None if pd.isna(number) else number
                elif pd.isna(value):
                    row[key] = None
                else:
                    row[key] = value
            records.append(row)
        return json.dumps(records, separators=(",", ":"), ensure_ascii=False)

    def _join_unique(series: pd.Series) -> str:
        values = [
            str(value)
            for value in series.dropna().astype(str)
            if str(value) and str(value).lower() != "nan"
        ]
        return ", ".join(sorted(set(values)))

    detail_frame = pd.DataFrame()
    if not daily_pnl.empty:
        detail_frame = daily_pnl.merge(
            daily_greeks,
            on=["scenario", "product", "date"],
            how="left",
            suffixes=("", "_greek"),
        )
        if not hedge_actions.empty and "trade_type" in hedge_actions.columns:
            trades_only = hedge_actions[hedge_actions["trade_type"].notna()].copy()
            if not trades_only.empty:
                trades_only["abs_quantity"] = trades_only["quantity"].abs()
                trades_only["abs_notional"] = trades_only["notional"].abs()
                trade_agg = trades_only.groupby(["scenario", "product", "date"]).agg(
                    trade_count=("trade_type", "count"),
                    trade_quantity=("quantity", "sum"),
                    trade_abs_quantity=("abs_quantity", "sum"),
                    trade_notional=("abs_notional", "sum"),
                    trade_types=("trade_type", _join_unique),
                )
                detail_frame = detail_frame.merge(
                    trade_agg.reset_index(),
                    on=["scenario", "product", "date"],
                    how="left",
                )
        if not lifecycle_events.empty:
            event_agg = lifecycle_events.groupby(["scenario", "product", "date"]).agg(
                lifecycle_actions=("action_type", _join_unique),
                lifecycle_cashflow=("cashflow", "sum"),
                lifecycle_barrier=("barrier", "first"),
            )
            detail_frame = detail_frame.merge(
                event_agg.reset_index(),
                on=["scenario", "product", "date"],
                how="left",
            )
        for column in [
            "trade_count",
            "trade_quantity",
            "trade_abs_quantity",
            "trade_notional",
            "lifecycle_cashflow",
        ]:
            if column not in detail_frame.columns:
                detail_frame[column] = 0.0
            detail_frame[column] = detail_frame[column].fillna(0.0)
        for column in ["trade_types", "lifecycle_actions"]:
            if column not in detail_frame.columns:
                detail_frame[column] = ""
            detail_frame[column] = detail_frame[column].fillna("")

    detail_columns = [
        "scenario",
        "product",
        "date",
        "spot",
        "active_contract",
        "futures_contracts",
        "product_mtm",
        "hedge_mtm",
        "cash",
        "daily_product_pnl",
        "daily_hedge_pnl",
        "daily_total_pnl",
        "total_pnl",
        "product_delta",
        "product_gamma",
        "post_hedge_delta",
        "post_hedge_delta_cash_1pct",
        "post_hedge_gamma_cash_1pct",
        "pricing_q",
        "implied_q",
        "trade_abs_quantity",
        "trade_notional",
        "trade_types",
        "lifecycle_actions",
        "lifecycle_cashflow",
    ]
    detail_payload = json_records(detail_frame[[col for col in detail_columns if col in detail_frame.columns]])
    summary_payload = json_records(summary_for_display)
    try:
        initial_spot_level = float(terms_map.get("initial_spot", np.nan))
        ko_level = initial_spot_level * float(terms_map.get("ko_ratio", np.nan))
        ki_level = initial_spot_level * float(terms_map.get("ki_ratio", np.nan))
    except (TypeError, ValueError):
        initial_spot_level = ko_level = ki_level = np.nan
    levels_payload = json.dumps(
        {
            "initial_spot": initial_spot_level if math.isfinite(initial_spot_level) else None,
            "ko": ko_level if math.isfinite(ko_level) else None,
            "ki": ki_level if math.isfinite(ki_level) else None,
        },
        separators=(",", ":"),
    )

    trade_chart_script = """
  <script id="trade-chart-data" type="application/json">__CHART_PAYLOAD__</script>
  <script id="summary-chart-data" type="application/json">__SUMMARY_PAYLOAD__</script>
  <script id="barrier-level-data" type="application/json">__LEVELS_PAYLOAD__</script>
  <script>
    (() => {
      const rows = JSON.parse(document.getElementById("trade-chart-data")?.textContent || "[]");
      const summaries = JSON.parse(document.getElementById("summary-chart-data")?.textContent || "[]");
      const levels = JSON.parse(document.getElementById("barrier-level-data")?.textContent || "{}");
      const scenarios = [...new Set(rows.map((row) => row.scenario).filter(Boolean))].sort();
      const products = [...new Set(rows.map((row) => row.product).filter(Boolean))].sort();
      const freqLabels = { D: "Day", W: "Week", M: "Month" };
      const chartKindLabels = { candle: "Candle", line: "Line" };
      const palette = {
        ink: "#17345d",
        green: "#087443",
        red: "#a63838",
        amber: "#9a6500",
        blue: "#4267ac",
        soft: "rgba(250,252,255,0.94)"
      };
      const seriesColors = [
        "#17345d", "#087443", "#a63838", "#9a6500", "#4267ac",
        "#7a4bd9", "#008b9a", "#c4507a", "#667085", "#d17a00"
      ];
      const rootStyle = getComputedStyle(document.documentElement);
      const plotFont = rootStyle.getPropertyValue("--font-body").trim() || "Aptos, Segoe UI, sans-serif";

      const fmt = (value) => {
        const number = Number(value);
        if (!Number.isFinite(number)) return "";
        return number.toLocaleString(undefined, { maximumFractionDigits: 0 });
      };

      const escapeHtml = (value) => String(value).replace(/[&<>"']/g, (char) => ({
        "&": "&amp;",
        "<": "&lt;",
        ">": "&gt;",
        '"': "&quot;",
        "'": "&#39;"
      })[char]);

      const setOptions = (select, values, preferred) => {
        if (!select) return;
        select.innerHTML = values.map((value) => `<option value="${escapeHtml(value)}">${escapeHtml(value)}</option>`).join("");
        const preferredValues = Array.isArray(preferred) ? preferred : [preferred];
        [...select.options].forEach((option) => {
          option.selected = preferredValues.includes(option.value);
        });
        if (!select.selectedOptions.length && select.options.length) select.options[0].selected = true;
      };

      const selectedValuesRaw = (select) => {
        if (!select) return [];
        return [...select.selectedOptions].map((option) => option.value).filter(Boolean);
      };

      const selectedValues = (select, fallbackValues) => {
        if (!select) return fallbackValues;
        const values = selectedValuesRaw(select);
        return values.length ? values : fallbackValues;
      };

      const filterLabel = (values, allValues) => {
        if (!values.length) return "None";
        if (values.length === allValues.length) return "All";
        if (values.length <= 2) return values.join(", ");
        return `${values.slice(0, 2).join(", ")} +${values.length - 2}`;
      };

      const updateMultiSelectControl = (select, allValues) => {
        const control = select?.closest("label")?.querySelector(".multi-select");
        if (!control) return;
        const selected = selectedValuesRaw(select);
        const effective = selected.length ? selected : allValues;
        const button = control.querySelector(".multi-select-button");
        if (button) button.querySelector(".multi-select-text").textContent = filterLabel(effective, allValues);
        control.querySelectorAll('input[type="checkbox"]').forEach((checkbox) => {
          const option = [...select.options].find((item) => item.value === checkbox.value);
          checkbox.checked = Boolean(option?.selected);
        });
      };

      const installMultiSelectControl = (select, allValues) => {
        const host = select?.closest("label");
        if (!host) return;
        host.classList.add("filter-label");
        if (host.querySelector(".multi-select")) {
          updateMultiSelectControl(select, allValues);
          return;
        }
        const control = document.createElement("div");
        control.className = "multi-select";
        control.innerHTML = `
          <button class="multi-select-button" type="button" aria-expanded="false">
            <span class="multi-select-text"></span>
          </button>
          <div class="multi-select-menu">
            <div class="multi-select-actions">
              <button class="multi-select-action" type="button" data-action="all">All</button>
            </div>
            ${allValues.map((value) => `
              <div class="multi-option">
                <input type="checkbox" value="${escapeHtml(value)}">
                <span>${escapeHtml(value)}</span>
              </div>
            `).join("")}
          </div>
        `;
        host.appendChild(control);
        const button = control.querySelector(".multi-select-button");
        button.addEventListener("click", (event) => {
          event.preventDefault();
          event.stopPropagation();
          document.querySelectorAll(".multi-select.open").forEach((other) => {
            if (other !== control) {
              other.classList.remove("open");
              other.querySelector(".multi-select-button")?.setAttribute("aria-expanded", "false");
            }
          });
          const isOpen = control.classList.toggle("open");
          button.setAttribute("aria-expanded", String(isOpen));
        });
        control.querySelector('[data-action="all"]')?.addEventListener("click", (event) => {
          event.preventDefault();
          [...select.options].forEach((option) => {
            option.selected = true;
          });
          updateMultiSelectControl(select, allValues);
          select.dispatchEvent(new Event("change", { bubbles: true }));
        });
        control.querySelectorAll('input[type="checkbox"]').forEach((checkbox) => {
          checkbox.addEventListener("change", () => {
            const option = [...select.options].find((item) => item.value === checkbox.value);
            if (option) option.selected = checkbox.checked;
            if (![...select.options].some((item) => item.selected)) {
              if (option) option.selected = true;
              checkbox.checked = true;
            }
            updateMultiSelectControl(select, allValues);
            select.dispatchEvent(new Event("change", { bubbles: true }));
          });
        });
        updateMultiSelectControl(select, allValues);
      };

      document.addEventListener("click", (event) => {
        if (event.target.closest(".multi-select")) return;
        document.querySelectorAll(".multi-select.open").forEach((control) => {
          control.classList.remove("open");
          control.querySelector(".multi-select-button")?.setAttribute("aria-expanded", "false");
        });
      });

      document.addEventListener("keydown", (event) => {
        if (event.key !== "Escape") return;
        document.querySelectorAll(".multi-select.open").forEach((control) => {
          control.classList.remove("open");
          control.querySelector(".multi-select-button")?.setAttribute("aria-expanded", "false");
        });
      });

      const isoDate = (date) => new Date(date).toISOString().slice(0, 10);
      const bucketKey = (dateText, freq) => {
        const date = new Date(dateText);
        if (freq === "M") return `${date.getUTCFullYear()}-${String(date.getUTCMonth() + 1).padStart(2, "0")}`;
        if (freq === "W") {
          const monday = new Date(Date.UTC(date.getUTCFullYear(), date.getUTCMonth(), date.getUTCDate()));
          const day = monday.getUTCDay() || 7;
          monday.setUTCDate(monday.getUTCDate() - day + 1);
          return isoDate(monday);
        }
        return isoDate(date);
      };

      const aggregateOhlc = (series, valueKey, freq) => {
        const sorted = series
          .filter((row) => row[valueKey] !== null && row[valueKey] !== undefined && Number.isFinite(Number(row[valueKey])))
          .sort((a, b) => new Date(a.date) - new Date(b.date));
        const buckets = [];
        let previousClose = null;
        for (const row of sorted) {
          const close = Number(row[valueKey]);
          const open = previousClose === null ? close : previousClose;
          previousClose = close;
          const high = Math.max(open, close);
          const low = Math.min(open, close);
          const key = bucketKey(row.date, freq);
          let bucket = buckets[buckets.length - 1];
          if (!bucket || bucket.key !== key) {
            bucket = { key, x: row.date, open, high, low, close, rows: 1 };
            buckets.push(bucket);
          } else {
            bucket.x = row.date;
            bucket.high = Math.max(bucket.high, high);
            bucket.low = Math.min(bucket.low, low);
            bucket.close = close;
            bucket.rows += 1;
          }
        }
        return buckets;
      };

      const layoutBase = (title, yTitle) => ({
        template: "plotly_white",
        font: { family: plotFont, size: 12, color: palette.ink },
        dragmode: "zoom",
        hovermode: "x unified",
        margin: { l: 58, r: 24, t: 42, b: 44 },
        paper_bgcolor: "rgba(0,0,0,0)",
        plot_bgcolor: palette.soft,
        title: { text: title, x: 0, xanchor: "left", font: { size: 15 } },
        showlegend: false,
        xaxis: {
          title: "",
          type: "date",
          rangeslider: { visible: true, thickness: 0.08 }
        },
        yaxis: { title: yTitle, zeroline: true, zerolinecolor: "#95a3b8" }
      });

      const activeFreq = (panel) => panel.querySelector(".freq-button.active")?.dataset.freq || "D";
      const updateButtonState = (panel, freq) => {
        panel.querySelectorAll(".freq-button").forEach((button) => {
          button.classList.toggle("active", button.dataset.freq === freq);
        });
      };
      const activeChartKind = (panel) => panel.querySelector(".chart-kind-button.active")?.dataset.kind || panel.dataset.defaultKind || "candle";
      const updateChartKindState = (panel, kind) => {
        panel.querySelectorAll(".chart-kind-button").forEach((button) => {
          button.classList.toggle("active", button.dataset.kind === kind);
        });
      };

      const deDuplicateByDate = (series) => {
        const byDate = new Map();
        series
          .filter((row) => row.date)
          .sort((a, b) => new Date(a.date) - new Date(b.date))
          .forEach((row) => {
            if (!byDate.has(row.date)) byDate.set(row.date, row);
          });
        return [...byDate.values()];
      };

      const buildSeriesSpecs = (scenarioValues, productValues, valueKey) => {
        if (valueKey === "spot") {
          return scenarioValues
            .map((scenario) => ({
              name: scenario,
              rows: deDuplicateByDate(rows.filter((row) => row.scenario === scenario && productValues.includes(row.product)))
            }))
            .filter((spec) => spec.rows.length);
        }
        const specs = [];
        scenarioValues.forEach((scenario) => {
          productValues.forEach((product) => {
            const seriesRows = rows.filter((row) => row.scenario === scenario && row.product === product);
            if (seriesRows.length) specs.push({ name: `${scenario} / ${product}`, rows: seriesRows });
          });
        });
        return specs;
      };

      const renderTradingChart = (panel) => {
        const plot = panel.querySelector(".trade-plot");
        const scenarioValues = selectedValues(panel.querySelector('[data-role="scenario"]'), scenarios);
        const productValues = selectedValues(panel.querySelector('[data-role="product"]'), products);
        const valueKey = panel.dataset.value;
        const chartKind = activeChartKind(panel);
        const freq = activeFreq(panel);
        const selected = rows.filter((row) => scenarioValues.includes(row.scenario) && productValues.includes(row.product));
        const seriesSpecs = buildSeriesSpecs(scenarioValues, productValues, valueKey);
        const title = `${panel.dataset.title} - ${filterLabel(scenarioValues, scenarios)} / ${filterLabel(productValues, products)} (${freqLabels[freq]}, ${chartKindLabels[chartKind] || chartKind})`;
        const yTitle = panel.dataset.ytitle || valueKey;
        const lineShape = panel.dataset.lineShape || "linear";
        const traces = seriesSpecs.map((spec, idx) => {
          const ohlc = aggregateOhlc(spec.rows, valueKey, freq);
          if (!ohlc.length) return null;
          const color = seriesColors[idx % seriesColors.length];
          if (chartKind === "line") {
            return {
              type: "scatter",
              name: spec.name,
              mode: "lines",
              line: { color, width: 1.8, shape: lineShape },
              x: ohlc.map((bar) => bar.x),
              y: ohlc.map((bar) => bar.close),
              hovertemplate: "%{x|%Y-%m-%d}<br>%{fullData.name}<br>Value=%{y:,.0f}<extra></extra>"
            };
          }
          return {
              type: "candlestick",
              name: spec.name,
              x: ohlc.map((bar) => bar.x),
              open: ohlc.map((bar) => bar.open),
              high: ohlc.map((bar) => bar.high),
              low: ohlc.map((bar) => bar.low),
              close: ohlc.map((bar) => bar.close),
              opacity: seriesSpecs.length > 1 ? 0.72 : 1,
              increasing: { line: { color: palette.green, width: 1.2 }, fillcolor: "rgba(8,116,67,0.2)" },
              decreasing: { line: { color: palette.red, width: 1.2 }, fillcolor: "rgba(166,56,56,0.2)" },
              hovertemplate: "%{x|%Y-%m-%d}<br>%{fullData.name}<br>O=%{open:,.0f}<br>H=%{high:,.0f}<br>L=%{low:,.0f}<br>C=%{close:,.0f}<extra></extra>"
          };
        }).filter(Boolean);
        const empty = !traces.length;
        const layout = layoutBase(empty ? `${title} - no data` : title, yTitle);
        layout.showlegend = traces.length > 1;
        if (traces.length > 1) {
          layout.margin.r = 190;
          layout.legend = {
            orientation: "v",
            x: 1.02,
            xanchor: "left",
            y: 1,
            yanchor: "top",
            font: { size: 10 }
          };
        }
        if (panel.dataset.levels === "barriers") {
          layout.shapes = [];
          layout.annotations = [];
          [
            ["S0", levels.initial_spot, "#475467", "solid"],
            ["KO 103%", levels.ko, palette.green, "dash"],
            ["KI 75%", levels.ki, palette.red, "dash"]
          ].forEach(([label, value, color, dash]) => {
            if (value === null || value === undefined || !Number.isFinite(Number(value))) return;
            layout.shapes.push({ type: "line", xref: "paper", x0: 0, x1: 1, yref: "y", y0: value, y1: value, line: { color, width: 1, dash } });
            layout.annotations.push({ xref: "paper", x: 0, yref: "y", y: value, text: label, showarrow: false, xanchor: "left", yanchor: label === "KI 75%" ? "top" : "bottom", font: { size: 11, color } });
          });
        }
        if (plot && window.Plotly) Plotly.react(plot, traces, layout, { responsive: true, displaylogo: false, scrollZoom: true });
        const caption = panel.querySelector(".chart-caption");
        const barCount = traces.reduce((count, trace) => count + (trace.x?.length || 0), 0);
        if (caption) caption.textContent = `${selected.length.toLocaleString()} daily rows, ${barCount.toLocaleString()} ${freqLabels[freq].toLowerCase()} bars across ${traces.length.toLocaleString()} series. Drag to zoom, double-click to reset.`;
      };

      const renderSummaryBars = (panel) => {
        const plot = panel.querySelector(".trade-plot");
        const scenarioValues = selectedValues(panel.querySelector('[data-role="scenario"]'), scenarios);
        const productValues = selectedValues(panel.querySelector('[data-role="product"]'), products);
        const selected = summaries.filter((row) => scenarioValues.includes(row.scenario) && productValues.includes(row.product));
        const metricKeys = (panel.dataset.metrics || "").split(",").filter(Boolean);
        const labels = selected.map((row) => `${row.scenario} / ${row.product}`);
        const metricNames = metricKeys.map((key) => key.replaceAll("_", " "));
        const metricColors = [palette.blue, palette.green, palette.amber, palette.red, "#667085"];
        const layout = {
          template: "plotly_white",
          font: { family: plotFont, size: 12, color: palette.ink },
          margin: { l: 70, r: 20, t: 64, b: 96 },
          paper_bgcolor: "rgba(0,0,0,0)",
          plot_bgcolor: palette.soft,
          title: { text: `${panel.dataset.title} - ${filterLabel(scenarioValues, scenarios)} / ${filterLabel(productValues, products)}`, x: 0, xanchor: "left", font: { size: 15 } },
          showlegend: metricKeys.length > 1,
          legend: { orientation: "h", y: 1.12, x: 0, xanchor: "left", font: { size: 10 } },
          barmode: "group",
          xaxis: { tickangle: -30, automargin: true },
          yaxis: { zeroline: true, zerolinecolor: "#95a3b8" }
        };
        const traces = selected.length
          ? metricKeys.map((key, idx) => ({
              type: "bar",
              name: metricNames[idx],
              x: labels,
              y: selected.map((row) => Number(row[key] || 0)),
              marker: { color: metricColors[idx % metricColors.length] },
              hovertemplate: "%{x}<br>%{fullData.name}: %{y:,.0f}<extra></extra>"
            }))
          : [];
        if (plot && window.Plotly) Plotly.react(plot, traces, layout, { responsive: true, displaylogo: false });
        const caption = panel.querySelector(".chart-caption");
        if (caption) {
          if (selected.length === 1) {
            caption.textContent = metricKeys.map((key, idx) => `${metricNames[idx]}: ${fmt(selected[0][key])}`).join(" | ");
          } else {
            caption.textContent = `${selected.length.toLocaleString()} runs selected across ${filterLabel(scenarioValues, scenarios)} / ${filterLabel(productValues, products)}.`;
          }
        }
      };

      document.querySelectorAll(".trade-chart").forEach((panel) => {
        const scenarioSelect = panel.querySelector('[data-role="scenario"]');
        const productSelect = panel.querySelector('[data-role="product"]');
        setOptions(scenarioSelect, scenarios, "historical");
        setOptions(productSelect, products, "PPP-DKI");
        panel.querySelectorAll("select").forEach((select) => {
          select.addEventListener("change", () => {
            const values = select.dataset.role === "scenario" ? scenarios : products;
            updateMultiSelectControl(select, values);
            panel.dataset.panel === "summary" ? renderSummaryBars(panel) : renderTradingChart(panel);
          });
        });
        installMultiSelectControl(scenarioSelect, scenarios);
        installMultiSelectControl(productSelect, products);
        panel.querySelectorAll(".freq-button").forEach((button) => {
          button.addEventListener("click", () => {
            updateButtonState(panel, button.dataset.freq);
            renderTradingChart(panel);
          });
        });
        panel.querySelectorAll(".chart-kind-button").forEach((button) => {
          button.addEventListener("click", () => {
            updateChartKindState(panel, button.dataset.kind);
            renderTradingChart(panel);
          });
        });
        updateButtonState(panel, "D");
        updateChartKindState(panel, panel.dataset.defaultKind || "candle");
        if (panel.dataset.panel === "summary") renderSummaryBars(panel);
        else renderTradingChart(panel);
      });
    })();
  </script>
""".replace("__CHART_PAYLOAD__", detail_payload.replace("</", "<\\/")).replace(
        "__SUMMARY_PAYLOAD__", summary_payload.replace("</", "<\\/")
    ).replace("__LEVELS_PAYLOAD__", levels_payload.replace("</", "<\\/"))

    detail_script = """
  <script id="daily-detail-data" type="application/json">__DETAIL_PAYLOAD__</script>
  <script>
    (() => {
      const raw = document.getElementById("daily-detail-data")?.textContent || "[]";
      const rows = JSON.parse(raw);
      const scenarioSelect = document.getElementById("detailScenario");
      const productSelect = document.getElementById("detailProduct");
      const tableHead = document.getElementById("detailTableHead");
      const tableBody = document.getElementById("detailTableBody");
      const kpiHost = document.getElementById("detailKpis");
      const plotHost = document.getElementById("detailPlot");
      const plotFont = getComputedStyle(document.documentElement).getPropertyValue("--font-body").trim() || "Aptos, Segoe UI, sans-serif";
      const columns = [
        ["date", "Date", "text"],
        ["spot", "Spot", "number"],
        ["active_contract", "Contract", "text"],
        ["futures_contracts", "Fut Pos", "number"],
        ["product_mtm", "Product MTM", "money"],
        ["hedge_mtm", "Hedge MTM", "money"],
        ["cash", "Cash", "money"],
        ["daily_product_pnl", "Daily Product", "money"],
        ["daily_hedge_pnl", "Daily Hedge", "money"],
        ["daily_total_pnl", "Daily Total", "money"],
        ["total_pnl", "Total PnL", "money"],
        ["product_delta", "Prod Delta", "number"],
        ["product_gamma", "Prod Gamma", "number4"],
        ["post_hedge_delta", "Net Delta", "number"],
        ["post_hedge_delta_cash_1pct", "Delta Cash 1%", "money"],
        ["post_hedge_gamma_cash_1pct", "Gamma Cash 1%", "money"],
        ["pricing_q", "Pricing q", "percent"],
        ["implied_q", "Implied q", "percent"],
        ["trade_abs_quantity", "Abs Trade Qty", "number"],
        ["trade_notional", "Trade Notional", "money"],
        ["trade_types", "Trades", "text"],
        ["lifecycle_actions", "Event", "text"],
        ["lifecycle_cashflow", "Event CF", "money"]
      ];

      const unique = (key) => [...new Set(rows.map((row) => row[key]).filter(Boolean))];
      const signedClass = (value) => Number(value) < 0 ? "num neg" : Number(value) > 0 ? "num pos" : "num";
      const fmt = (value, type) => {
        if (value === null || value === undefined || value === "") return "";
        if (type === "text") return String(value);
        const number = Number(value);
        if (!Number.isFinite(number)) return "";
        if (type === "percent") return `${(number * 100).toFixed(2)}%`;
        if (type === "number4") return number.toLocaleString(undefined, { maximumFractionDigits: 4, minimumFractionDigits: 4 });
        if (type === "number") return number.toLocaleString(undefined, { maximumFractionDigits: 2 });
        return number.toLocaleString(undefined, { maximumFractionDigits: 0 });
      };

      const setOptions = (select, values, preferred) => {
        select.innerHTML = values.map((value) => `<option value="${value}">${value}</option>`).join("");
        if (values.includes(preferred)) select.value = preferred;
      };

      const renderKpis = (selected) => {
        if (!selected.length) {
          kpiHost.innerHTML = "";
          return;
        }
        const last = selected[selected.length - 1];
        const worstDay = selected.reduce((best, row) => Number(row.daily_total_pnl || 0) < Number(best.daily_total_pnl || 0) ? row : best, selected[0]);
        const maxDelta = selected.reduce((best, row) => Math.abs(Number(row.post_hedge_delta_cash_1pct || 0)) > Math.abs(Number(best.post_hedge_delta_cash_1pct || 0)) ? row : best, selected[0]);
        const maxTrade = selected.reduce((sum, row) => sum + Math.abs(Number(row.trade_abs_quantity || 0)), 0);
        const events = selected.filter((row) => row.lifecycle_actions).map((row) => `${row.date}: ${row.lifecycle_actions}`).join(" | ") || "none";
        const cards = [
          ["Final Total PnL", fmt(last.total_pnl, "money"), last.total_pnl < 0 ? "bad" : "good"],
          ["Worst Daily PnL", `${fmt(worstDay.daily_total_pnl, "money")} on ${worstDay.date}`, "bad"],
          ["Max Delta Cash 1%", `${fmt(maxDelta.post_hedge_delta_cash_1pct, "money")} on ${maxDelta.date}`, "warn"],
          ["Final Futures Pos", `${fmt(last.futures_contracts, "number")} ${last.active_contract || ""}`, ""],
          ["Abs Contracts Traded", fmt(maxTrade, "number"), ""],
          ["Events", events, ""]
        ];
        kpiHost.innerHTML = cards.map(([label, value, tone]) => `<div class="metric compact ${tone}"><span>${label}</span><strong>${value}</strong></div>`).join("");
      };

      const renderPlot = (selected) => {
        if (!plotHost || !window.Plotly || !selected.length) return;
        const x = selected.map((row) => row.date);
        Plotly.react(plotHost, [
          { x, y: selected.map((row) => row.total_pnl), name: "Total PnL", type: "scatter", mode: "lines", line: { color: "#17345d", width: 2 } },
          { x, y: selected.map((row) => row.daily_total_pnl), name: "Daily PnL", type: "bar", yaxis: "y2", marker: { color: selected.map((row) => Number(row.daily_total_pnl || 0) < 0 ? "#c95b59" : "#2b8f62") }, opacity: 0.42 },
          { x, y: selected.map((row) => row.post_hedge_delta_cash_1pct), name: "Delta Cash 1%", type: "scatter", mode: "lines", line: { color: "#9a6500", width: 1.5, dash: "dot" } }
        ], {
          template: "plotly_white",
          font: { family: plotFont, size: 12, color: "#172033" },
          margin: { l: 58, r: 58, t: 36, b: 42 },
          paper_bgcolor: "rgba(0,0,0,0)",
          plot_bgcolor: "rgba(250,252,255,0.94)",
          legend: { orientation: "h", y: 1.1, x: 0 },
          xaxis: { title: "Date" },
          yaxis: { title: "Cumulative / Delta Cash" },
          yaxis2: { title: "Daily PnL", overlaying: "y", side: "right", showgrid: false }
        }, { responsive: true, displaylogo: false });
      };

      const renderTable = (selected) => {
        tableHead.innerHTML = columns.map(([, label]) => `<th>${label}</th>`).join("");
        tableBody.innerHTML = selected.map((row) => {
          const cells = columns.map(([key, , type]) => {
            const cls = type === "text" ? "" : signedClass(row[key]);
            return `<td class="${cls}">${fmt(row[key], type)}</td>`;
          }).join("");
          return `<tr>${cells}</tr>`;
        }).join("");
      };

      const render = () => {
        const selected = rows.filter((row) => row.scenario === scenarioSelect.value && row.product === productSelect.value);
        renderKpis(selected);
        renderPlot(selected);
        renderTable(selected);
      };

      if (rows.length && scenarioSelect && productSelect) {
        setOptions(scenarioSelect, unique("scenario"), "historical");
        setOptions(productSelect, unique("product"), "PPP-DKI");
        scenarioSelect.addEventListener("change", render);
        productSelect.addEventListener("change", render);
        render();
      }
    })();
  </script>
""".replace("__DETAIL_PAYLOAD__", detail_payload.replace("</", "<\\/"))

    observation_schedule_payload_text = json.dumps(
        observation_schedules,
        separators=(",", ":"),
    )
    observation_schedule_script = """
  <script id="observation-schedules-data" type="application/json">__OBSERVATION_SCHEDULES_PAYLOAD__</script>
  <script>
    (() => {
      const schedules = JSON.parse(document.getElementById("observation-schedules-data")?.textContent || "{}");
      const modal = document.getElementById("observationScheduleModal");
      const title = document.getElementById("observationScheduleTitle");
      const summary = document.getElementById("observationScheduleSummary");
      const body = document.getElementById("observationScheduleBody");
      const pageLabel = document.getElementById("observationSchedulePage");
      const prevButton = document.getElementById("observationSchedulePrev");
      const nextButton = document.getElementById("observationScheduleNext");
      const pageSize = 30;
      let activeKey = "";
      let activePage = 0;

      const render = () => {
        const dates = schedules[activeKey] || [];
        const pageCount = Math.max(1, Math.ceil(dates.length / pageSize));
        activePage = Math.min(Math.max(activePage, 0), pageCount - 1);
        const start = activePage * pageSize;
        const visibleDates = dates.slice(start, start + pageSize);
        if (title) title.textContent = activeKey || "Observation Schedule";
        if (summary) summary.textContent = `${dates.length.toLocaleString()} observation dates`;
        if (pageLabel) pageLabel.textContent = `Page ${activePage + 1} / ${pageCount}`;
        if (prevButton) prevButton.disabled = activePage <= 0;
        if (nextButton) nextButton.disabled = activePage >= pageCount - 1;
        if (body) {
          body.innerHTML = visibleDates.map((date, idx) => `
            <tr>
              <td class="num">${start + idx + 1}</td>
              <td>${date}</td>
            </tr>
          `).join("") || '<tr><td colspan="2">No schedule dates</td></tr>';
        }
      };

      const open = (key) => {
        activeKey = key;
        activePage = 0;
        render();
        if (modal) {
          modal.hidden = false;
          modal.querySelector(".modal-close")?.focus();
        }
      };

      const close = () => {
        if (modal) modal.hidden = true;
      };

      document.querySelectorAll(".obs-schedule-button").forEach((button) => {
        button.addEventListener("click", () => open(button.dataset.observationKey || ""));
      });
      prevButton?.addEventListener("click", () => {
        activePage -= 1;
        render();
      });
      nextButton?.addEventListener("click", () => {
        activePage += 1;
        render();
      });
      modal?.querySelectorAll("[data-close-modal]").forEach((button) => {
        button.addEventListener("click", close);
      });
      modal?.addEventListener("click", (event) => {
        if (event.target === modal) close();
      });
      document.addEventListener("keydown", (event) => {
        if (event.key === "Escape" && modal && !modal.hidden) close();
      });
    })();
  </script>
""".replace("__OBSERVATION_SCHEDULES_PAYLOAD__", observation_schedule_payload_text.replace("</", "<\\/"))

    plotly_divs = [
        plot_div(pnl_heatmap, include_plotlyjs=True),
        plot_div(drawdown_heatmap),
        plot_div(pnl_trajectory),
        plot_div(attribution),
        plot_div(hedge_contracts),
        plot_div(hedge_turnover),
        plot_div(delta_cash),
        plot_div(gamma_cash),
        plot_div(lifecycle_fig),
        plot_div(price_paths),
    ]

    html = f"""<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>PPP/NPP Snowball Backtest Case Study</title>
  <style>
    :root {{
      --ink: #172033;
      --muted: #667085;
      --line: #d9e1ee;
      --panel: #ffffff;
      --soft: #f5f7fb;
      --navy: #17345d;
      --blue: #4267ac;
      --green: #087443;
      --red: #a63838;
      --amber: #9a6500;
      --font-body: "IBM Plex Sans", "Source Sans 3", Aptos, "Segoe UI Variable", "Segoe UI", sans-serif;
      --font-display: "Source Sans 3", "IBM Plex Sans", "Aptos Display", Aptos, "Segoe UI Variable", "Segoe UI", sans-serif;
      --font-number: "IBM Plex Mono", "Cascadia Mono", "SFMono-Regular", Menlo, Consolas, monospace;
    }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0;
      font-family: var(--font-body);
      font-size: 14px;
      line-height: 1.5;
      text-rendering: optimizeLegibility;
      -webkit-font-smoothing: antialiased;
      font-feature-settings: "kern" 1;
      color: var(--ink);
      background: linear-gradient(180deg, #eef3f9 0%, #f8fafc 48%, #eef2f6 100%);
    }}
    header {{
      background: #10233e;
      color: #fff;
      padding: 28px 34px 22px;
      border-bottom: 4px solid #d5a84f;
    }}
    header h1 {{
      margin: 0 0 8px;
      font-family: var(--font-display);
      font-size: 30px;
      letter-spacing: 0;
      font-weight: 760;
    }}
    header p {{ margin: 0; color: #d9e4f2; max-width: 1180px; line-height: 1.45; }}
    nav {{
      position: sticky;
      top: 0;
      z-index: 10;
      display: flex;
      gap: 10px;
      align-items: center;
      padding: 10px 34px;
      background: rgba(248, 250, 252, 0.94);
      border-bottom: 1px solid var(--line);
      backdrop-filter: blur(10px);
    }}
    nav a {{
      color: var(--navy);
      text-decoration: none;
      font-size: 13px;
      font-weight: 680;
      padding: 7px 10px;
      border: 1px solid transparent;
    }}
    nav a:hover {{ border-color: var(--line); background: #fff; }}
    main {{ padding: 24px 34px 44px; }}
    section {{ margin: 0 auto 28px; max-width: 1480px; }}
    h2 {{
      margin: 0 0 12px;
      font-family: var(--font-display);
      font-size: 22px;
      line-height: 1.2;
      font-weight: 760;
    }}
    h3 {{
      margin: 0 0 10px;
      font-family: var(--font-display);
      font-size: 15px;
      line-height: 1.25;
      font-weight: 720;
      color: var(--navy);
    }}
    .subtle {{ color: var(--muted); font-size: 13.5px; line-height: 1.55; margin: 0 0 14px; }}
    .metrics {{
      display: grid;
      grid-template-columns: repeat(6, minmax(150px, 1fr));
      gap: 12px;
      margin: 16px 0 18px;
    }}
    .metric {{
      background: var(--panel);
      border: 1px solid var(--line);
      padding: 14px 14px 12px;
      min-height: 98px;
    }}
    .metric span {{ display: block; color: var(--muted); font-size: 12px; font-weight: 680; margin-bottom: 8px; }}
    .metric strong {{
      display: block;
      font-family: var(--font-number);
      font-size: 23px;
      line-height: 1.12;
      font-weight: 720;
      font-variant-numeric: tabular-nums;
    }}
    .metric small {{ display: block; color: var(--muted); margin-top: 8px; line-height: 1.35; }}
    .metric.good strong {{ color: var(--green); }}
    .metric.bad strong {{ color: var(--red); }}
    .metric.warn strong {{ color: var(--amber); }}
    .grid-2 {{ display: grid; grid-template-columns: minmax(0, 1fr) minmax(0, 1fr); gap: 16px; }}
    .grid-3 {{ display: grid; grid-template-columns: repeat(3, minmax(0, 1fr)); gap: 16px; }}
    .panel {{
      background: rgba(255,255,255,0.96);
      border: 1px solid var(--line);
      padding: 16px;
      box-shadow: 0 12px 32px rgba(28, 47, 79, 0.06);
      overflow: hidden;
    }}
    .plot {{ min-height: 420px; }}
    .trade-chart {{
      display: grid;
      grid-template-rows: auto minmax(420px, 1fr) auto;
      gap: 10px;
      min-height: 520px;
    }}
    .chart-toolbar {{
      display: flex;
      flex-wrap: wrap;
      align-items: end;
      gap: 12px;
      padding: 10px;
      background: #eef3f8;
      border: 1px solid var(--line);
    }}
    .chart-toolbar label {{
      display: grid;
      gap: 5px;
      color: var(--muted);
      font-size: 12px;
      font-weight: 680;
    }}
    .chart-toolbar .filter-label {{
      min-width: 210px;
      position: relative;
    }}
    .chart-toolbar select {{
      min-width: 190px;
      padding: 8px 10px;
      border: 1px solid #bcc8d9;
      background: #fff;
      color: var(--ink);
      font: inherit;
      font-size: 13.5px;
    }}
    .chart-toolbar select[multiple] {{
      position: absolute;
      width: 1px;
      height: 1px;
      opacity: 0;
      pointer-events: none;
    }}
    .multi-select {{ position: relative; min-width: 210px; }}
    .multi-select-button {{
      width: 100%;
      min-height: 38px;
      border: 1px solid #bcc8d9;
      background: #fff;
      color: var(--ink);
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 10px;
      padding: 8px 10px;
      font: inherit;
      font-size: 13.5px;
      font-weight: 680;
      cursor: pointer;
      text-align: left;
    }}
    .multi-select-button::after {{
      content: "";
      width: 0;
      height: 0;
      border-left: 4px solid transparent;
      border-right: 4px solid transparent;
      border-top: 5px solid currentColor;
      flex: 0 0 auto;
    }}
    .multi-select.open .multi-select-button::after {{
      transform: rotate(180deg);
    }}
    .multi-select-text {{
      overflow: hidden;
      text-overflow: ellipsis;
      white-space: nowrap;
    }}
    .multi-select-menu {{
      display: none;
      position: absolute;
      top: calc(100% + 4px);
      left: 0;
      z-index: 40;
      width: 100%;
      max-height: 245px;
      overflow: auto;
      border: 1px solid #a9b8cd;
      background: #fff;
      box-shadow: 0 18px 36px rgba(25, 42, 70, 0.18);
      padding: 6px;
    }}
    .multi-select.open .multi-select-menu {{ display: grid; gap: 2px; }}
    .multi-select-actions {{
      display: flex;
      gap: 6px;
      padding: 0 0 6px;
      margin-bottom: 4px;
      border-bottom: 1px solid #e1e7f0;
    }}
    .multi-select-action {{
      border: 1px solid #c7d2e2;
      background: #f6f8fb;
      color: var(--navy);
      padding: 4px 8px;
      font: inherit;
      font-size: 11px;
      font-weight: 760;
      cursor: pointer;
    }}
    .multi-option {{
      display: flex;
      align-items: center;
      gap: 7px;
      padding: 6px 5px;
      color: var(--ink);
      font-size: 12.5px;
      font-weight: 640;
      cursor: pointer;
    }}
    .multi-option:hover {{ background: #eef3f8; }}
    .multi-option input {{ margin: 0; }}
    .segmented {{
      display: inline-grid;
      grid-auto-flow: column;
      border: 1px solid #bcc8d9;
      background: #fff;
    }}
    .freq-button,
    .chart-kind-button {{
      min-width: 42px;
      border: 0;
      border-right: 1px solid #d9e1ee;
      padding: 8px 10px;
      background: #fff;
      color: var(--navy);
      font: inherit;
      font-size: 12px;
      font-weight: 760;
      cursor: pointer;
    }}
    .chart-kind-button {{
      min-width: 62px;
    }}
    .freq-button:last-child,
    .chart-kind-button:last-child {{ border-right: 0; }}
    .freq-button.active,
    .chart-kind-button.active {{
      background: var(--navy);
      color: #fff;
    }}
    .trade-plot {{ min-height: 390px; }}
    .chart-caption {{
      color: var(--muted);
      font-size: 12.5px;
      line-height: 1.45;
      min-height: 18px;
    }}
    .detail-controls {{
      display: flex;
      flex-wrap: wrap;
      gap: 12px;
      align-items: end;
      margin-bottom: 14px;
      padding: 12px;
      background: #eef3f8;
      border: 1px solid var(--line);
    }}
    .detail-controls label {{
      display: grid;
      gap: 5px;
      color: var(--muted);
      font-size: 12px;
      font-weight: 680;
    }}
    .detail-controls select {{
      min-width: 210px;
      padding: 8px 10px;
      border: 1px solid #bcc8d9;
      background: #fff;
      color: var(--ink);
      font: inherit;
      font-size: 13.5px;
    }}
    .metric.compact {{
      min-height: 72px;
      padding: 12px;
    }}
    .metric.compact strong {{
      font-size: 16px;
      line-height: 1.22;
      overflow-wrap: anywhere;
    }}
    .detail-plot {{ min-height: 360px; margin: 8px 0 14px; }}
    .detail-scroll {{ max-height: 560px; }}
    #detailTable th {{ top: 0; }}
    #detailTable td {{ font-size: 11.5px; }}
    table {{
      border-collapse: collapse;
      width: 100%;
      margin: 8px 0 0;
      background: #fff;
    }}
    th, td {{
      border-bottom: 1px solid #e4e9f2;
      padding: 7px 9px;
      font-size: 12.5px;
      line-height: 1.35;
      text-align: left;
      vertical-align: top;
      white-space: nowrap;
    }}
    th {{
      background: #eef3f8;
      color: #33415c;
      font-family: var(--font-display);
      font-weight: 760;
      position: sticky;
      top: 42px;
    }}
    .table-scroll {{ max-height: 430px; overflow: auto; border: 1px solid var(--line); }}
    .num {{
      text-align: right;
      font-family: var(--font-number);
      font-variant-numeric: tabular-nums;
      letter-spacing: 0;
    }}
    .term-name {{
      font-family: var(--font-number);
      font-size: 12px;
    }}
    .term-value-action {{
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 12px;
      min-width: 420px;
    }}
    .term-value-action span {{
      white-space: normal;
      line-height: 1.4;
    }}
    .obs-schedule-button {{
      flex: 0 0 auto;
      border: 1px solid #bcc8d9;
      background: #fff;
      color: var(--navy);
      padding: 5px 9px;
      font: inherit;
      font-size: 12px;
      font-weight: 760;
      cursor: pointer;
    }}
    .obs-schedule-button:hover {{
      background: #eef3f8;
    }}
    .modal-backdrop[hidden] {{ display: none; }}
    .modal-backdrop {{
      position: fixed;
      inset: 0;
      z-index: 100;
      display: grid;
      place-items: center;
      padding: 28px;
      background: rgba(16, 35, 62, 0.46);
    }}
    .modal-dialog {{
      width: min(720px, 100%);
      max-height: min(760px, 90vh);
      display: grid;
      grid-template-rows: auto auto minmax(0, 1fr) auto;
      background: #fff;
      border: 1px solid #a9b8cd;
      box-shadow: 0 28px 70px rgba(10, 28, 52, 0.28);
    }}
    .modal-header {{
      display: flex;
      align-items: start;
      justify-content: space-between;
      gap: 14px;
      padding: 16px 18px 10px;
      border-bottom: 1px solid var(--line);
    }}
    .modal-header h3 {{ margin: 0; }}
    .modal-close {{
      border: 1px solid #bcc8d9;
      background: #fff;
      color: var(--ink);
      width: 30px;
      height: 30px;
      font: inherit;
      font-size: 18px;
      line-height: 1;
      cursor: pointer;
    }}
    .modal-summary {{
      margin: 0;
      padding: 10px 18px;
      color: var(--muted);
      font-size: 12.5px;
    }}
    .modal-table-wrap {{
      overflow: auto;
      margin: 0 18px;
      border: 1px solid var(--line);
    }}
    .modal-table-wrap table {{ margin: 0; }}
    .modal-table-wrap th {{ top: 0; }}
    .modal-footer {{
      display: flex;
      justify-content: flex-end;
      align-items: center;
      gap: 10px;
      padding: 12px 18px 16px;
    }}
    .modal-footer button {{
      border: 1px solid #bcc8d9;
      background: #fff;
      color: var(--navy);
      min-width: 80px;
      padding: 7px 10px;
      font: inherit;
      font-size: 12px;
      font-weight: 760;
      cursor: pointer;
    }}
    .modal-footer button:disabled {{
      color: #98a2b3;
      cursor: not-allowed;
    }}
    .modal-page-label {{
      color: var(--muted);
      font-size: 12.5px;
      min-width: 92px;
      text-align: center;
    }}
    .pos {{ color: var(--green); }}
    .neg {{ color: var(--red); }}
    .empty-state {{ color: var(--muted); padding: 14px; border: 1px dashed var(--line); background: #fff; }}
    .caption {{ color: var(--muted); font-size: 12px; margin-top: 8px; }}
    footer {{ max-width: 1480px; margin: 0 auto; color: var(--muted); font-size: 12px; padding: 18px 34px 36px; }}
    @media (max-width: 1100px) {{
      .metrics {{ grid-template-columns: repeat(2, minmax(0, 1fr)); }}
      .grid-2, .grid-3 {{ grid-template-columns: 1fr; }}
      nav {{ flex-wrap: wrap; }}
    }}
  </style>
</head>
<body>
  <header>
    <h1>PPP/NPP Snowball Backtest Control Dashboard</h1>
    <p>Generated {esc(generated_at)}. Underlying {esc(terms_map.get("underlying", ""))}, hedge {esc(terms_map.get("hedging_instrument", ""))}, notional {fmt_number(terms_map.get("notional", 0.0))}, KO {fmt_percent(terms_map.get("ko_ratio", 0.0))}, KI {fmt_percent(terms_map.get("ki_ratio", 0.0))}, pricing q {fmt_percent(terms_map.get("dividend_yield", 0.0))}.</p>
  </header>
  <nav>
    <a href="#executive">General Manager</a>
    <a href="#price-paths">Price Paths</a>
    <a href="#trader">Trader</a>
    <a href="#risk">Risk Manager</a>
    <a href="#daily-detail">Daily Detail</a>
    <a href="#events">Events</a>
    <a href="#terms">Terms</a>
  </nav>
  <main>
    <section id="executive">
      <h2>General Manager View</h2>
      <p class="subtle">Portfolio-level comparison across products and scenarios: which structure makes or loses money, where stress is concentrated, and whether lifecycle outcomes are clean.</p>
      <div class="metrics">{executive_cards}</div>
      <div class="grid-2">
        <div class="panel plot">{plotly_divs[0]}</div>
        <div class="panel plot">{plotly_divs[1]}</div>
      </div>
      <div class="grid-2" style="margin-top:16px;">
        <div class="panel trade-chart" data-value="total_pnl" data-title="Total PnL K-Line" data-ytitle="Total PnL">
          <div class="chart-toolbar">
            <label>Scenario<select data-role="scenario" multiple size="5" title="Use Ctrl/Cmd-click for multiple selections"></select></label>
            <label>Product<select data-role="product" multiple size="3" title="Use Ctrl/Cmd-click for multiple selections"></select></label>
            <div class="segmented">
              <button class="freq-button active" type="button" data-freq="D">D</button>
              <button class="freq-button" type="button" data-freq="W">W</button>
              <button class="freq-button" type="button" data-freq="M">M</button>
            </div>
            <div class="segmented">
              <button class="chart-kind-button active" type="button" data-kind="candle">Candle</button>
              <button class="chart-kind-button" type="button" data-kind="line">Line</button>
            </div>
          </div>
          <div class="trade-plot"></div>
          <div class="chart-caption"></div>
        </div>
        <div class="panel trade-chart" data-panel="summary" data-title="Final Product vs Hedge PnL" data-metrics="final_product_pnl,final_hedge_pnl,final_total_pnl,worst_drawdown">
          <div class="chart-toolbar">
            <label>Scenario<select data-role="scenario" multiple size="5" title="Use Ctrl/Cmd-click for multiple selections"></select></label>
            <label>Product<select data-role="product" multiple size="3" title="Use Ctrl/Cmd-click for multiple selections"></select></label>
          </div>
          <div class="trade-plot"></div>
          <div class="chart-caption"></div>
        </div>
      </div>
      <div class="panel" style="margin-top:16px;">
        <h3>Run-Level PnL and Lifecycle Summary</h3>
        <div class="table-scroll">{executive_table}</div>
      </div>
    </section>
    <section id="price-paths">
      <h2>Scenario Price Paths</h2>
      <p class="subtle">Underlying spot paths for the historical, rapid-up, rapid-down, high-oscillation, and low-oscillation scenarios, with initial spot, KO, and KI reference levels.</p>
      <div class="panel trade-chart" data-value="spot" data-title="Underlying Price K-Line" data-ytitle="000852.SH spot" data-levels="barriers">
        <div class="chart-toolbar">
          <label>Scenario<select data-role="scenario" multiple size="5" title="Use Ctrl/Cmd-click for multiple selections"></select></label>
          <label>Product<select data-role="product" multiple size="3" title="Use Ctrl/Cmd-click for multiple selections"></select></label>
          <div class="segmented">
            <button class="freq-button active" type="button" data-freq="D">D</button>
            <button class="freq-button" type="button" data-freq="W">W</button>
            <button class="freq-button" type="button" data-freq="M">M</button>
          </div>
          <div class="segmented">
            <button class="chart-kind-button active" type="button" data-kind="candle">Candle</button>
            <button class="chart-kind-button" type="button" data-kind="line">Line</button>
          </div>
        </div>
        <div class="trade-plot"></div>
        <div class="chart-caption"></div>
      </div>
    </section>
    <section id="trader">
      <h2>Trader View</h2>
      <p class="subtle">Hedge inventory, turnover, roll load, and data-quality exceptions that directly affect desk execution and explain realized hedge PnL.</p>
      <div class="grid-2">
        <div class="panel trade-chart" data-value="futures_contracts" data-title="Futures Inventory K-Line" data-ytitle="Contracts" data-default-kind="line" data-line-shape="hv">
          <div class="chart-toolbar">
            <label>Scenario<select data-role="scenario" multiple size="5" title="Use Ctrl/Cmd-click for multiple selections"></select></label>
            <label>Product<select data-role="product" multiple size="3" title="Use Ctrl/Cmd-click for multiple selections"></select></label>
            <div class="segmented">
              <button class="freq-button active" type="button" data-freq="D">D</button>
              <button class="freq-button" type="button" data-freq="W">W</button>
              <button class="freq-button" type="button" data-freq="M">M</button>
            </div>
            <div class="segmented">
              <button class="chart-kind-button" type="button" data-kind="candle">Candle</button>
              <button class="chart-kind-button active" type="button" data-kind="line">Line</button>
            </div>
          </div>
          <div class="trade-plot"></div>
          <div class="chart-caption"></div>
        </div>
        <div class="panel trade-chart" data-panel="summary" data-title="Hedge Workload" data-metrics="gross_trade_notional,max_trade_notional,num_trades,roll_count">
          <div class="chart-toolbar">
            <label>Scenario<select data-role="scenario" multiple size="5" title="Use Ctrl/Cmd-click for multiple selections"></select></label>
            <label>Product<select data-role="product" multiple size="3" title="Use Ctrl/Cmd-click for multiple selections"></select></label>
          </div>
          <div class="trade-plot"></div>
          <div class="chart-caption"></div>
        </div>
      </div>
      <div class="grid-2" style="margin-top:16px;">
        <div class="panel">
          <h3>Hedge Workload by Run</h3>
          <div class="table-scroll">{trader_summary_table}</div>
        </div>
        <div class="panel">
          <h3>Roll/Data Exceptions</h3>
          <div class="table-scroll">{trader_exceptions_table}</div>
          <p class="caption">Rows with reason `futures_roll_missing_old_contract` used the selected active futures price as proxy close because the held old contract was absent from that date's futures chain.</p>
        </div>
      </div>
    </section>
    <section id="risk">
      <h2>Risk Manager View</h2>
      <p class="subtle">Residual delta and gamma cash exposure after hedging, largest drawdowns, and loss days by product/scenario.</p>
      <div class="grid-2">
        <div class="panel trade-chart" data-value="post_hedge_delta_cash_1pct" data-title="Residual Delta Cash K-Line" data-ytitle="Cash PnL / 1%">
          <div class="chart-toolbar">
            <label>Scenario<select data-role="scenario" multiple size="5" title="Use Ctrl/Cmd-click for multiple selections"></select></label>
            <label>Product<select data-role="product" multiple size="3" title="Use Ctrl/Cmd-click for multiple selections"></select></label>
            <div class="segmented">
              <button class="freq-button active" type="button" data-freq="D">D</button>
              <button class="freq-button" type="button" data-freq="W">W</button>
              <button class="freq-button" type="button" data-freq="M">M</button>
            </div>
            <div class="segmented">
              <button class="chart-kind-button active" type="button" data-kind="candle">Candle</button>
              <button class="chart-kind-button" type="button" data-kind="line">Line</button>
            </div>
          </div>
          <div class="trade-plot"></div>
          <div class="chart-caption"></div>
        </div>
        <div class="panel trade-chart" data-value="post_hedge_gamma_cash_1pct" data-title="Gamma Cash K-Line" data-ytitle="Cash Gamma / 1%">
          <div class="chart-toolbar">
            <label>Scenario<select data-role="scenario" multiple size="5" title="Use Ctrl/Cmd-click for multiple selections"></select></label>
            <label>Product<select data-role="product" multiple size="3" title="Use Ctrl/Cmd-click for multiple selections"></select></label>
            <div class="segmented">
              <button class="freq-button active" type="button" data-freq="D">D</button>
              <button class="freq-button" type="button" data-freq="W">W</button>
              <button class="freq-button" type="button" data-freq="M">M</button>
            </div>
            <div class="segmented">
              <button class="chart-kind-button active" type="button" data-kind="candle">Candle</button>
              <button class="chart-kind-button" type="button" data-kind="line">Line</button>
            </div>
          </div>
          <div class="trade-plot"></div>
          <div class="chart-caption"></div>
        </div>
      </div>
      <div class="grid-2" style="margin-top:16px;">
        <div class="panel">
          <h3>Stress Ranking</h3>
          <div class="table-scroll">{risk_summary_table}</div>
        </div>
        <div class="panel">
          <h3>Worst Daily PnL Moves</h3>
          <div class="table-scroll">{loss_days_table}</div>
        </div>
      </div>
    </section>
    <section id="daily-detail">
      <h2>Daily Detail</h2>
      <div class="panel">
        <div class="detail-controls">
          <label>Scenario
            <select id="detailScenario"></select>
          </label>
          <label>Product
            <select id="detailProduct"></select>
          </label>
        </div>
        <div class="metrics" id="detailKpis"></div>
        <div id="detailPlot" class="detail-plot"></div>
        <div class="table-scroll detail-scroll">
          <table id="detailTable">
            <thead><tr id="detailTableHead"></tr></thead>
            <tbody id="detailTableBody"></tbody>
          </table>
        </div>
      </div>
    </section>
    <section id="events">
      <h2>Lifecycle Events</h2>
      <p class="subtle">KO, KI, and maturity actions are shown with before/after state so event ordering can be audited.</p>
      <div class="grid-2">
        <div class="panel plot">{plotly_divs[8]}</div>
        <div class="panel">
          <h3>Lifecycle Blotter</h3>
          <div class="table-scroll">{lifecycle_table}</div>
        </div>
      </div>
    </section>
    <section id="terms">
      <h2>Terms and Fair Coupons</h2>
      <div class="panel">{terms_html}</div>
    </section>
  </main>
  <div class="modal-backdrop" id="observationScheduleModal" hidden>
    <div class="modal-dialog" role="dialog" aria-modal="true" aria-labelledby="observationScheduleTitle">
      <div class="modal-header">
        <h3 id="observationScheduleTitle">Observation Schedule</h3>
        <button class="modal-close" type="button" aria-label="Close" data-close-modal>&times;</button>
      </div>
      <p class="modal-summary" id="observationScheduleSummary"></p>
      <div class="modal-table-wrap">
        <table>
          <thead><tr><th>#</th><th>Date</th></tr></thead>
          <tbody id="observationScheduleBody"></tbody>
        </table>
      </div>
      <div class="modal-footer">
        <button id="observationSchedulePrev" type="button">Prev</button>
        <span class="modal-page-label" id="observationSchedulePage"></span>
        <button id="observationScheduleNext" type="button">Next</button>
        <button type="button" data-close-modal>Close</button>
      </div>
    </div>
  </div>
  <footer>Source tables: `data/summary_metrics.csv`, `data/daily_pnl.csv`, `data/daily_greeks.csv`, `data/hedge_actions.csv`, `data/lifecycle_events.csv`.</footer>
{trade_chart_script}
{detail_script}
{observation_schedule_script}
</body>
</html>"""
    path.write_text(html, encoding="utf-8")
    return path


def _excel_safe_value(value: Any) -> Any:
    if isinstance(value, (np.integer, np.floating)):
        return value.item()
    if isinstance(value, (pd.Timestamp, datetime)):
        return str(value)
    return value


def run_case_study(args: argparse.Namespace) -> dict[str, Any]:
    terms = CaseStudyTerms(
        notional=float(args.notional),
        ko_ratio=float(args.ko),
        ki_ratio=float(args.ki),
        ppp_protection_rate=float(args.ppp_protection),
        rate=float(args.rate),
        dividend_yield=float(args.dividend_yield),
        volatility=float(args.volatility),
    )
    if terms.notional <= 0:
        raise CaseStudyError("notional must be positive")
    if not 0 <= terms.ppp_protection_rate <= 1:
        raise CaseStudyError("ppp-protection must be in [0, 1]")

    output_dir = Path(args.output_dir)
    cache_dir = output_dir / "cache"
    frames = load_market_cache_or_fetch(
        cache_dir=cache_dir,
        refresh_data=bool(args.refresh_data),
        cache_only=bool(args.cache_only),
        synthetic_only=bool(args.synthetic_only),
        history_years=float(args.history_years),
        start_date=args.start_date,
        end_date=args.end_date,
        terms=terms,
        scenario_days=int(args.scenario_days or 756),
    )
    start, end = latest_window(
        frames.spot_data,
        history_years=float(args.history_years),
        start_date=args.start_date,
        end_date=args.end_date,
    )
    base_spot = restrict_window(frames.spot_data, start, end)
    if args.scenario_days and int(args.scenario_days) > 0:
        base_spot = base_spot.head(int(args.scenario_days)).reset_index(drop=True)
    if len(base_spot) < 3:
        raise CaseStudyError("At least three spot observations are required")

    validate_quad_grid_for_case(args, base_spot)
    scenarios = generate_scenario_spot_data(base_spot)
    engine_config = create_engine_config(args)
    initial_spot = float(base_spot["spot"].iloc[0])
    issue_date = pd.Timestamp(base_spot["date"].iloc[0]).normalize()
    coupons, coupon_results = solve_case_study_coupons(
        dates=base_spot["date"],
        initial_spot=initial_spot,
        issue_date=issue_date,
        terms=terms,
        engine_config=engine_config,
    )

    run_results: dict[tuple[str, str], Any] = {}
    summaries: list[dict[str, Any]] = []
    for scenario, spot_data in scenarios.items():
        futures = (
            restrict_window(frames.futures_data, start, end)
            if scenario == "historical" and not frames.futures_data.empty
            else synthetic_futures_chain(spot_data, terms.rate, terms.dividend_yield)
        )
        if args.scenario_days and int(args.scenario_days) > 0:
            futures = futures[futures["date"].isin(spot_data["date"])].reset_index(drop=True)
        market_data = build_market_dataset(spot_data, futures, terms)
        products = build_case_study_products(
            initial_spot=initial_spot,
            issue_date=issue_date,
            dates=base_spot["date"],
            terms=terms,
            coupons=coupons,
        )
        for product_label, product in products.items():
            print(f"[run] {scenario} / {product_label}")
            results, summary = run_single_backtest(
                product_label=product_label,
                product=product,
                scenario=scenario,
                market_data=market_data,
                terms=terms,
                engine_config=engine_config,
                args=args,
                output_dir=output_dir,
            )
            run_results[(scenario, product_label)] = results
            summaries.append({key: _excel_safe_value(value) for key, value in summary.items()})

    consolidated = build_consolidated_frames(run_results)
    term_table = terms_frame(
        terms,
        coupons,
        initial_spot,
        coupon_results,
        issue_date=issue_date,
        dates=base_spot["date"],
    )
    write_csv_outputs(output_dir, summaries, consolidated)
    workbook_path = output_dir / "case_study_results.xlsx"
    write_excel_output(workbook_path, summaries, term_table, consolidated)
    charts = generate_charts(output_dir, summaries, consolidated, term_table)
    docx_path = write_docx_report(output_dir / "case_study_report.docx", summaries, term_table, charts)
    html_path = write_html_dashboard(output_dir / "case_study_dashboard.html", summaries, term_table, consolidated)

    manifest = {
        "output_dir": str(output_dir),
        "workbook": str(workbook_path),
        "docx": str(docx_path),
        "html": str(html_path),
        "charts": {key: str(value) for key, value in charts.items()},
        "coupons": coupons,
        "coupon_pv": {key: value.pv for key, value in coupon_results.items()},
        "initial_spot": initial_spot,
        "pv_convention": "principal_excluded_zero_upfront",
        "num_runs": len(summaries),
    }
    output_dir.mkdir(parents=True, exist_ok=True)
    (output_dir / "manifest.json").write_text(
        json.dumps(manifest, indent=2, default=str), encoding="utf-8"
    )
    return manifest


def main(argv: Optional[list[str]] = None) -> None:
    args = parse_args(argv)
    try:
        manifest = run_case_study(args)
    except (CaseStudyError, ValidationError) as exc:
        raise SystemExit(str(exc)) from exc
    print("\nPPP/NPP Snowball Backtest Case Study")
    print("=" * 72)
    print(f"Runs:     {manifest['num_runs']}")
    print(f"Output:   {manifest['output_dir']}")
    print(f"Workbook: {manifest['workbook']}")
    print(f"DOCX:     {manifest['docx']}")
    print(f"HTML:     {manifest['html']}")


if __name__ == "__main__":
    main()
