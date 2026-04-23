"""
Bilingual risk comparison report for normalized snowball structures.

This module compares matched snowball structures under exact discrete
observation schedules and produces a bilingual DOCX report.
"""

from __future__ import annotations

import argparse
import importlib.util
import os
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, Iterable, Mapping, Optional, Sequence

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from asset.equity.engine.base_engine import BaseEngine
from asset.equity.engine.mc.snowball_mc_engine import SnowballMCEngine
from asset.equity.param import EngineParams, MCParams, PDEParams, QuadParams
from asset.equity.product.option.snowball_config import (
    AccrualConfig,
    BarrierConfig,
    PayoffConfig,
)
from asset.equity.product.option.snowball_helpers import (
    create_european_ki_snowball,
    create_parachute_snowball,
)
from asset.equity.product.option.snowball_option import SnowballOption
from asset.equity.report.autocallable_risk_report import (
    _barrier_distance_metrics,
    _clone_env,
    _scale_vol_surface,
    _select_snowball_pricing_engine,
    _shift_dividend_yield,
)
from asset.equity.report.plotting import save_heatmap
from asset.equity.report.term_structure import SkewSmileVolSurface
from asset.equity.riskmeasures.greeks_calculator import GreeksCalculator
from param import FlatVolSurface, SpotQuote
from param.div import ContinuousDividendYield
from param.rrf import FlatRateCurve
from priceenv import PricingEnvironment
from util.calendar import (
    BusinessDayConvention,
    CalendarType,
    create_calendar,
)
from util.enum import CouponPayType, ObservationType, ProtectionType
from util.enum.engine_enums import EngineType, MonteCarloMethod
from util.exceptions import ValidationError


@dataclass(frozen=True)
class SnowballRiskComparisonConfig:
    valuation_date: datetime = datetime(2026, 1, 5)
    output_dir: Path = Path("output/doc/snowball_risk_comparison")
    report_filename: str = "snowball_risk_comparison_bilingual.docx"
    bilingual_layout: str = "english_then_chinese"
    initial_price: float = 100.0
    strike: float = 100.0
    tenor_months: int = 12
    ko_start_month: int = 3
    ko_barrier: float = 103.0
    ki_barrier: float = 70.0
    annual_coupon: float = 0.12
    rate: float = 0.02
    dividend_yield: float = 0.10
    volatility: float = 0.22
    business_days_in_year: int = 244
    calendar_type: CalendarType = CalendarType.CHINA
    protection_rate: float = 0.20
    contract_multiplier: float = 1.0
    num_paths: int = 20000
    seed: int = 42
    mc_method: MonteCarloMethod = MonteCarloMethod.PSEUDO
    engine_preference: Sequence[str] = ("quad", "pde", "mc")
    quad_params: Optional[QuadParams] = field(
        default_factory=lambda: QuadParams(grid_points=301)
    )
    pde_params: Optional[PDEParams] = field(default_factory=PDEParams)
    stress_spot_shocks: Sequence[float] = (
        -0.20,
        -0.15,
        -0.10,
        -0.05,
        0.00,
        0.05,
        0.10,
    )
    stress_vol_shocks: Sequence[float] = (-0.10, -0.05, 0.00, 0.05, 0.10)
    stress_div_shifts: Sequence[float] = (-0.02, -0.01, 0.00, 0.03, 0.05)
    stress_skew_shocks: Sequence[float] = (-0.12, -0.06, 0.0, 0.06, 0.12)
    stress_smile_shocks: Sequence[float] = (-0.06, -0.03, 0.0, 0.03, 0.06)
    ladder_spot_shocks: Sequence[float] = (-0.15, -0.10, -0.05, 0.0, 0.05)
    ladder_vol_shocks: Sequence[float] = (-0.05, 0.0, 0.05)
    terminal_cliff_band: float = 0.02
    greek_spot_multipliers: Sequence[float] = (
        0.60,
        0.65,
        0.70,
        0.75,
        0.80,
        0.85,
        0.90,
        0.95,
        1.00,
        1.03,
        1.05,
        1.10,
        1.15,
        1.20,
    )
    greek_key_spots: Sequence[float] = (65.0, 70.0, 75.0, 85.0, 100.0, 103.0)
    greek_q_shifts: Sequence[float] = (-0.02, -0.015, -0.01, -0.005, 0.0, 0.01, 0.02, 0.03, 0.05)
    greek_key_q_shifts: Sequence[float] = (-0.02, 0.0, 0.03, 0.05)
    greek_q_slice_spots: Mapping[str, float] = field(
        default_factory=lambda: {"near_ki": 70.0, "near_ko": 103.0}
    )
    generate_plots: bool = True


@dataclass(frozen=True)
class SnowballRiskComparisonArtifacts:
    report_path: Path
    output_dir: Path
    plot_paths: Mapping[str, Path]


@dataclass(frozen=True)
class StructuralDeltas:
    delta_monitoring: float
    delta_protection: float
    delta_parachute_dki: float
    delta_parachute_eki: float


@dataclass(frozen=True)
class StructureMetrics:
    label: str
    pv: float
    std_error: float
    expected_life: float
    no_ko_probability: float
    ki_probability: float
    ki_no_ko_probability: float
    loss_prob_5: float
    loss_prob_10: float
    loss_prob_20: float
    es95: float
    es99: float
    rebound_band_probability: float
    rebound_band_loss_probability: float
    parachute_rescue_band_probability: float
    parachute_rescue_loss_probability: float
    terminal_cliff_band_probability: float
    terminal_cliff_loss_probability: float
    conditional_ko_cashflows: pd.DataFrame


@dataclass(frozen=True)
class DeterministicCaseOutcome:
    state: str
    knocked_out: bool
    knocked_in: bool
    payoff: float
    payoff_ratio: float
    terminal_spot: float
    notes: str


@dataclass(frozen=True)
class DeterministicCaseResult:
    name: str
    narrative: str
    anchors: Sequence[tuple[float, float]]
    outcomes: Mapping[str, DeterministicCaseOutcome]


@dataclass(frozen=True)
class GreekCurves:
    spot_grid: np.ndarray
    delta: pd.DataFrame
    gamma: pd.DataFrame
    vega: pd.DataFrame
    rhoq: pd.DataFrame
    key_spot_table: pd.DataFrame
    q_grid: np.ndarray
    delta_q_curve: pd.DataFrame
    gamma_q_curve: pd.DataFrame
    vega_q_curve: pd.DataFrame
    rhoq_q_curve: pd.DataFrame
    key_q_table: pd.DataFrame
    q_slice_curves: Mapping[str, Mapping[str, pd.DataFrame]]


@dataclass
class _PathSnapshot:
    metrics: StructureMetrics
    discounted_payoff: np.ndarray
    payoff: np.ndarray
    settlement_times: np.ndarray
    terminal_spots: np.ndarray
    min_spots: np.ndarray
    is_ko: np.ndarray
    is_v0: np.ndarray
    is_v1: np.ndarray
    first_ko_idx: np.ndarray
    first_ki_idx: np.ndarray
    ko_times: np.ndarray


def build_default_snowball_risk_comparison_config(
    output_dir: Optional[Path] = None,
) -> SnowballRiskComparisonConfig:
    config = SnowballRiskComparisonConfig()
    if output_dir is None:
        return config
    return SnowballRiskComparisonConfig(output_dir=Path(output_dir))


def _add_months(base_date: datetime, months: int) -> datetime:
    year = base_date.year + (base_date.month - 1 + months) // 12
    month = (base_date.month - 1 + months) % 12 + 1
    day = min(
        base_date.day,
        [31, 29 if year % 4 == 0 and (year % 100 != 0 or year % 400 == 0) else 28, 31,
         30, 31, 30, 31, 31, 30, 31, 30, 31][month - 1],
    )
    return datetime(year, month, day)


def _business_day_time(
    calendar,
    start_date: datetime,
    end_date: datetime,
    bus_days_in_year: int,
) -> float:
    days = calendar.count_business_days(
        start_date, end_date, include_start=True, include_end=True
    )
    return days / float(bus_days_in_year)


def _build_ko_schedule(
    config: SnowballRiskComparisonConfig,
):
    calendar = create_calendar(
        config.calendar_type,
        year_range=(config.valuation_date.year, config.valuation_date.year + 2),
    )
    ko_dates = []
    for month in range(config.ko_start_month, config.tenor_months + 1):
        raw = _add_months(config.valuation_date, month)
        adj = calendar.adjust_date(raw, BusinessDayConvention.FOLLOWING)
        ko_dates.append(adj)
    maturity_date = ko_dates[-1]
    ko_times = [
        _business_day_time(
            calendar, config.valuation_date, dt, config.business_days_in_year
        )
        for dt in ko_dates
    ]
    daily_ki_dates = []
    current = config.valuation_date + timedelta(days=1)
    while current <= maturity_date:
        if calendar.is_business_day(current):
            daily_ki_dates.append(current)
        current += timedelta(days=1)
    daily_ki_times = [
        _business_day_time(
            calendar, config.valuation_date, dt, config.business_days_in_year
        )
        for dt in daily_ki_dates
    ]
    maturity_time = _business_day_time(
        calendar, config.valuation_date, maturity_date, config.business_days_in_year
    )
    return calendar, ko_dates, ko_times, daily_ki_dates, daily_ki_times, maturity_date, maturity_time


def _build_pricing_environment(
    config: SnowballRiskComparisonConfig,
    calendar,
) -> PricingEnvironment:
    return PricingEnvironment(
        spot_quote=SpotQuote(spot=config.initial_price),
        vol_surface=FlatVolSurface(volatility=config.volatility),
        rate_curve=FlatRateCurve(rate=config.rate),
        div_yield=ContinuousDividendYield(div_yield=config.dividend_yield),
        valuation_date=config.valuation_date,
        bus_days_in_year=config.business_days_in_year,
        calendar=calendar,
    )


def _base_payoff_config(
    config: SnowballRiskComparisonConfig,
    *,
    protection_type: ProtectionType,
) -> PayoffConfig:
    return PayoffConfig(
        rebate_rate=config.annual_coupon,
        include_principal=False,
        participation_rate=1.0,
        protection_type=protection_type,
        protection_rate=config.protection_rate if protection_type == ProtectionType.PARTIAL else 0.0,
    )


def _build_structures(
    config: SnowballRiskComparisonConfig,
    ko_times: Sequence[float],
    daily_ki_times: Sequence[float],
    maturity_time: float,
) -> Dict[str, SnowballOption]:
    ko_times_list = [float(t) for t in ko_times]
    daily_ki_list = [float(t) for t in daily_ki_times]

    ppp_dki = SnowballOption(
        initial_price=config.initial_price,
        strike=config.strike,
        maturity=maturity_time,
        contract_multiplier=config.contract_multiplier,
        barrier_config=BarrierConfig(
            ko_barrier=config.ko_barrier,
            ko_rate=config.annual_coupon,
            ko_observation_type=ObservationType.DISCRETE,
            ko_observation_dates=ko_times_list,
            ki_barrier=config.ki_barrier,
            ki_observation_type=ObservationType.DISCRETE,
            ki_observation_dates=daily_ki_list,
            ki_continuous=False,
        ),
        payoff_config=_base_payoff_config(
            config, protection_type=ProtectionType.PARTIAL
        ),
        accrual_config=AccrualConfig(
            coupon_pay_type=CouponPayType.INSTANT,
            is_annualized=True,
        ),
    )

    ppp_eki = create_european_ki_snowball(
        initial_price=config.initial_price,
        strike=config.strike,
        maturity=maturity_time,
        contract_multiplier=config.contract_multiplier,
        ko_barrier=config.ko_barrier,
        ko_rate=config.annual_coupon,
        ki_barrier=config.ki_barrier,
        num_ko_observations=len(ko_times_list),
        ko_observation_dates=ko_times_list,
        include_principal=False,
        participation_rate=1.0,
        protection_type=ProtectionType.PARTIAL,
        protection_rate=config.protection_rate,
        rebate_rate=config.annual_coupon,
    )

    npp_dki = SnowballOption(
        initial_price=config.initial_price,
        strike=config.strike,
        maturity=maturity_time,
        contract_multiplier=config.contract_multiplier,
        barrier_config=BarrierConfig(
            ko_barrier=config.ko_barrier,
            ko_rate=config.annual_coupon,
            ko_observation_type=ObservationType.DISCRETE,
            ko_observation_dates=ko_times_list,
            ki_barrier=config.ki_barrier,
            ki_observation_type=ObservationType.DISCRETE,
            ki_observation_dates=daily_ki_list,
            ki_continuous=False,
        ),
        payoff_config=_base_payoff_config(config, protection_type=ProtectionType.NONE),
        accrual_config=AccrualConfig(
            coupon_pay_type=CouponPayType.INSTANT,
            is_annualized=True,
        ),
    )

    ppp_dki_parachute = create_parachute_snowball(
        initial_price=config.initial_price,
        strike=config.strike,
        maturity=maturity_time,
        contract_multiplier=config.contract_multiplier,
        ko_barrier=config.ko_barrier,
        ko_rate=config.annual_coupon,
        ki_barrier=config.ki_barrier,
        num_observations=len(ko_times_list),
        ko_observation_dates=ko_times_list,
        ki_observation_type=ObservationType.DISCRETE,
        ki_observation_dates=daily_ki_list,
        ki_continuous=False,
        include_principal=False,
        participation_rate=1.0,
        protection_type=ProtectionType.PARTIAL,
        protection_rate=config.protection_rate,
        rebate_rate=config.annual_coupon,
    )

    ppp_eki_parachute = create_parachute_snowball(
        initial_price=config.initial_price,
        strike=config.strike,
        maturity=maturity_time,
        contract_multiplier=config.contract_multiplier,
        ko_barrier=config.ko_barrier,
        ko_rate=config.annual_coupon,
        ki_barrier=config.ki_barrier,
        num_observations=len(ko_times_list),
        ko_observation_dates=ko_times_list,
        ki_observation_type=ObservationType.DISCRETE,
        ki_observation_dates=[maturity_time],
        ki_continuous=False,
        include_principal=False,
        participation_rate=1.0,
        protection_type=ProtectionType.PARTIAL,
        protection_rate=config.protection_rate,
        rebate_rate=config.annual_coupon,
    )

    return {
        "PPP-DKI": ppp_dki,
        "PPP-EKI": ppp_eki,
        "NPP-DKI": npp_dki,
        "PPP-DKI-Parachute": ppp_dki_parachute,
        "PPP-EKI-Parachute": ppp_eki_parachute,
    }


def _get_principal(product: SnowballOption) -> float:
    return product.initial_price * product.contract_multiplier


def _compute_ko_cashflow_table(
    discounted_payoff: np.ndarray,
    is_ko: np.ndarray,
    first_ko_idx: np.ndarray,
    ko_times: np.ndarray,
) -> pd.DataFrame:
    rows = []
    cumulative = 0.0
    for idx, ko_time in enumerate(ko_times):
        hit = is_ko & (first_ko_idx == idx)
        p_ko = float(np.mean(hit))
        cumulative += p_ko
        rows.append(
            {
                "ko_time": float(ko_time),
                "p_ko": p_ko,
                "p_survive": max(0.0, 1.0 - cumulative),
                "ed_ko_cf": float(np.mean(np.where(hit, discounted_payoff, 0.0))),
            }
        )
    return pd.DataFrame(rows)


def _loss_metrics(payoff_ratio: np.ndarray) -> tuple[float, float, float, float, float]:
    losses = -np.minimum(payoff_ratio, 0.0)
    loss_prob_5 = float(np.mean(losses > 0.05))
    loss_prob_10 = float(np.mean(losses > 0.10))
    loss_prob_20 = float(np.mean(losses > 0.20))
    q95 = float(np.quantile(losses, 0.95))
    q99 = float(np.quantile(losses, 0.99))
    es95 = float(np.mean(losses[losses >= q95])) if np.any(losses >= q95) else 0.0
    es99 = float(np.mean(losses[losses >= q99])) if np.any(losses >= q99) else 0.0
    return loss_prob_5, loss_prob_10, loss_prob_20, es95, es99


def _classify_paths(
    *,
    config: SnowballRiskComparisonConfig,
    product: SnowballOption,
    pricing_env: PricingEnvironment,
    engine: SnowballMCEngine,
    paths: np.ndarray,
    all_times: np.ndarray,
    rng_seed: int,
) -> _PathSnapshot:
    ko_profile = product.get_ko_observation_profile(pricing_env)
    ko_times = np.array(ko_profile["observation_times"], dtype=float)
    ko_barriers = np.array(ko_profile["barriers"], dtype=float)
    ko_indices = np.searchsorted(all_times, ko_times)

    ki_profile = product.get_ki_observation_profile(pricing_env)
    ki_times = np.array(ki_profile["observation_times"], dtype=float)
    ki_barriers = np.array(ki_profile["barriers"], dtype=float)
    ki_indices = np.searchsorted(all_times, ki_times)

    ko_triggered, first_ko_idx = engine._check_ko_barriers(
        paths, ko_indices, ko_barriers, product.is_reverse
    )
    ki_triggered, first_ki_idx = engine._check_ki_barriers(
        paths, ki_indices, ki_barriers, product.is_reverse
    )

    is_ko = ko_triggered
    is_v0 = ~is_ko & ~ki_triggered
    is_v1 = ~is_ko & ki_triggered

    maturity = product.get_maturity(pricing_env)
    sigma = pricing_env.get_vol(product.strike, maturity)
    payoffs, settlement_times, stats = engine._compute_payoffs(
        product=product,
        pricing_env=pricing_env,
        paths=paths,
        all_times=all_times,
        ko_indices=ko_indices,
        ki_indices=ki_indices,
        r=pricing_env.get_rate(maturity),
        T=maturity,
        sigma=sigma,
        rng_seed=rng_seed,
    )
    discounted_payoff = np.array(
        [
            payoff * pricing_env.get_discount_factor(float(t))
            for payoff, t in zip(payoffs, settlement_times)
        ],
        dtype=float,
    )

    principal = _get_principal(product)
    payoff_ratio = payoffs / principal
    loss_prob_5, loss_prob_10, loss_prob_20, es95, es99 = _loss_metrics(payoff_ratio)
    min_spots = np.min(paths, axis=1)
    terminal_spots = paths[:, -1]
    ki_level = float(product.barrier_config.ki_barrier)
    standard_final_ko = float(config.ko_barrier)
    rebound_band = (~is_ko) & (min_spots < ki_level) & (terminal_spots > ki_level) & (
        terminal_spots < product.strike
    )
    rescue_band = (~is_ko) & (min_spots < ki_level) & (terminal_spots >= ki_level) & (
        terminal_spots < standard_final_ko
    )
    cliff_band = (~is_ko) & (
        np.abs(terminal_spots / ki_level - 1.0) <= float(config.terminal_cliff_band)
    )
    losses = payoff_ratio < 0.0

    metrics = StructureMetrics(
        label="",
        pv=float(np.mean(discounted_payoff)),
        std_error=float(np.std(discounted_payoff, ddof=1) / np.sqrt(len(discounted_payoff))),
        expected_life=float(np.mean(settlement_times)),
        no_ko_probability=float(np.mean(~is_ko)),
        ki_probability=float(np.mean(ki_triggered)),
        ki_no_ko_probability=float(np.mean(is_v1)),
        loss_prob_5=loss_prob_5,
        loss_prob_10=loss_prob_10,
        loss_prob_20=loss_prob_20,
        es95=es95,
        es99=es99,
        rebound_band_probability=float(np.mean(rebound_band)),
        rebound_band_loss_probability=float(np.mean(rebound_band & losses)),
        parachute_rescue_band_probability=float(np.mean(rescue_band)),
        parachute_rescue_loss_probability=float(np.mean(rescue_band & losses)),
        terminal_cliff_band_probability=float(np.mean(cliff_band)),
        terminal_cliff_loss_probability=float(np.mean(cliff_band & losses)),
        conditional_ko_cashflows=_compute_ko_cashflow_table(
            discounted_payoff=discounted_payoff,
            is_ko=is_ko,
            first_ko_idx=first_ko_idx,
            ko_times=ko_times,
        ),
    )
    return _PathSnapshot(
        metrics=metrics,
        discounted_payoff=discounted_payoff,
        payoff=payoffs,
        settlement_times=settlement_times,
        terminal_spots=terminal_spots,
        min_spots=min_spots,
        is_ko=is_ko,
        is_v0=is_v0,
        is_v1=is_v1,
        first_ko_idx=first_ko_idx,
        first_ki_idx=first_ki_idx,
        ko_times=ko_times,
    )


def _simulate_common_paths(
    *,
    config: SnowballRiskComparisonConfig,
    pricing_env: PricingEnvironment,
    products: Mapping[str, SnowballOption],
) -> tuple[np.ndarray, np.ndarray, Dict[str, _PathSnapshot]]:
    all_times_set = {float(products["PPP-DKI"].get_maturity(pricing_env))}
    for product in products.values():
        ko_profile = product.get_ko_observation_profile(pricing_env)
        all_times_set.update(float(t) for t in ko_profile["observation_times"])
        ki_profile = product.get_ki_observation_profile(pricing_env)
        all_times_set.update(float(t) for t in ki_profile["observation_times"])
    all_times = np.array(sorted(all_times_set), dtype=float)
    dt_array = np.diff(np.concatenate([[0.0], all_times]))
    mc_engine = SnowballMCEngine(
        params=MCParams(
            num_paths=config.num_paths,
            seed=config.seed,
            bus_days_in_year=config.business_days_in_year,
        ),
        method=EngineType.MONTE_CARLO(config.mc_method),
    )
    generator = mc_engine._create_path_generator(
        S=pricing_env.spot,
        r=pricing_env.get_rate(float(all_times[-1])),
        q=pricing_env.get_div_yield(float(all_times[-1])),
        sigma=pricing_env.get_vol(config.strike, float(all_times[-1])),
        T=float(all_times[-1]),
        dt_array=dt_array,
    )
    paths, _ = generator.generate_paths(return_aux=False)

    snapshots: Dict[str, _PathSnapshot] = {}
    for idx, (label, product) in enumerate(products.items()):
        snapshot = _classify_paths(
            config=config,
            product=product,
            pricing_env=pricing_env,
            engine=mc_engine,
            paths=paths,
            all_times=all_times,
            rng_seed=config.seed + 1000 + idx,
        )
        snapshot.metrics = StructureMetrics(
            label=label,
            pv=snapshot.metrics.pv,
            std_error=snapshot.metrics.std_error,
            expected_life=snapshot.metrics.expected_life,
            no_ko_probability=snapshot.metrics.no_ko_probability,
            ki_probability=snapshot.metrics.ki_probability,
            ki_no_ko_probability=snapshot.metrics.ki_no_ko_probability,
            loss_prob_5=snapshot.metrics.loss_prob_5,
            loss_prob_10=snapshot.metrics.loss_prob_10,
            loss_prob_20=snapshot.metrics.loss_prob_20,
            es95=snapshot.metrics.es95,
            es99=snapshot.metrics.es99,
            rebound_band_probability=snapshot.metrics.rebound_band_probability,
            rebound_band_loss_probability=snapshot.metrics.rebound_band_loss_probability,
            parachute_rescue_band_probability=snapshot.metrics.parachute_rescue_band_probability,
            parachute_rescue_loss_probability=snapshot.metrics.parachute_rescue_loss_probability,
            terminal_cliff_band_probability=snapshot.metrics.terminal_cliff_band_probability,
            terminal_cliff_loss_probability=snapshot.metrics.terminal_cliff_loss_probability,
            conditional_ko_cashflows=snapshot.metrics.conditional_ko_cashflows,
        )
        snapshots[label] = snapshot
    return all_times, paths, snapshots


def _interpolate_path(
    anchors: Sequence[tuple[float, float]],
    observation_times: np.ndarray,
) -> np.ndarray:
    anchor_times = np.array([float(t) for t, _ in anchors], dtype=float)
    anchor_spots = np.array([float(s) for _, s in anchors], dtype=float)
    return np.interp(observation_times, anchor_times, anchor_spots)


def _evaluate_deterministic_path(
    *,
    product: SnowballOption,
    pricing_env: PricingEnvironment,
    observation_times: np.ndarray,
    path_spots: np.ndarray,
) -> DeterministicCaseOutcome:
    ko_profile = product.get_ko_observation_profile(pricing_env)
    ko_times = np.array(ko_profile["observation_times"], dtype=float)
    ko_barriers = np.array(ko_profile["barriers"], dtype=float)
    ko_payoffs = np.array(ko_profile["payoffs"], dtype=float)
    ko_indices = np.searchsorted(observation_times, ko_times)
    ko_prices = path_spots[ko_indices]
    ko_hit = ko_prices >= ko_barriers
    if ko_hit.any():
        idx = int(np.argmax(ko_hit))
        payoff = float(ko_payoffs[idx])
        return DeterministicCaseOutcome(
            state="KO",
            knocked_out=True,
            knocked_in=False,
            payoff=payoff,
            payoff_ratio=payoff / _get_principal(product),
            terminal_spot=float(path_spots[-1]),
            notes=f"KO at observation {idx + 1}",
        )

    ki_profile = product.get_ki_observation_profile(pricing_env)
    ki_times = np.array(ki_profile["observation_times"], dtype=float)
    ki_barriers = np.array(ki_profile["barriers"], dtype=float)
    ki_indices = np.searchsorted(observation_times, ki_times)
    ki_prices = path_spots[ki_indices]
    ki_hit = ki_prices <= ki_barriers
    knocked_in = bool(ki_hit.any())
    terminal_spot = float(path_spots[-1])
    if knocked_in:
        payoff = float(product.get_maturity_payoff_v1(terminal_spot, pricing_env))
        state = "KI / No KO"
    else:
        payoff = float(product.get_maturity_payoff_v0(terminal_spot, pricing_env))
        state = "No KI / No KO"
    return DeterministicCaseOutcome(
        state=state,
        knocked_out=False,
        knocked_in=knocked_in,
        payoff=payoff,
        payoff_ratio=payoff / _get_principal(product),
        terminal_spot=terminal_spot,
        notes="Terminal payoff",
    )


def _build_deterministic_cases(maturity_time: float) -> Sequence[tuple[str, str, Sequence[tuple[float, float]]]]:
    return [
        (
            "Crash Then Full Rebound",
            "Sharp selloff followed by a full recovery above the KO region.",
            [(0.0, 100.0), (0.35 * maturity_time, 65.0), (maturity_time, 110.0)],
        ),
        (
            "Crash Then Partial Rebound 85",
            "Rebound path used to highlight the monitoring-rule advantage of EKI.",
            [(0.0, 100.0), (0.35 * maturity_time, 65.0), (maturity_time, 85.0)],
        ),
        (
            "Parachute Rescue 74",
            "Rebound into the KI-to-standard-final-KO band where parachute should matter most.",
            [(0.0, 100.0), (0.35 * maturity_time, 65.0), (maturity_time, 74.0)],
        ),
        (
            "Slow Bleed Lower",
            "Persistent drawdown finishing below KI to show convergence of protected structures.",
            [
                (0.0, 100.0),
                (0.25 * maturity_time, 92.0),
                (0.50 * maturity_time, 84.0),
                (0.75 * maturity_time, 76.0),
                (maturity_time, 68.0),
            ],
        ),
        (
            "High-Vol Sideways Around KI",
            "Sideways market oscillating around the KI level without recovering materially.",
            [
                (0.0, 100.0),
                (0.25 * maturity_time, 74.0),
                (0.50 * maturity_time, 69.0),
                (0.75 * maturity_time, 73.0),
                (maturity_time, 71.0),
            ],
        ),
        (
            "Late Selloff Final Window",
            "Late move near maturity to expose EKI terminal-state concentration.",
            [
                (0.0, 100.0),
                (0.70 * maturity_time, 101.0),
                (0.92 * maturity_time, 95.0),
                (maturity_time, 69.0),
            ],
        ),
    ]


def _run_deterministic_cases(
    *,
    products: Mapping[str, SnowballOption],
    pricing_env: PricingEnvironment,
    observation_times: np.ndarray,
) -> Sequence[DeterministicCaseResult]:
    cases = []
    maturity_time = float(observation_times[-1])
    for name, narrative, anchors in _build_deterministic_cases(maturity_time):
        path_spots = _interpolate_path(anchors, observation_times)
        outcomes = {
            label: _evaluate_deterministic_path(
                product=product,
                pricing_env=pricing_env,
                observation_times=observation_times,
                path_spots=path_spots,
            )
            for label, product in products.items()
        }
        cases.append(
            DeterministicCaseResult(
                name=name,
                narrative=narrative,
                anchors=anchors,
                outcomes=outcomes,
            )
        )
    return cases


def _select_stress_engine(
    config: SnowballRiskComparisonConfig,
) -> tuple[BaseEngine, str]:
    return _select_snowball_pricing_engine(
        preference=config.engine_preference,
        quad_params=config.quad_params,
        pde_params=config.pde_params,
        mc_params=MCParams(
            num_paths=max(4000, min(config.num_paths, 12000)),
            seed=config.seed,
            bus_days_in_year=config.business_days_in_year,
        ),
    )


def _compute_stress_curve(
    *,
    base_product: SnowballOption,
    pricing_env: PricingEnvironment,
    engine: BaseEngine,
    shocks: Sequence[float],
    kind: str,
) -> pd.DataFrame:
    rows = []
    base_pv = float(engine.price(base_product, pricing_env))
    for shock in shocks:
        if kind == "spot":
            env = _clone_env(pricing_env, spot=pricing_env.spot * (1.0 + float(shock)))
        elif kind == "vol":
            vol_surface = _scale_vol_surface(pricing_env.vol_surface, 1.0 + float(shock))
            env = _clone_env(pricing_env, vol_surface=vol_surface)
        elif kind == "div":
            div = _shift_dividend_yield(pricing_env.div_yield, float(shock))
            env = _clone_env(pricing_env, div_yield=div)
        else:
            raise ValidationError(f"Unknown stress curve kind: {kind}")
        pv = float(engine.price(base_product, env))
        rows.append({"shock": float(shock), "pv": pv, "pnl": pv - base_pv})
    return pd.DataFrame(rows)


def _compute_skew_smile_heatmap(
    *,
    products: Mapping[str, SnowballOption],
    pricing_env: PricingEnvironment,
    engine: BaseEngine,
    skew_shocks: Sequence[float],
    smile_shocks: Sequence[float],
) -> np.ndarray:
    ppp_eki = products["PPP-EKI"]
    ppp_dki = products["PPP-DKI"]
    z = np.zeros((len(skew_shocks), len(smile_shocks)), dtype=float)
    if pricing_env.vol_surface is None:
        return z
    for i, skew in enumerate(skew_shocks):
        for j, smile in enumerate(smile_shocks):
            skew_surface = SkewSmileVolSurface(
                base=pricing_env.vol_surface,
                skew=float(skew),
                smile=float(smile),
            )
            env = _clone_env(pricing_env, vol_surface=skew_surface)
            z[i, j] = float(engine.price(ppp_eki, env) - engine.price(ppp_dki, env))
    return z


def _compute_delta_heatmap(
    *,
    left_product: SnowballOption,
    right_product: SnowballOption,
    pricing_env: PricingEnvironment,
    engine: BaseEngine,
    spot_shocks: Sequence[float],
    vol_shocks: Sequence[float],
) -> np.ndarray:
    z = np.zeros((len(spot_shocks), len(vol_shocks)), dtype=float)
    if pricing_env.vol_surface is None:
        return z
    for i, s_shock in enumerate(spot_shocks):
        spot = pricing_env.spot * (1.0 + float(s_shock))
        for j, v_shock in enumerate(vol_shocks):
            vol_surface = _scale_vol_surface(
                pricing_env.vol_surface, 1.0 + float(v_shock)
            )
            env = _clone_env(pricing_env, spot=spot, vol_surface=vol_surface)
            z[i, j] = float(engine.price(left_product, env) - engine.price(right_product, env))
    return z


def _require_matplotlib():
    os.environ.setdefault("MPLCONFIGDIR", "/tmp/matplotlib")
    import matplotlib.pyplot as plt  # type: ignore

    return plt


def _save_multi_line_plot(
    *,
    curves: Mapping[str, pd.DataFrame],
    x_col: str,
    y_col: str,
    title: str,
    xlabel: str,
    ylabel: str,
    path: Path,
) -> None:
    plt = _require_matplotlib()
    path.parent.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(8, 4.5))
    for label, df in curves.items():
        ax.plot(df[x_col].to_numpy(), df[y_col].to_numpy(), linewidth=2, label=label)
    ax.set_title(title)
    ax.set_xlabel(xlabel)
    ax.set_ylabel(ylabel)
    ax.grid(True, alpha=0.3)
    ax.legend()
    fig.tight_layout()
    fig.savefig(path, dpi=150)
    plt.close(fig)


def _save_terminal_cliff_plot(
    *,
    snapshots: Mapping[str, _PathSnapshot],
    ki_barrier: float,
    path: Path,
) -> None:
    plt = _require_matplotlib()
    path.parent.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(8, 4.5))
    bins = np.linspace(0.90 * ki_barrier, 1.10 * ki_barrier, 17)
    for label in ("PPP-EKI", "PPP-DKI"):
        snapshot = snapshots[label]
        mask = ~snapshot.is_ko
        if not np.any(mask):
            continue
        x = snapshot.terminal_spots[mask]
        y = snapshot.payoff[mask]
        centers = 0.5 * (bins[1:] + bins[:-1])
        values = []
        for left, right in zip(bins[:-1], bins[1:]):
            bin_mask = (x >= left) & (x < right)
            values.append(float(np.mean(y[bin_mask])) if np.any(bin_mask) else np.nan)
        ax.plot(centers, values, marker="o", linewidth=2, label=label)
    ax.axvline(ki_barrier, linestyle="--", color="black", alpha=0.6, label="KI")
    ax.set_title("Terminal Cliff Analysis")
    ax.set_xlabel("Terminal Spot")
    ax.set_ylabel("Average Payoff")
    ax.grid(True, alpha=0.3)
    ax.legend()
    fig.tight_layout()
    fig.savefig(path, dpi=150)
    plt.close(fig)


def _generate_plot_artifacts(
    *,
    config: SnowballRiskComparisonConfig,
    products: Mapping[str, SnowballOption],
    pricing_env: PricingEnvironment,
    snapshots: Mapping[str, _PathSnapshot],
    greek_curves: GreekCurves,
    output_dir: Path,
) -> Dict[str, Path]:
    if not config.generate_plots:
        return {}

    plot_dir = output_dir / "plots"
    plot_dir.mkdir(parents=True, exist_ok=True)
    engine, _ = _select_stress_engine(config)

    spot_curves = {
        label: _compute_stress_curve(
            base_product=product,
            pricing_env=pricing_env,
            engine=engine,
            shocks=config.stress_spot_shocks,
            kind="spot",
        )
        for label, product in products.items()
        if label in {"PPP-EKI", "PPP-DKI", "NPP-DKI"}
    }
    vol_curves = {
        label: _compute_stress_curve(
            base_product=product,
            pricing_env=pricing_env,
            engine=engine,
            shocks=config.stress_vol_shocks,
            kind="vol",
        )
        for label, product in products.items()
        if label in {"PPP-EKI", "PPP-DKI", "NPP-DKI"}
    }
    div_curves = {
        label: _compute_stress_curve(
            base_product=product,
            pricing_env=pricing_env,
            engine=engine,
            shocks=config.stress_div_shifts,
            kind="div",
        )
        for label, product in products.items()
        if label in {"PPP-EKI", "PPP-DKI", "NPP-DKI"}
    }

    plot_paths = {
        "spot_stress": plot_dir / "spot_stress.png",
        "vol_stress": plot_dir / "vol_stress.png",
        "div_stress": plot_dir / "div_stress.png",
        "skew_smile": plot_dir / "skew_smile_delta_monitoring.png",
        "delta_monitoring": plot_dir / "delta_monitoring_heatmap.png",
        "delta_protection": plot_dir / "delta_protection_heatmap.png",
        "terminal_cliff": plot_dir / "terminal_cliff.png",
        "greek_delta": plot_dir / "greek_delta_vs_spot.png",
        "greek_gamma": plot_dir / "greek_gamma_vs_spot.png",
        "greek_vega": plot_dir / "greek_vega_vs_spot.png",
        "greek_rhoq": plot_dir / "greek_rhoq_vs_spot.png",
    }
    _save_multi_line_plot(
        curves=spot_curves,
        x_col="shock",
        y_col="pnl",
        title="Spot Shock PnL",
        xlabel="Spot Shock",
        ylabel="PnL",
        path=plot_paths["spot_stress"],
    )
    _save_multi_line_plot(
        curves=vol_curves,
        x_col="shock",
        y_col="pnl",
        title="Volatility Shock PnL",
        xlabel="Vol Shock",
        ylabel="PnL",
        path=plot_paths["vol_stress"],
    )
    _save_multi_line_plot(
        curves=div_curves,
        x_col="shock",
        y_col="pnl",
        title="Dividend Shock PnL",
        xlabel="Dividend Yield Shift",
        ylabel="PnL",
        path=plot_paths["div_stress"],
    )

    skew_smile = _compute_skew_smile_heatmap(
        products=products,
        pricing_env=pricing_env,
        engine=engine,
        skew_shocks=config.stress_skew_shocks,
        smile_shocks=config.stress_smile_shocks,
    )
    save_heatmap(
        x=np.array(config.stress_skew_shocks, dtype=float),
        y=np.array(config.stress_smile_shocks, dtype=float),
        z=skew_smile,
        title="Delta Monitoring under Skew / Smile Shocks",
        xlabel="Skew Shock",
        ylabel="Smile Shock",
        colorbar_label="PPP-EKI minus PPP-DKI",
        path=plot_paths["skew_smile"],
    )

    monitoring_heatmap = _compute_delta_heatmap(
        left_product=products["PPP-EKI"],
        right_product=products["PPP-DKI"],
        pricing_env=pricing_env,
        engine=engine,
        spot_shocks=config.stress_spot_shocks,
        vol_shocks=config.stress_vol_shocks,
    )
    save_heatmap(
        x=np.array(config.stress_spot_shocks, dtype=float),
        y=np.array(config.stress_vol_shocks, dtype=float),
        z=monitoring_heatmap,
        title="Monitoring Delta Heatmap",
        xlabel="Spot Shock",
        ylabel="Vol Shock",
        colorbar_label="PPP-EKI minus PPP-DKI",
        path=plot_paths["delta_monitoring"],
    )

    protection_heatmap = _compute_delta_heatmap(
        left_product=products["PPP-DKI"],
        right_product=products["NPP-DKI"],
        pricing_env=pricing_env,
        engine=engine,
        spot_shocks=config.stress_spot_shocks,
        vol_shocks=config.stress_vol_shocks,
    )
    save_heatmap(
        x=np.array(config.stress_spot_shocks, dtype=float),
        y=np.array(config.stress_vol_shocks, dtype=float),
        z=protection_heatmap,
        title="Protection Delta Heatmap",
        xlabel="Spot Shock",
        ylabel="Vol Shock",
        colorbar_label="PPP-DKI minus NPP-DKI",
        path=plot_paths["delta_protection"],
    )
    _save_terminal_cliff_plot(
        snapshots=snapshots,
        ki_barrier=config.ki_barrier,
        path=plot_paths["terminal_cliff"],
    )
    _save_greek_curve_plot(
        greek_df=greek_curves.delta,
        x_col="spot",
        title="Delta vs Spot",
        xlabel="Spot",
        ylabel="Delta",
        path=plot_paths["greek_delta"],
    )
    _save_greek_curve_plot(
        greek_df=greek_curves.gamma,
        x_col="spot",
        title="Gamma vs Spot",
        xlabel="Spot",
        ylabel="Gamma",
        path=plot_paths["greek_gamma"],
    )
    _save_greek_curve_plot(
        greek_df=greek_curves.vega,
        x_col="spot",
        title="Vega vs Spot",
        xlabel="Spot",
        ylabel="Vega",
        path=plot_paths["greek_vega"],
    )
    _save_greek_curve_plot(
        greek_df=greek_curves.rhoq,
        x_col="spot",
        title="RhoQ vs Spot",
        xlabel="Spot",
        ylabel="Dividend Rho (RhoQ)",
        path=plot_paths["greek_rhoq"],
    )
    plot_paths.update(
        {
            "greek_delta_q": plot_dir / "greek_delta_vs_q.png",
            "greek_gamma_q": plot_dir / "greek_gamma_vs_q.png",
            "greek_vega_q": plot_dir / "greek_vega_vs_q.png",
            "greek_rhoq_q": plot_dir / "greek_rhoq_vs_q.png",
        }
    )
    _save_greek_curve_plot(
        greek_df=greek_curves.delta_q_curve,
        x_col="q",
        title="Delta vs Dividend Yield",
        xlabel="Dividend Yield q",
        ylabel="Delta",
        path=plot_paths["greek_delta_q"],
    )
    _save_greek_curve_plot(
        greek_df=greek_curves.gamma_q_curve,
        x_col="q",
        title="Gamma vs Dividend Yield",
        xlabel="Dividend Yield q",
        ylabel="Gamma",
        path=plot_paths["greek_gamma_q"],
    )
    _save_greek_curve_plot(
        greek_df=greek_curves.vega_q_curve,
        x_col="q",
        title="Vega vs Dividend Yield",
        xlabel="Dividend Yield q",
        ylabel="Vega",
        path=plot_paths["greek_vega_q"],
    )
    _save_greek_curve_plot(
        greek_df=greek_curves.rhoq_q_curve,
        x_col="q",
        title="RhoQ vs Dividend Yield",
        xlabel="Dividend Yield q",
        ylabel="Dividend Rho (RhoQ)",
        path=plot_paths["greek_rhoq_q"],
    )
    plot_paths.update(
        {
            "greek_q_near_ki": plot_dir / "greek_q_near_ki_panel.png",
            "greek_q_near_ko": plot_dir / "greek_q_near_ko_panel.png",
        }
    )
    _save_greek_q_slice_panel(
        slice_curves=greek_curves.q_slice_curves["near_ki"],
        title=f"Greeks vs q at Spot = KI ({config.ki_barrier:.4f})",
        path=plot_paths["greek_q_near_ki"],
    )
    _save_greek_q_slice_panel(
        slice_curves=greek_curves.q_slice_curves["near_ko"],
        title=f"Greeks vs q at Spot = KO ({config.ko_barrier:.4f})",
        path=plot_paths["greek_q_near_ko"],
    )
    return plot_paths


def _fmt_pct(value: float) -> str:
    return f"{value:.4%}"


def _fmt_num(value: float) -> str:
    return f"{value:,.4f}"


def _add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def _add_paragraph(document: Document, text: str, *, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold


def _add_dataframe_table(document: Document, df: pd.DataFrame, title: Optional[str] = None) -> None:
    if title:
        _add_paragraph(document, title, bold=True)
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for idx, col in enumerate(df.columns):
        hdr_cells[idx].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(df.columns):
            value = row[col]
            if isinstance(value, float):
                cells[idx].text = f"{value:.4f}"
            else:
                cells[idx].text = str(value)


def _add_picture_if_exists(document: Document, path: Optional[Path], width: float = 6.6) -> None:
    if path is None or not path.exists():
        return
    document.add_picture(str(path), width=Inches(width))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _build_metrics_table(
    snapshots: Mapping[str, _PathSnapshot],
) -> pd.DataFrame:
    rows = []
    order = ["PPP-EKI", "PPP-DKI", "NPP-DKI", "PPP-DKI-Parachute", "PPP-EKI-Parachute"]
    for label in order:
        metrics = snapshots[label].metrics
        rows.append(
            {
                "Structure": label,
                "PV": metrics.pv,
                "StdErr": metrics.std_error,
                "P(No KO)": metrics.no_ko_probability,
                "P(KI)": metrics.ki_probability,
                "P(KI & No KO)": metrics.ki_no_ko_probability,
                "P(loss>5%)": metrics.loss_prob_5,
                "P(loss>10%)": metrics.loss_prob_10,
                "P(loss>20%)": metrics.loss_prob_20,
                "ES95": metrics.es95,
                "ES99": metrics.es99,
            }
        )
    return pd.DataFrame(rows)


def _build_contract_terms_table(
    *,
    config: SnowballRiskComparisonConfig,
    products: Mapping[str, SnowballOption],
    language: str,
) -> pd.DataFrame:
    is_english = language == "en"
    rows = []
    order = ["PPP-EKI", "PPP-DKI", "NPP-DKI", "PPP-DKI-Parachute", "PPP-EKI-Parachute"]
    for label in order:
        product = products[label]
        ko_barrier = product.barrier_config.ko_barrier
        final_ko_barrier = (
            ko_barrier[-1] if isinstance(ko_barrier, list) else ko_barrier
        )
        protection_type = product.payoff_config.protection_type.name
        ki_dates = product.barrier_config.ki_observation_dates or []
        if len(ki_dates) == 1:
            monitoring = "European final-only KI" if is_english else "欧式终值敲入"
        else:
            monitoring = "Daily business-day KI" if is_english else "交易日日度敲入"
        rows.append(
            {
                ("Structure" if is_english else "结构"): label,
                ("Monitoring" if is_english else "敲入监控"): monitoring,
                ("Tenor" if is_english else "期限"): (
                    f"{config.tenor_months}M" if is_english else f"{config.tenor_months}个月"
                ),
                ("KO Schedule" if is_english else "敲出安排"): (
                    f"Monthly, months {config.ko_start_month}-{config.tenor_months}"
                    if is_english
                    else f"月度观察，第{config.ko_start_month}-{config.tenor_months}个月"
                ),
                ("KO Barrier" if is_english else "敲出价"): f"{float(config.ko_barrier):.4f}",
                ("Final KO" if is_english else "最终敲出价"): f"{float(final_ko_barrier):.4f}",
                ("KI Barrier" if is_english else "敲入价"): f"{float(config.ki_barrier):.4f}",
                ("Coupon / Rebate" if is_english else "票息 / 红利"): _fmt_pct(config.annual_coupon),
                ("PV Principal" if is_english else "PV本金处理"): (
                    "Excluded" if is_english else "不含本金"
                ),
                ("Protection" if is_english else "保本方式"): (
                    f"{protection_type} ({_fmt_pct(config.protection_rate)} floor)"
                    if is_english and protection_type == "PARTIAL"
                    else (
                        protection_type
                        if is_english
                        else ("部分保本" if protection_type == "PARTIAL" else "非保本")
                    )
                ),
                ("Post-KI KO" if is_english else "敲入后可敲出"): (
                    "Allowed" if is_english else "允许"
                ),
            }
        )
    return pd.DataFrame(rows)


def _build_delta_table(deltas: StructuralDeltas) -> pd.DataFrame:
    return pd.DataFrame(
        [
            {"Measure": "Δ_monitoring", "Value": deltas.delta_monitoring},
            {"Measure": "Δ_protection", "Value": deltas.delta_protection},
            {"Measure": "Δ_parachute_DKI", "Value": deltas.delta_parachute_dki},
            {"Measure": "Δ_parachute_EKI", "Value": deltas.delta_parachute_eki},
        ]
    )


def _build_band_table(snapshots: Mapping[str, _PathSnapshot]) -> pd.DataFrame:
    rows = []
    for label in ("PPP-EKI", "PPP-DKI", "NPP-DKI", "PPP-DKI-Parachute", "PPP-EKI-Parachute"):
        metrics = snapshots[label].metrics
        rows.append(
            {
                "Structure": label,
                "Rebound Band": metrics.rebound_band_probability,
                "Rebound Band Loss": metrics.rebound_band_loss_probability,
                "Parachute Rescue Band": metrics.parachute_rescue_band_probability,
                "Parachute Rescue Loss": metrics.parachute_rescue_loss_probability,
                "Terminal Cliff Band": metrics.terminal_cliff_band_probability,
                "Terminal Cliff Loss": metrics.terminal_cliff_loss_probability,
            }
        )
    return pd.DataFrame(rows)


def _build_deterministic_table(cases: Sequence[DeterministicCaseResult]) -> pd.DataFrame:
    rows = []
    for case in cases:
        for structure, outcome in case.outcomes.items():
            rows.append(
                {
                    "Case": case.name,
                    "Structure": structure,
                    "State": outcome.state,
                    "Payoff": outcome.payoff,
                    "Return": outcome.payoff_ratio,
                    "Terminal": outcome.terminal_spot,
                }
            )
    return pd.DataFrame(rows)


def _build_terminology_table(language: str) -> pd.DataFrame:
    is_english = language == "en"
    terms = [
        (
            "Snowball",
            "雪球",
            "Autocallable structure with KO observations, KI condition, coupons, and a maturity payoff that depends on the path state.",
            "带有敲出观察、敲入条件、票息以及按路径状态决定到期收益的自动赎回结构。",
        ),
        (
            "KI",
            "敲入",
            "Knock-in event. Once activated, the note moves into the downside payoff regime if it survives to maturity without KO.",
            "敲入事件。一旦触发，若产品未提前敲出并持有到期，则进入下跌损失收益状态。",
        ),
        (
            "EKI",
            "欧式敲入",
            "European knock-in. The KI barrier is checked only at the final fixing.",
            "欧式敲入。仅在最终观察日检查敲入障碍。",
        ),
        (
            "DKI",
            "日度敲入",
            "Daily knock-in. The KI barrier is checked on each business-day close.",
            "日度敲入。每个交易日收盘检查敲入障碍。",
        ),
        (
            "KO",
            "敲出",
            "Knock-out event. If triggered on an observation date, the note terminates early and pays the contractual KO payoff.",
            "敲出事件。在观察日触发后，产品提前终止并支付约定的敲出收益。",
        ),
        (
            "PPP",
            "部分保本",
            "Partial principal protection. Losses are floored by a contractual protection rate rather than fully exposed.",
            "部分保本。损失受到保本比例下限约束，而不是完全暴露。",
        ),
        (
            "NPP",
            "非保本",
            "Non-principal protected. Downside is not floored by a protection level.",
            "非保本。下行损失没有保本下限约束。",
        ),
        (
            "Parachute",
            "降落伞",
            "Feature that lowers the final KO barrier to the KI barrier to rescue some KI paths that recover by maturity.",
            "一种将最终敲出价降至敲入价的条款，用于挽救部分到期前反弹的敲入路径。",
        ),
        (
            "Rebound Band",
            "反弹带",
            "State bucket defined by interim breach below KI, no KO, and terminal spot between KI and strike.",
            "一种状态区间：路径中途跌破 KI、未敲出、且终值位于 KI 与行权价之间。",
        ),
        (
            "Parachute Rescue Band",
            "降落伞救援带",
            "State bucket where the path breached KI, avoided standard KO, and finished between KI and the standard final KO barrier.",
            "一种状态区间：路径曾跌破 KI、未达到标准敲出、且终值位于 KI 与标准最终敲出价之间。",
        ),
        (
            "Terminal Cliff",
            "终值悬崖",
            "Concentrated terminal-state risk when spot finishes close to the KI barrier near maturity.",
            "到期前终值贴近 KI 时形成的集中终值状态风险。",
        ),
        (
            "PV",
            "现值",
            "Present value of discounted expected cashflows under the report model.",
            "在本报告模型下，贴现后的期望现金流现值。",
        ),
        (
            "ES95 / ES99",
            "ES95 / ES99",
            "Expected shortfall at the 95% / 99% confidence tail, measuring average tail loss severity.",
            "95% / 99% 置信尾部的期望短缺，用于衡量平均尾部损失严重度。",
        ),
        (
            "Common Random Numbers",
            "共同随机数",
            "Matched Monte Carlo draws reused across structures to make structural deltas less noisy.",
            "在不同结构之间复用相同蒙特卡洛随机路径，以减少结构差值的噪声。",
        ),
    ]
    rows = []
    for en_term, zh_term, en_def, zh_def in terms:
        rows.append(
            {
                ("Term" if is_english else "术语"): en_term if is_english else zh_term,
                ("English" if is_english else "英文"): en_term,
                ("Chinese" if is_english else "中文"): zh_term,
                ("Definition" if is_english else "定义"): en_def if is_english else zh_def,
            }
        )
    return pd.DataFrame(rows)


def _compute_greek_curves(
    *,
    config: SnowballRiskComparisonConfig,
    pricing_env: PricingEnvironment,
    products: Mapping[str, SnowballOption],
) -> GreekCurves:
    engine, _ = _select_stress_engine(config)
    greek_params = EngineParams(bus_days_in_year=config.business_days_in_year)
    calculator = GreeksCalculator(params=greek_params)
    labels = ("PPP-EKI", "PPP-DKI", "NPP-DKI")
    spot_grid = np.array(
        [config.initial_price * float(mult) for mult in config.greek_spot_multipliers],
        dtype=float,
    )
    base_q = float(config.dividend_yield)
    q_grid = np.array(
        [max(0.0, base_q + float(shift)) for shift in config.greek_q_shifts],
        dtype=float,
    )
    greek_frames: dict[str, pd.DataFrame] = {}
    for greek_name in ("delta", "gamma", "vega", "rhoq"):
        rows = []
        for spot in spot_grid:
            row = {"spot": float(spot)}
            env = _clone_env(pricing_env, spot=float(spot))
            for label in labels:
                values = calculator.calculate(
                    products[label],
                    env,
                    engine,
                    greeks=[
                        "delta",
                        "gamma",
                        "vega",
                        "dividend_rho",
                    ],
                )
                key = "dividend_rho" if greek_name == "rhoq" else greek_name
                row[label] = float(values[key])
            rows.append(row)
        greek_frames[greek_name] = pd.DataFrame(rows)

    greek_q_frames: dict[str, pd.DataFrame] = {}
    for greek_name in ("delta", "gamma", "vega", "rhoq"):
        rows = []
        for q_val in q_grid:
            row = {"q": float(q_val)}
            env = _clone_env(
                pricing_env,
                div_yield=ContinuousDividendYield(div_yield=float(q_val)),
            )
            for label in labels:
                values = calculator.calculate(
                    products[label],
                    env,
                    engine,
                    greeks=[
                        "delta",
                        "gamma",
                        "vega",
                        "dividend_rho",
                    ],
                )
                key = "dividend_rho" if greek_name == "rhoq" else greek_name
                row[label] = float(values[key])
            rows.append(row)
        greek_q_frames[greek_name] = pd.DataFrame(rows)

    q_slice_curves: dict[str, dict[str, pd.DataFrame]] = {}
    for slice_name, slice_spot in config.greek_q_slice_spots.items():
        slice_frames: dict[str, pd.DataFrame] = {}
        for greek_name in ("delta", "gamma", "vega", "rhoq"):
            rows = []
            for q_val in q_grid:
                row = {"q": float(q_val)}
                env = _clone_env(
                    pricing_env,
                    spot=float(slice_spot),
                    div_yield=ContinuousDividendYield(div_yield=float(q_val)),
                )
                for label in labels:
                    values = calculator.calculate(
                        products[label],
                        env,
                        engine,
                        greeks=[
                            "delta",
                            "gamma",
                            "vega",
                            "dividend_rho",
                        ],
                    )
                    key = "dividend_rho" if greek_name == "rhoq" else greek_name
                    row[label] = float(values[key])
                rows.append(row)
            slice_frames[greek_name] = pd.DataFrame(rows)
        q_slice_curves[slice_name] = slice_frames

    key_rows = []
    for spot in config.greek_key_spots:
        env = _clone_env(pricing_env, spot=float(spot))
        for label in labels:
            values = calculator.calculate(
                products[label],
                env,
                engine,
                greeks=[
                    "delta",
                    "gamma",
                    "vega",
                    "dividend_rho",
                ],
            )
            key_rows.append(
                {
                    "Spot": float(spot),
                    "Structure": label,
                    "Delta": float(values["delta"]),
                    "Gamma": float(values["gamma"]),
                    "Vega": float(values["vega"]),
                    "RhoQ": float(values["dividend_rho"]),
                }
            )

    key_q_rows = []
    for q_shift in config.greek_key_q_shifts:
        q_val = max(0.0, base_q + float(q_shift))
        env = _clone_env(
            pricing_env,
            div_yield=ContinuousDividendYield(div_yield=float(q_val)),
        )
        for label in labels:
            values = calculator.calculate(
                products[label],
                env,
                engine,
                greeks=[
                    "delta",
                    "gamma",
                    "vega",
                    "dividend_rho",
                ],
            )
            key_q_rows.append(
                {
                    "QShift": float(q_shift),
                    "QLevel": float(q_val),
                    "Structure": label,
                    "Delta": float(values["delta"]),
                    "Gamma": float(values["gamma"]),
                    "Vega": float(values["vega"]),
                    "RhoQ": float(values["dividend_rho"]),
                }
            )

    return GreekCurves(
        spot_grid=spot_grid,
        delta=greek_frames["delta"],
        gamma=greek_frames["gamma"],
        vega=greek_frames["vega"],
        rhoq=greek_frames["rhoq"],
        key_spot_table=pd.DataFrame(key_rows),
        q_grid=q_grid,
        delta_q_curve=greek_q_frames["delta"],
        gamma_q_curve=greek_q_frames["gamma"],
        vega_q_curve=greek_q_frames["vega"],
        rhoq_q_curve=greek_q_frames["rhoq"],
        key_q_table=pd.DataFrame(key_q_rows),
        q_slice_curves=q_slice_curves,
    )


def _save_greek_curve_plot(
    *,
    greek_df: pd.DataFrame,
    x_col: str,
    title: str,
    xlabel: str,
    ylabel: str,
    path: Path,
) -> None:
    plt = _require_matplotlib()
    path.parent.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(8, 4.5))
    for label in ("PPP-EKI", "PPP-DKI", "NPP-DKI"):
        ax.plot(
            greek_df[x_col].to_numpy(),
            greek_df[label].to_numpy(),
            linewidth=2,
            label=label,
        )
    ax.set_title(title)
    ax.set_xlabel(xlabel)
    ax.set_ylabel(ylabel)
    ax.grid(True, alpha=0.3)
    ax.legend()
    fig.tight_layout()
    fig.savefig(path, dpi=150)
    plt.close(fig)


def _save_greek_q_slice_panel(
    *,
    slice_curves: Mapping[str, pd.DataFrame],
    title: str,
    path: Path,
) -> None:
    plt = _require_matplotlib()
    path.parent.mkdir(parents=True, exist_ok=True)
    fig, axes = plt.subplots(2, 2, figsize=(10, 7))
    specs = [
        ("delta", "Delta vs q", "Delta"),
        ("gamma", "Gamma vs q", "Gamma"),
        ("vega", "Vega vs q", "Vega"),
        ("rhoq", "RhoQ vs q", "Dividend Rho (RhoQ)"),
    ]
    for ax, (greek_key, subtitle, ylabel) in zip(axes.flatten(), specs):
        greek_df = slice_curves[greek_key]
        for label in ("PPP-EKI", "PPP-DKI", "NPP-DKI"):
            ax.plot(
                greek_df["q"].to_numpy(),
                greek_df[label].to_numpy(),
                linewidth=2,
                label=label,
            )
        ax.set_title(subtitle)
        ax.set_xlabel("Dividend Yield q")
        ax.set_ylabel(ylabel)
        ax.grid(True, alpha=0.3)
    axes[0, 0].legend()
    fig.suptitle(title)
    fig.tight_layout()
    fig.savefig(path, dpi=150)
    plt.close(fig)


def _compose_english_summary(
    *,
    snapshots: Mapping[str, _PathSnapshot],
    deltas: StructuralDeltas,
) -> str:
    ppp_eki = snapshots["PPP-EKI"].metrics
    ppp_dki = snapshots["PPP-DKI"].metrics
    npp_dki = snapshots["NPP-DKI"].metrics
    return (
        "Under normalized China-style terms, PPP-EKI is lower risk than PPP-DKI on "
        f"raw loss probability ({_fmt_pct(ppp_eki.loss_prob_5)} vs {_fmt_pct(ppp_dki.loss_prob_5)} for loss >5%) "
        f"and NPP-DKI is materially worse on tail loss severity (ES95 {_fmt_num(npp_dki.es95)} vs {_fmt_num(ppp_dki.es95)}). "
        f"The monitoring value Δ_monitoring is {_fmt_num(deltas.delta_monitoring)}, while Δ_protection is {_fmt_num(deltas.delta_protection)}. "
        "PV is quoted ex-principal, consistent with current China index-note convention. "
        "The residual risk in EKI is more terminally concentrated: near the KI barrier at maturity, small terminal moves "
        "can flip the note between the rebate state and the downside state."
    )


def _compose_chinese_summary(
    *,
    snapshots: Mapping[str, _PathSnapshot],
    deltas: StructuralDeltas,
) -> str:
    ppp_eki = snapshots["PPP-EKI"].metrics
    ppp_dki = snapshots["PPP-DKI"].metrics
    npp_dki = snapshots["NPP-DKI"].metrics
    return (
        "在归一化的中资雪球条款下，部分保本欧式敲入雪球在本金亏损概率上低于部分保本日度敲入雪球，"
        f"其中亏损超过5%的概率分别为{_fmt_pct(ppp_eki.loss_prob_5)}与{_fmt_pct(ppp_dki.loss_prob_5)}；"
        f"非保本日度敲入雪球的尾部损失显著更差，ES95分别为{_fmt_num(npp_dki.es95)}与{_fmt_num(ppp_dki.es95)}。"
        f"监控方式增量价值 Δ_monitoring 为{_fmt_num(deltas.delta_monitoring)}，保本增量价值 Δ_protection 为{_fmt_num(deltas.delta_protection)}。"
        "PV 采用当前中国指数雪球常见的不含本金口径。欧式敲入并非在所有维度都更安全，其剩余风险更集中于终值附近，尤其是在到期前临近敲入价时，终值的小幅波动即可触发状态切换。"
    )


def _add_report_sections(
    *,
    document: Document,
    language: str,
    config: SnowballRiskComparisonConfig,
    pricing_env: PricingEnvironment,
    products: Mapping[str, SnowballOption],
    snapshots: Mapping[str, _PathSnapshot],
    greek_curves: GreekCurves,
    deltas: StructuralDeltas,
    deterministic_cases: Sequence[DeterministicCaseResult],
    plot_paths: Mapping[str, Path],
) -> None:
    is_english = language == "en"
    title = (
        "Snowball Monitoring / Protection / Parachute Risk Comparison"
        if is_english
        else "雪球监控方式 / 保本 / 降落伞特征风险对比报告"
    )
    subtitle = (
        "Normalized China-style structures, English first."
        if is_english
        else "归一化中资雪球结构，对应英文版镜像内容。"
    )
    _add_heading(document, title, level=0)
    _add_paragraph(document, subtitle)

    _add_heading(
        document,
        "Executive Conclusion" if is_english else "执行摘要",
        level=1,
    )
    _add_paragraph(
        document,
        _compose_english_summary(snapshots=snapshots, deltas=deltas)
        if is_english
        else _compose_chinese_summary(snapshots=snapshots, deltas=deltas),
    )

    _add_heading(
        document,
        "Structure Normalization and Methodology" if is_english else "结构归一与方法说明",
        level=1,
    )
    methodology = (
        f"Baseline assumptions: S0={config.initial_price:.0f}, strike={config.strike:.0f}, "
        f"KI={config.ki_barrier:.0f}, KO={config.ko_barrier:.0f}, coupon={_fmt_pct(config.annual_coupon)}, "
        f"r={_fmt_pct(config.rate)}, q={_fmt_pct(config.dividend_yield)}, vol={_fmt_pct(config.volatility)}, "
        f"tenor={config.tenor_months}M, business-day count={config.business_days_in_year}. "
        "Dividend yield q is calibrated to the current China equity-index convention and analyzed over the practical band [8.0000%, 15.0000%]. "
        "Daily-KI uses exact business-day discrete observations; EKI uses final-only KI. "
        "PV is quoted ex-principal and state probabilities come from common-random-number Monte Carlo on one shared observation grid."
        if is_english
        else f"基准假设为：S0={config.initial_price:.0f}，行权价={config.strike:.0f}，敲入价={config.ki_barrier:.0f}，"
        f"敲出价={config.ko_barrier:.0f}，票息={_fmt_pct(config.annual_coupon)}，r={_fmt_pct(config.rate)}，"
        f"q={_fmt_pct(config.dividend_yield)}，波动率={_fmt_pct(config.volatility)}，期限={config.tenor_months}个月，"
        f"年交易日={config.business_days_in_year}。股息率 q 采用当前中国股票指数常见口径，并在 [8.0000%, 15.0000%] 区间内分析。"
        "日度敲入采用精确离散交易日观察，欧式敲入仅在终值观察。"
        "PV 采用不含本金口径，状态概率基于统一观测网格与共同随机数的蒙特卡洛路径。"
    )
    _add_paragraph(document, methodology)
    _add_dataframe_table(
        document,
        _build_contract_terms_table(
            config=config,
            products=products,
            language=language,
        ),
        title="Contract Terms Comparison" if is_english else "合约条款对比表",
    )

    _add_heading(
        document,
        "Matched-Structure Pricing Comparison" if is_english else "匹配结构定价对比",
        level=1,
    )
    _add_dataframe_table(document, _build_metrics_table(snapshots))
    _add_dataframe_table(document, _build_delta_table(deltas))

    _add_heading(
        document,
        "Greeks Comparison" if is_english else "Greeks 对比",
        level=1,
    )
    _add_paragraph(
        document,
        (
            "This section compares how delta, gamma, vega, and rhoq evolve as spot moves through the KI region, the rebound band, and the KO region. "
            "It now includes both spot sweeps and dividend-yield (q) sweeps, plus dedicated q panels at the KI and KO barrier neighborhoods. PPP-EKI should usually show less path-driven downside sensitivity than PPP-DKI after rebound states, while NPP-DKI should retain the same directional shape but with worse downside convexity."
            if is_english
            else "本节比较 delta、gamma、vega 与 rhoq 随现货穿越 KI 区间、反弹区间与 KO 区间时的变化，"
            "并新增了随股息率 q 变化的敏感度曲线，以及贴近 KI / KO 障碍位置的 q 切片图。一般而言，PPP-EKI 在反弹状态后的路径依赖型下行敏感度会弱于 PPP-DKI，而 NPP-DKI 的方向形态类似，但下行凸性风险更差。"
        ),
    )
    _add_dataframe_table(
        document,
        greek_curves.key_spot_table,
        title="Key Spot Greek Levels" if is_english else "关键点位 Greek 对比",
    )
    _add_dataframe_table(
        document,
        greek_curves.key_q_table,
        title="Key Dividend Yield Greek Levels"
        if is_english
        else "关键 q 水平 Greek 对比",
    )
    for key in (
        "greek_delta",
        "greek_gamma",
        "greek_vega",
        "greek_rhoq",
        "greek_delta_q",
        "greek_gamma_q",
        "greek_vega_q",
        "greek_rhoq_q",
    ):
        _add_picture_if_exists(document, plot_paths.get(key))
    _add_paragraph(
        document,
        "Near-KI and Near-KO q slices" if is_english else "近 KI 与近 KO 的 q 切片",
        bold=True,
    )
    for key in ("greek_q_near_ki", "greek_q_near_ko"):
        _add_picture_if_exists(document, plot_paths.get(key))

    _add_heading(
        document,
        "Path Template Evidence" if is_english else "路径模板证据",
        level=1,
    )
    _add_paragraph(
        document,
        "Deterministic path tests show where monitoring rules matter most. "
        "The 100→65→85 case is the clean rebound band where PPP-EKI avoids the loss regime that still applies to PPP-DKI. "
        "The 100→65→74 case isolates parachute rescue value and shows why DKI benefits more than EKI."
        if is_english
        else "确定性路径测试展示了监控规则最有差异的区间。100→65→85 对应典型反弹带，在该区间中 PPP-EKI 通常可以避免 PPP-DKI 仍会触发的亏损状态；"
        "100→65→74 则单独刻画了降落伞的救援价值，并显示其对 DKI 的提升明显大于对 EKI 的提升。",
    )
    _add_dataframe_table(document, _build_deterministic_table(deterministic_cases))

    _add_heading(
        document,
        "Event Probability and Cashflow Decomposition"
        if is_english
        else "事件概率与现金流分解",
        level=1,
    )
    _add_dataframe_table(
        document,
        snapshots["PPP-EKI"].metrics.conditional_ko_cashflows,
        title="PPP-EKI" if is_english else "PPP-EKI 现金流分解",
    )
    _add_dataframe_table(
        document,
        snapshots["PPP-DKI"].metrics.conditional_ko_cashflows,
        title="PPP-DKI" if is_english else "PPP-DKI 现金流分解",
    )

    _add_heading(
        document,
        "Loss Tail and Expected Shortfall" if is_english else "损失尾部与期望短缺",
        level=1,
    )
    _add_dataframe_table(document, _build_band_table(snapshots))

    _add_heading(
        document,
        "Barrier / Terminal-Cliff Concentration"
        if is_english
        else "障碍附近 / 到期悬崖风险集中",
        level=1,
    )
    ko_barrier, ko_pct, ko_sigma = _barrier_distance_metrics(
        spot=pricing_env.spot,
        barrier=config.ko_barrier,
        time_to_barrier=float(snapshots["PPP-EKI"].ko_times[0]),
        pricing_env=pricing_env,
        product=products["PPP-EKI"],
    )
    _add_paragraph(
        document,
        (
            f"Initial KO barrier watch: level={ko_barrier:.2f}, pct distance={_fmt_pct(ko_pct)}, sigma distance={ko_sigma:.3f}. "
            f"Within the terminal cliff band (±{_fmt_pct(config.terminal_cliff_band)} around KI), "
            f"PPP-EKI loss probability is {_fmt_pct(snapshots['PPP-EKI'].metrics.terminal_cliff_loss_probability)} "
            f"versus {_fmt_pct(snapshots['PPP-DKI'].metrics.terminal_cliff_loss_probability)} for PPP-DKI."
            if is_english
            else f"初始敲出障碍监控：障碍水平={ko_barrier:.2f}，百分比距离={_fmt_pct(ko_pct)}，标准差距离={ko_sigma:.3f}。"
            f"在终值位于 KI 附近 ±{_fmt_pct(config.terminal_cliff_band)} 的悬崖区间内，PPP-EKI 的亏损概率为"
            f"{_fmt_pct(snapshots['PPP-EKI'].metrics.terminal_cliff_loss_probability)}，PPP-DKI 为"
            f"{_fmt_pct(snapshots['PPP-DKI'].metrics.terminal_cliff_loss_probability)}。"
        ),
    )
    _add_picture_if_exists(document, plot_paths.get("terminal_cliff"))

    _add_heading(
        document,
        "Parachute Feature Analysis" if is_english else "降落伞特征分析",
        level=1,
    )
    _add_paragraph(
        document,
        (
            f"Δ_parachute_DKI = {_fmt_num(deltas.delta_parachute_dki)} and "
            f"Δ_parachute_EKI = {_fmt_num(deltas.delta_parachute_eki)}. "
            "The MC evidence and deterministic 100→65→74 path both support the same conclusion: "
            "parachute mostly rescues paths that had already knocked in under daily monitoring."
            if is_english
            else f"Δ_parachute_DKI = {_fmt_num(deltas.delta_parachute_dki)}，"
            f"Δ_parachute_EKI = {_fmt_num(deltas.delta_parachute_eki)}。"
            "蒙特卡洛结果与 100→65→74 的确定性路径结论一致：降落伞的主要价值在于挽救那些在日度监控下已提前敲入、但终值又回到 KI 以上的路径。"
        )
    )

    _add_heading(document, "Stress Section" if is_english else "压力测试部分", level=1)
    _add_paragraph(
        document,
        "The stress layer uses fast repricing to show how the same structures redistribute risk under spot, vol, dividend, skew/smile, and combined trader ladders."
        if is_english
        else "压力测试层通过快速重定价展示相同结构在现货、波动率、股息、偏斜/微笑以及组合交易员阶梯情景下如何重新分配风险。",
    )
    for key in (
        "spot_stress",
        "vol_stress",
        "div_stress",
        "skew_smile",
        "delta_monitoring",
        "delta_protection",
    ):
        _add_picture_if_exists(document, plot_paths.get(key))

    _add_heading(
        document,
        "Limitations and Interpretation" if is_english else "局限性与解读边界",
        level=1,
    )
    _add_paragraph(
        document,
        "This report is a normalized matched-structure study. It does not model issuer spread, secondary-market liquidity, or a dedicated local-vol / jump-diffusion snowball process. "
        "The purpose is to isolate the marginal contribution of monitoring, protection, and parachute design under exact discrete barrier rules."
        if is_english
        else "本报告是归一化的匹配结构研究，不包含发行人信用利差、二级市场流动性，也未单独实现局部波动率或跳扩散雪球模型。"
        "其目标是基于精确离散障碍规则，分离监控方式、保本设置与降落伞设计的边际贡献。",
    )

    _add_heading(
        document,
        "Appendix: Terminology" if is_english else "附录：术语说明",
        level=1,
    )
    _add_paragraph(
        document,
        "Definitions for the key terms used throughout the report."
        if is_english
        else "以下列示报告中反复出现的核心术语及其定义。",
    )
    _add_dataframe_table(document, _build_terminology_table(language))


def _build_docx_report(
    *,
    config: SnowballRiskComparisonConfig,
    pricing_env: PricingEnvironment,
    products: Mapping[str, SnowballOption],
    snapshots: Mapping[str, _PathSnapshot],
    greek_curves: GreekCurves,
    deltas: StructuralDeltas,
    deterministic_cases: Sequence[DeterministicCaseResult],
    plot_paths: Mapping[str, Path],
    output_dir: Path,
) -> Path:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    _add_report_sections(
        document=document,
        language="en",
        config=config,
        pricing_env=pricing_env,
        products=products,
        snapshots=snapshots,
        greek_curves=greek_curves,
        deltas=deltas,
        deterministic_cases=deterministic_cases,
        plot_paths=plot_paths,
    )
    document.add_section(WD_SECTION.NEW_PAGE)
    _add_report_sections(
        document=document,
        language="zh",
        config=config,
        pricing_env=pricing_env,
        products=products,
        snapshots=snapshots,
        greek_curves=greek_curves,
        deltas=deltas,
        deterministic_cases=deterministic_cases,
        plot_paths=plot_paths,
    )
    output_dir.mkdir(parents=True, exist_ok=True)
    report_path = output_dir / config.report_filename
    document.save(str(report_path))
    return report_path


def _load_input_module(path: str):
    spec = importlib.util.spec_from_file_location("snowball_risk_input", path)
    if spec is None or spec.loader is None:
        raise ValidationError(f"Unable to load input module: {path}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def generate_snowball_risk_comparison_report(
    config: Optional[SnowballRiskComparisonConfig] = None,
) -> SnowballRiskComparisonArtifacts:
    if config is None:
        config = build_default_snowball_risk_comparison_config()
    if config.bilingual_layout != "english_then_chinese":
        raise ValidationError("Only english_then_chinese bilingual layout is supported.")

    output_dir = Path(config.output_dir)
    (
        calendar,
        _ko_dates,
        ko_times,
        _daily_ki_dates,
        daily_ki_times,
        _maturity_date,
        maturity_time,
    ) = _build_ko_schedule(config)
    pricing_env = _build_pricing_environment(config, calendar)
    products = _build_structures(config, ko_times, daily_ki_times, maturity_time)
    observation_times, _paths, snapshots = _simulate_common_paths(
        config=config,
        pricing_env=pricing_env,
        products=products,
    )
    deterministic_cases = _run_deterministic_cases(
        products=products,
        pricing_env=pricing_env,
        observation_times=observation_times,
    )
    greek_curves = _compute_greek_curves(
        config=config,
        pricing_env=pricing_env,
        products=products,
    )

    deltas = StructuralDeltas(
        delta_monitoring=snapshots["PPP-EKI"].metrics.pv
        - snapshots["PPP-DKI"].metrics.pv,
        delta_protection=snapshots["PPP-DKI"].metrics.pv
        - snapshots["NPP-DKI"].metrics.pv,
        delta_parachute_dki=snapshots["PPP-DKI-Parachute"].metrics.pv
        - snapshots["PPP-DKI"].metrics.pv,
        delta_parachute_eki=snapshots["PPP-EKI-Parachute"].metrics.pv
        - snapshots["PPP-EKI"].metrics.pv,
    )
    plot_paths = _generate_plot_artifacts(
        config=config,
        products=products,
        pricing_env=pricing_env,
        snapshots=snapshots,
        greek_curves=greek_curves,
        output_dir=output_dir,
    )
    report_path = _build_docx_report(
        config=config,
        pricing_env=pricing_env,
        products=products,
        snapshots=snapshots,
        greek_curves=greek_curves,
        deltas=deltas,
        deterministic_cases=deterministic_cases,
        plot_paths=plot_paths,
        output_dir=output_dir,
    )
    return SnowballRiskComparisonArtifacts(
        report_path=report_path,
        output_dir=output_dir,
        plot_paths=plot_paths,
    )


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate a bilingual DOCX risk comparison report for snowball structures."
    )
    parser.add_argument("--input", type=str, help="Optional Python input module path.")
    parser.add_argument("--out", type=str, help="Override output directory.")
    args = parser.parse_args()

    if args.input:
        module = _load_input_module(args.input)
        if hasattr(module, "build_config"):
            config = module.build_config()
        elif hasattr(module, "config"):
            config = module.config
        else:
            raise ValidationError("Input module must define build_config() or config.")
    else:
        config = build_default_snowball_risk_comparison_config()

    if args.out:
        config = SnowballRiskComparisonConfig(
            **{**config.__dict__, "output_dir": Path(args.out)}
        )

    result = generate_snowball_risk_comparison_report(config)
    print(f"Generated report: {result.report_path}")


if __name__ == "__main__":
    main()
